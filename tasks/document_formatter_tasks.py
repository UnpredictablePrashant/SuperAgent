from __future__ import annotations

import os
import re
import time
import uuid
from pathlib import Path
from typing import Any

from tasks.a2a_agent_utils import begin_agent_session, publish_agent_output
from tasks.utils import OUTPUT_DIR, log_task_update, write_text_file, resolve_output_path

AGENT_METADATA = {
    "document_formatter_agent": {
        "description": (
            "Takes any text/markdown content and produces a properly formatted document "
            "with downloadable exports: Markdown (.md), PDF (.pdf), and Word (.docx). "
            "Applies professional structure: title, table of contents, headings, tables, "
            "and consistent styling. Runs as a final pass after any report-generating workflow."
        ),
        "skills": ["formatting", "markdown", "pdf", "docx", "document"],
        "input_keys": [
            "current_objective",
            "draft_response",
            "final_output",
            "document_formatter_content",
            "document_formatter_title",
        ],
        "output_keys": [
            "document_formatter_md_path",
            "document_formatter_pdf_path",
            "document_formatter_docx_path",
            "document_formatter_title",
            "long_document_compiled_path",
            "long_document_compiled_pdf_path",
            "long_document_compiled_docx_path",
        ],
        "requirements": ["python-docx"],
    }
}


def _slugify(text: str, max_len: int = 40) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")
    return slug[:max_len] or "document"


def _extract_title(content: str, objective: str) -> str:
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
    if objective:
        first_sentence = objective.split(".")[0].strip()
        if first_sentence:
            return first_sentence[:80]
    return "Document"


def _ensure_proper_markdown(content: str, title: str) -> str:
    """Ensure the markdown has a clean title header and consistent structure."""
    lines = content.splitlines()
    has_h1 = any(l.strip().startswith("# ") for l in lines[:5])
    if not has_h1:
        content = f"# {title}\n\n{content}"
    content = re.sub(r"\n{4,}", "\n\n\n", content)
    return content.strip()


def _markdown_to_styled_html(markdown_text: str, title: str) -> str:
    css = """
    <style>
      body { font-family: Georgia, serif; max-width: 860px; margin: 0 auto; padding: 40px 50px;
             color: #222; line-height: 1.7; }
      h1 { font-size: 2em; border-bottom: 3px solid #333; padding-bottom: 8px; margin-top: 0; }
      h2 { font-size: 1.5em; border-bottom: 1px solid #aaa; padding-bottom: 4px; margin-top: 2em; color: #2a5298; }
      h3 { font-size: 1.2em; color: #333; margin-top: 1.5em; }
      h4 { font-size: 1.05em; color: #555; }
      p { margin: 0.8em 0; }
      ul, ol { margin: 0.6em 0 0.6em 2em; }
      li { margin: 0.3em 0; }
      code { background: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-size: 0.9em; }
      pre { background: #f4f4f4; padding: 14px; border-radius: 5px; overflow-x: auto; }
      pre code { background: none; padding: 0; }
      blockquote { border-left: 4px solid #2a5298; margin: 1em 0; padding: 0.5em 1em; background: #f0f4ff; }
      table { border-collapse: collapse; width: 100%; margin: 1em 0; }
      th { background: #2a5298; color: white; padding: 8px 12px; text-align: left; }
      td { padding: 7px 12px; border: 1px solid #ddd; }
      tr:nth-child(even) td { background: #f7f9ff; }
      strong { color: #111; }
      a { color: #2a5298; }
      .cover { text-align: center; padding: 60px 0; border-bottom: 2px solid #333; margin-bottom: 40px; }
      .cover h1 { border: none; font-size: 2.5em; }
      .cover .date { color: #666; margin-top: 10px; }
    </style>
    """
    from datetime import date
    date_str = date.today().strftime("%B %d, %Y")
    cover = f'<div class="cover"><h1>{title}</h1><div class="date">{date_str}</div></div>\n'
    body = _md_to_html_body(markdown_text)
    return f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{title}</title>{css}</head><body>{cover}{body}</body></html>"


def _md_to_html_body(md: str) -> str:
    lines = md.splitlines()
    out: list[str] = []
    in_code = False
    code_buf: list[str] = []
    in_ul = False
    in_ol = False
    in_table = False
    table_buf: list[str] = []

    def flush_lists() -> None:
        nonlocal in_ul, in_ol
        if in_ul:
            out.append("</ul>")
            in_ul = False
        if in_ol:
            out.append("</ol>")
            in_ol = False

    def flush_table() -> None:
        nonlocal in_table, table_buf
        if not in_table:
            return
        rows = table_buf[:]
        table_buf.clear()
        in_table = False
        if not rows:
            return
        out.append("<table>")
        header_cells = [c.strip() for c in rows[0].strip("|").split("|")]
        out.append("<tr>" + "".join(f"<th>{_inline(c)}</th>" for c in header_cells) + "</tr>")
        for row in rows[2:]:
            cells = [c.strip() for c in row.strip("|").split("|")]
            out.append("<tr>" + "".join(f"<td>{_inline(c)}</td>" for c in cells) + "</tr>")
        out.append("</table>")

    def _inline(text: str) -> str:
        text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
        text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
        text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
        text = re.sub(r"\[(.+?)\]\((.+?)\)", r'<a href="\2">\1</a>', text)
        return text

    for line in lines:
        raw = line

        if raw.strip().startswith("```"):
            if not in_code:
                flush_lists()
                flush_table()
                in_code = True
                out.append('<pre><code>')
            else:
                in_code = False
                out.append("</code></pre>")
            continue
        if in_code:
            import html as _html_mod
            out.append(_html_mod.escape(raw))
            continue

        if "|" in raw and raw.strip().startswith("|"):
            flush_lists()
            if not in_table:
                in_table = True
                table_buf.clear()
            table_buf.append(raw)
            continue
        if in_table:
            flush_table()

        stripped = raw.strip()
        m_h = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m_h:
            flush_lists()
            level = len(m_h.group(1))
            out.append(f"<h{level}>{_inline(m_h.group(2))}</h{level}>")
            continue

        if re.match(r"^[-*]\s+", stripped):
            flush_table()
            if not in_ul:
                if in_ol:
                    out.append("</ol>")
                    in_ol = False
                out.append("<ul>")
                in_ul = True
            ul_text = re.sub(r"^[-*]\s+", "", stripped)
            out.append("<li>" + _inline(ul_text) + "</li>")
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_table()
            if not in_ol:
                if in_ul:
                    out.append("</ul>")
                    in_ul = False
                out.append("<ol>")
                in_ol = True
            ol_text = re.sub(r"^\d+\.\s+", "", stripped)
            out.append("<li>" + _inline(ol_text) + "</li>")
            continue

        if re.match(r"^>", stripped):
            flush_lists()
            out.append(f"<blockquote>{_inline(stripped[1:].strip())}</blockquote>")
            continue

        flush_lists()
        flush_table()
        if stripped:
            out.append(f"<p>{_inline(stripped)}</p>")
        else:
            out.append("")

    flush_lists()
    flush_table()
    return "\n".join(out)


def _export_to_docx(content: str, title: str, out_path: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        title_para = doc.add_paragraph()
        title_para.style = doc.styles["Title"]
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x22, 0x55, 0x99)

        from datetime import date
        date_para = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
        date_para.paragraph_format.space_after = Pt(24)

        in_code = False
        code_buf: list[str] = []
        table_buf: list[str] = []
        in_table = False

        def flush_code() -> None:
            nonlocal in_code, code_buf
            if code_buf:
                p = doc.add_paragraph("\n".join(code_buf))
                p.style = doc.styles.get("Code") or doc.styles["Normal"]
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                code_buf.clear()
            in_code = False

        def flush_table_buf() -> None:
            nonlocal in_table, table_buf
            if not in_table or not table_buf:
                in_table = False
                table_buf.clear()
                return
            header = [c.strip() for c in table_buf[0].strip("|").split("|")]
            data_rows = []
            for row in table_buf[2:]:
                if re.match(r"^\s*\|?\s*[-:]+\s*\|", row):
                    continue
                data_rows.append([c.strip() for c in row.strip("|").split("|")])
            if not header:
                in_table = False
                table_buf.clear()
                return
            num_cols = len(header)
            tbl = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
            tbl.style = "Table Grid"
            hdr_cells = tbl.rows[0].cells
            for i, h in enumerate(header[:num_cols]):
                hdr_cells[i].text = h
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                tc = hdr_cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "225599")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)
            for ri, row_data in enumerate(data_rows):
                row_cells = tbl.rows[ri + 1].cells
                for ci, cell_val in enumerate(row_data[:num_cols]):
                    row_cells[ci].text = cell_val
            in_table = False
            table_buf.clear()

        lines = content.splitlines()
        for line in lines:
            raw = line

            if raw.strip().startswith("```"):
                if not in_code:
                    flush_table_buf()
                    in_code = True
                    code_buf.clear()
                else:
                    flush_code()
                continue
            if in_code:
                code_buf.append(raw)
                continue

            if "|" in raw and raw.strip().startswith("|"):
                if not in_table:
                    in_table = True
                    table_buf.clear()
                table_buf.append(raw)
                continue
            if in_table:
                flush_table_buf()

            stripped = raw.strip()
            m_h = re.match(r"^(#{1,6})\s+(.*)", stripped)
            if m_h:
                level = len(m_h.group(1))
                heading_style = f"Heading {min(level, 4)}"
                try:
                    p = doc.add_paragraph(m_h.group(2), style=heading_style)
                except Exception:
                    p = doc.add_paragraph(m_h.group(2))
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(max(20 - (level - 1) * 2, 12))
                continue

            if re.match(r"^[-*]\s+", stripped):
                text = re.sub(r"^[-*]\s+", "", stripped)
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
                text = re.sub(r"\*(.+?)\*", r"\1", text)
                text = re.sub(r"`(.+?)`", r"\1", text)
                p = doc.add_paragraph(text, style="List Bullet")
                continue

            if re.match(r"^\d+\.\s+", stripped):
                text = re.sub(r"^\d+\.\s+", "", stripped)
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
                p = doc.add_paragraph(text, style="List Number")
                continue

            if stripped:
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
                text = re.sub(r"\*(.+?)\*", r"\1", text)
                text = re.sub(r"`(.+?)`", r"\1", text)
                p = doc.add_paragraph(text)
            else:
                doc.add_paragraph("")

        flush_table_buf()
        if in_code:
            flush_code()

        doc.save(out_path)
        return True
    except Exception as exc:
        log_task_update("DocFormatter", f"DOCX export failed: {exc}")
        return False


def _export_to_pdf(html_text: str, pdf_path: str) -> bool:
    try:
        import weasyprint
        weasyprint.HTML(string=html_text).write_pdf(pdf_path)
        return True
    except Exception:
        pass
    try:
        from tasks.long_document_tasks import _markdown_to_plain_text, _render_pdf_bytes
        plain = _markdown_to_plain_text(html_text)
        Path(pdf_path).write_bytes(_render_pdf_bytes(plain))
        return True
    except Exception as exc:
        log_task_update("DocFormatter", f"PDF export failed: {exc}")
        return False


def document_formatter_agent(state: dict) -> dict:
    _, task_content, _ = begin_agent_session(state, "document_formatter_agent")
    log_task_update("DocFormatter", "Formatting document and exporting MD / PDF / DOCX …")

    content = str(
        state.get("document_formatter_content")
        or state.get("draft_response")
        or state.get("final_output")
        or task_content
        or ""
    ).strip()

    if not content:
        log_task_update("DocFormatter", "No content to format.")
        state["final_output"] = "No content was available to format."
        return publish_agent_output(state, "document_formatter_agent", state["final_output"])

    objective = str(state.get("current_objective") or state.get("user_query") or "").strip()
    title = str(state.get("document_formatter_title") or _extract_title(content, objective) or "Document")

    clean_md = _ensure_proper_markdown(content, title)

    slug = _slugify(title)
    ts = int(time.time())
    uid = str(uuid.uuid4())[:8]
    artifact_dir = f"doc_fmt_{slug}_{ts}_{uid}"
    artifact_path = Path(resolve_output_path(artifact_dir))
    artifact_path.mkdir(parents=True, exist_ok=True)

    md_path = artifact_path / f"{slug}.md"
    md_path.write_text(clean_md, encoding="utf-8")

    html_text = _markdown_to_styled_html(clean_md, title)
    html_path = artifact_path / f"{slug}.html"
    html_path.write_text(html_text, encoding="utf-8")

    docx_path = str(artifact_path / f"{slug}.docx")
    docx_ok = _export_to_docx(clean_md, title, docx_path)

    pdf_path = str(artifact_path / f"{slug}.pdf")
    pdf_ok = _export_to_pdf(html_text, pdf_path)

    state["document_formatter_md_path"] = str(md_path)
    state["document_formatter_pdf_path"] = pdf_path if pdf_ok else ""
    state["document_formatter_docx_path"] = docx_path if docx_ok else ""
    state["document_formatter_title"] = title

    state["long_document_compiled_path"] = str(md_path)
    state["long_document_compiled_pdf_path"] = pdf_path if pdf_ok else ""
    state["long_document_compiled_docx_path"] = docx_path if docx_ok else ""

    lines = [
        f"Document formatted successfully.",
        f"Title: {title}",
        f"Markdown: {md_path}",
    ]
    if docx_ok:
        lines.append(f"DOCX: {docx_path}")
    if pdf_ok:
        lines.append(f"PDF: {pdf_path}")
    lines.append("")
    lines.append(clean_md[:4000])
    summary = "\n".join(lines)

    state["final_output"] = summary
    state["draft_response"] = summary

    log_task_update("DocFormatter", f"Done — MD + {'DOCX ' if docx_ok else ''}{'PDF' if pdf_ok else ''} exported.")
    return publish_agent_output(state, "document_formatter_agent", summary)
