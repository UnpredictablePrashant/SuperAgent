from __future__ import annotations

import datetime as dt
import json
import math
import os
import re
import shutil
import subprocess
import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any, Mapping
from urllib.parse import urljoin, urlparse
import textwrap

from openai import OpenAI

from kendr.domain.deep_research import build_source_strategy, discover_research_intent
from kendr.domain.local_drive import (
    normalize_extension_set as _normalize_extension_set,
    resolve_paths as _resolve_local_drive_paths,
    scan_local_drive_tree as _scan_local_drive_tree,
)
from kendr.execution_trace import append_execution_event
from kendr.llm_router import supports_native_web_search
from kendr.rag_manager import build_research_grounding
from kendr.workflow_contract import approval_request_to_text, build_approval_request
from tasks.a2a_agent_utils import begin_agent_session, publish_agent_output
from tasks.coding_tasks import _extract_output_text
from tasks.file_memory import bootstrap_file_memory, update_planning_file
from tasks.planning_tasks import build_plan_approval_prompt, normalize_plan_data, plan_as_markdown
from tasks.plagiarism_checker import build_plagiarism_report as _dedicated_build_plagiarism_report
from tasks.research_infra import (
    fetch_search_results,
    fetch_url_content,
    llm_json,
    llm_text,
    normalize_research_search_backend,
    openai_analyze_image,
    parse_documents,
    strip_code_fences,
)
from tasks.research_output import render_artifact_lines
from tasks.utils import OUTPUT_DIR, log_task_update, model_selection_for_agent, write_text_file, resolve_output_path


DEFAULT_DEEP_RESEARCH_MODEL = os.getenv("KENDR_DEEP_RESEARCH_MODEL", os.getenv("OPENAI_DEEP_RESEARCH_MODEL", "o4-mini-deep-research"))
DEFAULT_RESEARCH_FORMATS = ["pdf", "docx", "html", "md"]
SUPPORTED_CITATION_STYLES = {"apa", "mla", "chicago", "ieee", "vancouver", "harvard"}
DEEP_RESEARCH_LABEL = "Deep Research"
OPENAI_WEB_SEARCH_TOOL = os.getenv("OPENAI_WEB_SEARCH_TOOL", "web_search").strip() or "web_search"
DEEP_RESEARCH_DEPTH_PRESETS = {
    "brief": {
        "pages": 10,
        "label": "Focused Brief",
        "summary": "Focused",
        "description": "Tight synthesis of the most important findings with a narrower evidence sweep.",
    },
    "standard": {
        "pages": 25,
        "label": "Standard Report",
        "summary": "Standard",
        "description": "Balanced multi-section research with enough room for evidence, tradeoffs, and conclusions.",
    },
    "comprehensive": {
        "pages": 50,
        "label": "Comprehensive Study",
        "summary": "Comprehensive",
        "description": "Broader source coverage and deeper cross-section synthesis for more complex topics.",
    },
    "exhaustive": {
        "pages": 100,
        "label": "Exhaustive Dossier",
        "summary": "Exhaustive",
        "description": "Maximum depth and breadth for heavy research, longer runtimes, and large evidence sets.",
    },
}


def _normalize_research_depth_mode(value: Any, target_pages: int) -> str:
    normalized = str(value or "").strip().lower().replace("-", "_").replace(" ", "_")
    if normalized in DEEP_RESEARCH_DEPTH_PRESETS:
        return normalized
    if target_pages >= 100:
        return "exhaustive"
    if target_pages >= 50:
        return "comprehensive"
    if target_pages >= 20:
        return "standard"
    return "brief"


def _research_depth_config(value: Any, target_pages: int) -> dict[str, Any]:
    mode = _normalize_research_depth_mode(value, target_pages)
    config = dict(DEEP_RESEARCH_DEPTH_PRESETS.get(mode, DEEP_RESEARCH_DEPTH_PRESETS["standard"]))
    config["mode"] = mode
    return config


def _trace_now() -> str:
    return dt.datetime.now(dt.timezone.utc).isoformat()


def _trace_research_event(
    state: dict,
    *,
    title: str,
    detail: str = "",
    command: str = "",
    status: str = "running",
    kind: str = "research_activity",
    started_at: str = "",
    completed_at: str = "",
    metadata: dict[str, Any] | None = None,
    subtask: str = "",
) -> dict[str, Any]:
    return append_execution_event(
        state,
        kind=kind,
        actor="long_document_agent",
        status=status,
        title=title,
        detail=detail,
        command=command,
        started_at=started_at,
        completed_at=completed_at,
        metadata=metadata or {},
        persist=True,
        active_agent="long_document_agent",
        task=str(state.get("current_objective") or state.get("user_query") or "").strip(),
        subtask=subtask,
    )


def _trace_url_list(urls: list[str] | None, *, limit: int = 6) -> list[str]:
    compact: list[str] = []
    for raw in urls or []:
        value = str(raw or "").strip()
        if not value or value in compact:
            continue
        compact.append(value)
        if len(compact) >= limit:
            break
    return compact


def _cancel_requested(state: dict) -> bool:
    if bool(state.get("user_cancelled", False)):
        return True
    path_value = str(state.get("kill_switch_file", "")).strip()
    if not path_value:
        policy = state.get("privileged_execution_policy", {})
        if isinstance(policy, dict):
            path_value = str(policy.get("kill_switch_file", "")).strip()
    return bool(path_value) and Path(path_value).exists()


def _raise_if_cancelled(state: dict, *, phase: str = "") -> None:
    if not _cancel_requested(state):
        return
    state["user_cancelled"] = True
    _trace_research_event(
        state,
        title="Deep research stop requested",
        detail=(
            f"Stop requested during {phase}. Ending the deep research workflow."
            if phase
            else "Stop requested. Ending the deep research workflow."
        ),
        status="cancelled",
        metadata={"phase": phase or "runtime", "cancelled": True},
        subtask="Stop deep research workflow",
    )
    raise RuntimeError("Kill switch triggered. Refusing further execution.")

AGENT_METADATA = {
    "long_document_agent": {
        "description": (
            "Builds deep research reports through tiered planning, source-backed evidence gathering, "
            "cross-section correlation, plagiarism checks, and multi-format export."
        ),
        "skills": ["deep-research", "reporting", "chaptering", "synthesis", "citations", "plagiarism"],
        "input_keys": [
            "current_objective",
            "deep_research_mode",
            "long_document_mode",
            "long_document_pages",
            "long_document_sections",
            "long_document_section_pages",
            "long_document_title",
            "long_document_collect_sources_first",
            "long_document_disable_visuals",
            "long_document_section_references",
            "long_document_section_search",
            "long_document_section_search_results",
            "research_depth_mode",
            "research_model",
            "research_instructions",
            "research_max_tool_calls",
            "research_max_output_tokens",
            "research_poll_interval_seconds",
            "research_max_wait_seconds",
            "research_heartbeat_seconds",
            "research_output_formats",
            "research_citation_style",
            "research_enable_plagiarism_check",
            "research_web_search_enabled",
            "research_date_range",
            "research_max_sources",
            "research_checkpoint_enabled",
            "research_kb_enabled",
            "research_kb_id",
            "research_kb_top_k",
            "deep_research_source_urls",
        ],
        "output_keys": [
            "deep_research_analysis",
            "deep_research_intent",
            "deep_research_source_strategy",
            "deep_research_coverage_report",
            "deep_research_evidence_ledger",
            "deep_research_quality_report",
            "deep_research_artifacts_manifest",
            "deep_research_result_card",
            "long_document_title",
            "long_document_outline",
            "long_document_sections_data",
            "long_document_artifact_dir",
            "long_document_compiled_path",
            "long_document_compiled_html_path",
            "long_document_compiled_docx_path",
            "long_document_compiled_pdf_path",
            "long_document_outline_md_path",
            "long_document_subplan_md_path",
            "long_document_coherence_base_path",
            "long_document_coherence_live_path",
            "long_document_references_path",
            "long_document_references_json_path",
            "long_document_visual_index_path",
            "long_document_visual_index_json_path",
            "long_document_source_manifest_path",
            "long_document_source_manifest_json_path",
            "long_document_summary",
            "research_kb_used",
            "research_kb_name",
            "research_kb_hit_count",
            "research_kb_citations",
            "research_kb_warning",
            "draft_response",
        ],
        "requirements": ["configured_llm"],
        "display_name": "Deep Research Report Agent",
        "category": "documents",
        "intent_patterns": [
            "deep research report", "write a complete document", "research and write", "create a report",
            "generate a handbook", "produce a whitepaper", "write a guide",
            "document about", "full report on", "comprehensive guide to",
        ],
        "active_when": [],
        "config_hint": "Choose a configured model in Studio. OpenAI is still required when Deep Research web search is enabled.",
    }
}


def _safe_int(value: Any, default: int, minimum: int, maximum: int) -> int:
    try:
        parsed = int(value)
    except Exception:
        parsed = default
    return max(minimum, min(maximum, parsed))


def _parallelism(value: Any, *, env_key: str, default: int, minimum: int, maximum: int) -> int:
    candidate = value
    if candidate is None or str(candidate).strip() == "":
        candidate = os.getenv(env_key, str(default))
    return _safe_int(candidate, default, minimum, maximum)


def _normalize_title(value: str, fallback: str) -> str:
    title = str(value or "").strip()
    if title:
        return title
    return fallback


def _objective_title_fallback(objective: str) -> str:
    value = re.sub(r"\s+", " ", str(objective or "").strip())
    if not value:
        return "Deep Research Report"

    patterns = [
        r"^(?:create|draft|write|produce|generate|prepare|build)\s+(?:a|an|the)?\s*(?:deep\s+research\s+)?(?:report|study|analysis|brief|dossier|document)\s+(?:on|about|for)\s+",
        r"^(?:what\s+is|explain|analyze|analyse|investigate|research|summarize|summarise)\s+",
    ]
    for pattern in patterns:
        value = re.sub(pattern, "", value, flags=re.IGNORECASE).strip()

    value = value.rstrip("?.! ")
    if not value:
        return "Deep Research Report"
    trimmed = _trim_text(value, 90)
    return trimmed[0].upper() + trimmed[1:] if trimmed else "Deep Research Report"


def _generate_report_title(objective: str, *, fallback: str = "Deep Research Report") -> str:
    heuristic = _objective_title_fallback(objective)
    try:
        payload = llm_json(
            f"""
You are naming a deep research report.

Objective:
{objective}

Return ONLY valid JSON:
{{
  "title": "A concise professional report title in title case"
}}

Constraints:
- 4 to 10 words
- no quotes
- no trailing punctuation
- do not use the generic title "Deep Research Report"
- reflect the actual topic
""",
            {"title": heuristic},
        )
        candidate = str(payload.get("title", "") or "").strip()
    except Exception:
        candidate = heuristic

    candidate = re.sub(r"\s+", " ", candidate).strip().strip("\"'")
    if not candidate or candidate.lower() == "deep research report":
        candidate = heuristic
    return _normalize_title(candidate, fallback)


def _strip_heading_attributes(value: str) -> str:
    return re.sub(r"\s*\{#[^}]+\}\s*$", "", str(value or "").strip())


def _slugify_heading(value: str, *, fallback: str = "section") -> str:
    cleaned = _strip_heading_attributes(value).lower()
    slug = re.sub(r"[^a-z0-9]+", "-", cleaned).strip("-")
    return slug or fallback


def _heading_with_anchor(level: int, title: str, anchor: str) -> str:
    normalized_title = _strip_heading_attributes(title)
    normalized_anchor = _slugify_heading(anchor or title, fallback=f"heading-{level}")
    return f"{'#' * max(1, level)} {normalized_title} {{#{normalized_anchor}}}"


def _normalized_heading_candidate(value: str) -> str:
    cleaned = _strip_heading_attributes(value)
    cleaned = re.sub(r"^\s*#+\s*", "", cleaned)
    cleaned = re.sub(r"^\s*section\s+\d+\s*[:.\-]\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"^\s*\d+\s*[:.\-]\s*", "", cleaned)
    return re.sub(r"[^a-z0-9]+", " ", cleaned.lower()).strip()


def _read_text_file(path_value: str, fallback: str = "") -> str:
    path = str(path_value or "").strip()
    if not path:
        return fallback
    try:
        return Path(path).read_text(encoding="utf-8")
    except Exception:
        return fallback


def _trim_text(text: str, limit: int) -> str:
    value = str(text or "").strip()
    if len(value) <= limit:
        return value
    return value[: max(0, limit - 3)] + "..."


def _local_source_log_label(path_value: str, manifest_entry: Mapping[str, Any] | None = None) -> str:
    entry = manifest_entry if isinstance(manifest_entry, Mapping) else {}
    relative_path = str(entry.get("relative_path", "") or "").strip().replace("\\", "/")
    if relative_path:
        return relative_path
    file_name = str(entry.get("name", "") or "").strip()
    if file_name:
        return file_name
    raw_path = str(path_value or "").strip()
    return Path(raw_path).name or raw_path or "local file"


def _web_source_log_label(url: str) -> str:
    value = str(url or "").strip()
    if not value:
        return "website"
    parsed = urlparse(value)
    if parsed.netloc:
        compact = parsed.netloc + (parsed.path or "")
        if parsed.query:
            compact += f"?{parsed.query}"
        return _trim_text(compact, 120)
    return _trim_text(value, 120)


def _log_local_file_review(
    path_value: str,
    *,
    status: str,
    position: int,
    total: int,
    manifest_entry: Mapping[str, Any] | None = None,
    payload: Mapping[str, Any] | None = None,
) -> None:
    label = _local_source_log_label(path_value, manifest_entry)
    if status == "started":
        log_task_update(DEEP_RESEARCH_LABEL, f"Reviewing local file [{position}/{total}]: {label}")
        return
    metadata = payload.get("metadata", {}) if isinstance(payload, Mapping) and isinstance(payload.get("metadata"), dict) else {}
    error_text = str(metadata.get("error", "") or (payload.get("error", "") if isinstance(payload, Mapping) else "") or "").strip()
    reader = str(metadata.get("reader", "") or "").strip()
    char_count = len(str((payload or {}).get("text", "") or "").strip()) if isinstance(payload, Mapping) else 0
    if error_text or status == "failed":
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"Local file review failed [{position}/{total}]: {label} ({_trim_text(error_text or 'document parse failed', 140)})",
        )
        return
    outcome = [f"{char_count} chars" if char_count > 0 else "no readable text"]
    if reader:
        outcome.append(f"reader={reader}")
    log_task_update(DEEP_RESEARCH_LABEL, f"Reviewed local file [{position}/{total}]: {label} ({', '.join(outcome)}).")


def _log_web_review(
    url: str,
    *,
    status: str,
    position: int,
    total: int,
    payload: Mapping[str, Any] | None = None,
    context: str = "website",
) -> None:
    label = _web_source_log_label(url)
    if status == "started":
        log_task_update(DEEP_RESEARCH_LABEL, f"Reviewing {context} [{position}/{total}]: {label}")
        return
    error_text = str((payload or {}).get("error", "") or "").strip() if isinstance(payload, Mapping) else ""
    char_count = int((payload or {}).get("char_count", 0) or 0) if isinstance(payload, Mapping) else 0
    if not char_count and isinstance(payload, Mapping):
        char_count = len(str(payload.get("text", "") or "").strip())
    content_type = str((payload or {}).get("content_type", "") or "").strip() if isinstance(payload, Mapping) else ""
    if error_text or status == "failed":
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"{context.capitalize()} review failed [{position}/{total}]: {label} ({_trim_text(error_text or 'fetch failed', 140)})",
        )
        return
    outcome = [f"{char_count} chars" if char_count > 0 else "no readable text"]
    if content_type:
        outcome.append(content_type)
    log_task_update(DEEP_RESEARCH_LABEL, f"Reviewed {context} [{position}/{total}]: {label} ({', '.join(outcome)}).")


def _log_search_collection(
    query: str,
    payload: Mapping[str, Any] | None,
    *,
    max_queries: int = 6,
    max_results: int = 8,
) -> None:
    search_payload = payload if isinstance(payload, Mapping) else {}
    provider = str(search_payload.get("provider", "") or "").strip()
    if not provider:
        providers_tried = search_payload.get("providers_tried", [])
        if isinstance(providers_tried, list):
            provider = str(providers_tried[0] or "").strip()
    provider = provider or "none"

    query_plan = search_payload.get("query_plan", [])
    if isinstance(query_plan, list) and query_plan:
        total_queries = len(query_plan)
        for index, plan in enumerate(query_plan[:max_queries], start=1):
            if not isinstance(plan, Mapping):
                continue
            plan_query = _trim_text(str(plan.get("query", "") or query).strip(), 180)
            if not plan_query:
                continue
            metadata: list[str] = []
            scope = str(plan.get("scope", "") or "").strip()
            timelimit = str(plan.get("timelimit", "") or "").strip()
            if scope:
                metadata.append(scope)
            if timelimit:
                metadata.append(f"timelimit={timelimit}")
            metadata_text = f" ({', '.join(metadata)})" if metadata else ""
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Search query [{index}/{total_queries}] via {provider}: {plan_query}{metadata_text}",
            )
        if total_queries > max_queries:
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Additional search queries omitted from log: {total_queries - max_queries}.",
            )
    elif str(query or "").strip():
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"Search query via {provider}: {_trim_text(str(query).strip(), 180)}",
        )

    results = search_payload.get("results", [])
    if isinstance(results, list) and results:
        total_results = len(results)
        for index, item in enumerate(results[:max_results], start=1):
            if not isinstance(item, Mapping):
                continue
            url = str(item.get("url", "") or "").strip()
            if not url:
                continue
            title = _trim_text(str(item.get("title", "") or "").strip(), 100)
            title_text = f" — {title}" if title else ""
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Collected search result [{index}/{total_results}] via {provider}: {_web_source_log_label(url)}{title_text}",
            )
        if total_results > max_results:
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Additional collected search results omitted from log: {total_results - max_results}.",
            )
        return

    error_text = _trim_text(str(search_payload.get("error", "") or "").strip(), 180)
    log_task_update(
        DEEP_RESEARCH_LABEL,
        (
            f"No search result URLs collected via {provider}: {error_text}"
            if error_text
            else f"No search result URLs collected via {provider}."
        ),
    )


def _normalize_research_formats(raw_value: Any) -> list[str]:
    if isinstance(raw_value, str):
        items = [part.strip().lower() for part in raw_value.split(",")]
    elif isinstance(raw_value, list):
        items = [str(part).strip().lower() for part in raw_value]
    else:
        items = []
    seen: set[str] = set()
    normalized: list[str] = []
    for item in items:
        if not item or item not in {"pdf", "docx", "html", "md"} or item in seen:
            continue
        seen.add(item)
        normalized.append(item)
    return normalized or list(DEFAULT_RESEARCH_FORMATS)


def _normalize_citation_style(raw_value: Any) -> str:
    style = str(raw_value or "").strip().lower()
    if style in SUPPORTED_CITATION_STYLES:
        return style
    return "apa"


def _deep_research_strategy(state: dict[str, Any], *, objective: str, max_files: int) -> tuple[dict[str, Any], dict[str, Any]]:
    intent = state.get("deep_research_intent", {}) if isinstance(state.get("deep_research_intent"), dict) else {}
    if not intent:
        intent = discover_research_intent(objective, state)
        state["deep_research_intent"] = intent
    strategy = state.get("deep_research_source_strategy", {}) if isinstance(state.get("deep_research_source_strategy"), dict) else {}
    if not strategy:
        strategy = build_source_strategy(
            intent,
            max_files=max_files,
            allow_web_search=bool(state.get("research_web_search_enabled", True)),
            local_paths_present=bool(state.get("local_drive_paths") or state.get("local_drive_files") or state.get("local_drive_document_summaries")),
        )
        state["deep_research_source_strategy"] = strategy
    return intent, strategy


def _copy_artifact_alias(source_path: str, destination_path: str) -> str:
    source = str(source_path or "").strip()
    destination = str(destination_path or "").strip()
    if not source or not destination:
        return ""
    try:
        src = Path(_output_file_path(source))
        dst = Path(_output_file_path(destination))
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(src, dst)
        return str(dst)
    except Exception:
        return ""


def _mirror_root_report_artifacts(
    compiled_path: str,
    export_paths: Mapping[str, Any] | None = None,
    *,
    root_report_dir: str = "reports",
) -> dict[str, str]:
    safe_exports = export_paths if isinstance(export_paths, Mapping) else {}
    target_map = {
        "md": _artifact_file(root_report_dir, "report.md"),
        "html": _artifact_file(root_report_dir, "report.html"),
        "pdf": _artifact_file(root_report_dir, "report.pdf"),
        "docx": _artifact_file(root_report_dir, "report.docx"),
    }
    source_map = {
        "md": compiled_path,
        "html": str(safe_exports.get("html", "") or ""),
        "pdf": str(safe_exports.get("pdf", "") or ""),
        "docx": str(safe_exports.get("docx", "") or ""),
    }
    mirrored: dict[str, str] = {}
    for fmt, destination in target_map.items():
        source = str(source_map.get(fmt, "") or "").strip()
        if not source:
            continue
        mirrored_path = _copy_artifact_alias(source, destination)
        if mirrored_path:
            mirrored[fmt] = mirrored_path
    return mirrored


def _normalize_reference_identity(reference: Mapping[str, Any], fallback_index: int) -> tuple[str, str, str]:
    url = str(reference.get("url", "")).strip()
    label = str(reference.get("label", "")).strip() or str(reference.get("title", "")).strip() or f"Source {fallback_index}"
    source_type = str(reference.get("type", "")).strip() or ("local_file" if url.startswith("file:") else "web")
    return url, label, source_type


def _build_evidence_ledger(
    *,
    consolidated_references: list[dict[str, Any]],
    section_outputs: list[dict[str, Any]],
    local_entries: list[dict[str, Any]],
    url_entries: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    by_key: dict[str, dict[str, Any]] = {}
    section_reference_keys: dict[str, set[str]] = {}
    for section in section_outputs or []:
        if not isinstance(section, dict):
            continue
        section_key = f"section-{int(section.get('index', 0) or 0):02d}"
        refs = section.get("references", []) if isinstance(section.get("references", []), list) else []
        normalized_keys: set[str] = set()
        for ref in refs:
            if not isinstance(ref, dict):
                continue
            ref_url = str(ref.get("url", "")).strip().lower()
            ref_label = str(ref.get("label", "")).strip().lower()
            if ref_url:
                normalized_keys.add(ref_url)
            if ref_label:
                normalized_keys.add(ref_label)
        section_reference_keys[section_key] = normalized_keys
    for index, reference in enumerate(consolidated_references or [], start=1):
        if not isinstance(reference, dict):
            continue
        url, label, source_type = _normalize_reference_identity(reference, index)
        key = url or label.lower()
        entry = by_key.get(key)
        if entry is None:
            entry = {
                "source_id": str(reference.get("id", "")).strip() or f"S{index}",
                "path": url if url.startswith("file:") else "",
                "url": "" if url.startswith("file:") else url,
                "label": label,
                "type": source_type,
                "extract_quality": "high" if url or label else "low",
                "used_in_sections": [],
                "claim_links": [],
                "confidence": 0.92 if url else 0.65,
                "reader": str(reference.get("reader", "")).strip(),
                "char_count": int(reference.get("char_count", 0) or 0),
                "error_kind": str(reference.get("error_kind", "")).strip(),
            }
            by_key[key] = entry
        normalized_keys = {url.lower(), label.lower()} - {""}
        for section_key, ref_keys in section_reference_keys.items():
            if normalized_keys & ref_keys:
                if section_key not in entry["used_in_sections"]:
                    entry["used_in_sections"].append(section_key)
                if section_key not in entry["claim_links"]:
                    entry["claim_links"].append(section_key)

    start_index = len(by_key) + 1
    for item in local_entries or []:
        path_value = str(item.get("path", "")).strip()
        if not path_value:
            continue
        key = path_value.lower()
        error_text = str(item.get("error", "")).strip()
        error_kind = str(item.get("error_kind", "")).strip()
        char_count = int(item.get("char_count", 0) or 0)
        extract_quality = "failed" if error_text else "high" if char_count > 400 else "medium" if char_count > 80 else "low"
        confidence = 0.18 if error_text and error_kind == "corrupt" else 0.35 if error_text else 0.9 if char_count > 400 else 0.72 if char_count > 80 else 0.55
        entry = by_key.setdefault(
            key,
            {
                "source_id": f"L{start_index}",
                "path": path_value,
                "url": _file_source_url(path_value),
                "label": str(item.get("file_name", "")).strip() or Path(path_value).name,
                "type": str(item.get("type", "")).strip() or "local_file",
                "extract_quality": extract_quality,
                "used_in_sections": [],
                "claim_links": [],
                "confidence": confidence,
                "reader": str(item.get("reader", "")).strip(),
                "char_count": char_count,
                "error_kind": error_kind,
            },
        )
        start_index += 1
        if error_text:
            entry["extract_quality"] = extract_quality
            entry["confidence"] = confidence

    for item in url_entries or []:
        url = str(item.get("url", "")).strip()
        if not url:
            continue
        key = url.lower()
        error_text = str(item.get("error", "")).strip()
        char_count = int(item.get("char_count", 0) or 0)
        entry = by_key.setdefault(
            key,
            {
                "source_id": f"U{start_index}",
                "path": "",
                "url": url,
                "label": str(item.get("label", "")).strip() or _source_label(url),
                "type": str(item.get("source_type", "")).strip() or "provided_url",
                "extract_quality": "failed" if error_text else "high" if char_count > 400 else "medium" if char_count > 80 else "low",
                "used_in_sections": [],
                "claim_links": [],
                "confidence": 0.9 if not error_text and char_count > 400 else 0.78 if not error_text else 0.3,
                "reader": str(item.get("content_type", "")).strip(),
                "char_count": char_count,
                "error_kind": "fetch_failed" if error_text else "",
            },
        )
        start_index += 1
        if error_text:
            entry["extract_quality"] = "failed"

    return sorted(by_key.values(), key=lambda item: str(item.get("source_id", "")))


def _build_coverage_report(
    *,
    objective: str,
    intent: Mapping[str, Any],
    source_strategy: Mapping[str, Any],
    local_manifest: Mapping[str, Any],
    local_entries: list[dict[str, Any]],
    url_entries: list[dict[str, Any]],
    kb_grounding: Mapping[str, Any],
    kb_warning: str,
    evidence_sources: list[dict[str, Any]],
    consolidated_references: list[dict[str, Any]],
) -> dict[str, Any]:
    discovered = int(local_manifest.get("file_count", 0) or 0)
    selected = int(local_manifest.get("selected_file_count", 0) or 0)
    selected_family_counts = local_manifest.get("selected_family_counts", {}) if isinstance(local_manifest.get("selected_family_counts", {}), Mapping) else {}
    excluded_reason_counts = local_manifest.get("excluded_reason_counts", {}) if isinstance(local_manifest.get("excluded_reason_counts", {}), Mapping) else {}
    manifest_files = local_manifest.get("files", []) if isinstance(local_manifest.get("files", []), list) else []
    failed_local_entries = [item for item in (local_entries or []) if str(item.get("error", "")).strip()]
    missing_families: list[str] = []
    gaps: list[str] = []
    needs = [str(item).strip().lower() for item in (intent.get("source_needs", []) or []) if str(item).strip()]
    if "tables" in needs and int(selected_family_counts.get("spreadsheet", 0) or 0) == 0:
        gaps.append("Table-heavy objective detected, but no spreadsheet sources were selected.")
        missing_families.append("spreadsheet")
    if "images" in needs and int(selected_family_counts.get("presentation", 0) or 0) == 0 and int(selected_family_counts.get("image", 0) or 0) == 0:
        gaps.append("Visual-heavy objective detected, but no slide or image sources were selected.")
        missing_families.append("presentation/image")
    if "code" in needs and int(selected_family_counts.get("code", 0) or 0) == 0 and int(selected_family_counts.get("config", 0) or 0) == 0:
        gaps.append("Technical objective detected, but no code/config sources were selected.")
        missing_families.append("code/config")
    if bool(source_strategy.get("web_search_needed", False)) and not evidence_sources and not url_entries:
        gaps.append("Web research was expected, but no external evidence sources were collected yet.")
        missing_families.append("web")
    if bool(kb_grounding.get("requested", False)) and int(kb_grounding.get("hit_count", 0) or 0) <= 0:
        gaps.append(
            "Knowledge base grounding was requested, but no relevant KB hits were retrieved."
            if not kb_warning
            else f"Knowledge base grounding warning: {kb_warning}"
        )
    if selected <= 0 and discovered > 0:
        gaps.append("Files were discovered locally, but none were selected for processing.")
    if failed_local_entries:
        gaps.append(f"{len(failed_local_entries)} selected file(s) failed extraction and may need alternate readers or user replacement.")

    skipped_examples: dict[str, list[str]] = {}
    for item in manifest_files:
        if not isinstance(item, dict):
            continue
        reason = str(item.get("exclusion_reason", "")).strip()
        if not reason:
            continue
        name = str(item.get("name", "")).strip() or str(item.get("relative_path", "")).strip()
        if not name:
            continue
        bucket = skipped_examples.setdefault(reason, [])
        if len(bucket) < 5:
            bucket.append(name)

    revisit_plan: list[dict[str, Any]] = []
    failed_extractions: list[dict[str, Any]] = []
    for family in missing_families:
        revisit_plan.append(
            {
                "target": family,
                "reason": f"Missing evidence family: {family}.",
                "action": (
                    "Expand local selection toward this family or enable web sources for a second pass."
                    if family != "web"
                    else "Enable web search or provide explicit URLs for a second pass."
                ),
            }
        )
    for item in failed_local_entries[:5]:
        target = str(item.get("file_name", "")).strip() or Path(str(item.get("path", "")).strip()).name
        reason = str(item.get("error_kind", "")).strip() or "extract_failed"
        message = str(item.get("error", "")).strip()
        failed_extractions.append({"target": target, "reason": reason, "message": message})
        revisit_plan.append(
            {
                "target": target,
                "reason": reason,
                "action": "Retry with alternate parser/OCR or replace the damaged/protected file.",
            }
        )

    status = "good"
    if gaps:
        status = "needs_attention"
    if discovered > 0 and selected == 0:
        status = "blocked"
    elif failed_local_entries and status == "good":
        status = "needs_attention"

    return {
        "objective": objective,
        "status": status,
        "discovered_files": discovered,
        "selected_files": selected,
        "read_files": len(local_entries),
        "failed_selected_files": len(failed_local_entries),
        "explicit_urls": len(url_entries),
        "evidence_sources": len(evidence_sources),
        "citations": len(consolidated_references),
        "selected_family_counts": dict(selected_family_counts),
        "excluded_reason_counts": dict(excluded_reason_counts),
        "remaining_work": max(0, selected - len(local_entries)),
        "kb_enabled": bool(kb_grounding.get("requested", False)),
        "kb_name": str(kb_grounding.get("kb_name", "")).strip(),
        "kb_status": str(kb_grounding.get("kb_status", "")).strip(),
        "kb_hit_count": int(kb_grounding.get("hit_count", 0) or 0),
        "kb_source_count": len(_sources_from_kb_grounding(kb_grounding)),
        "kb_warning": kb_warning,
        "missing_families": missing_families,
        "failed_extractions": failed_extractions,
        "skipped_examples": skipped_examples,
        "revisit_plan": revisit_plan,
        "gaps": gaps,
        "strategy_summary": str(source_strategy.get("summary", "")).strip(),
        "why_selected": dict(source_strategy.get("selection_notes", {})) if isinstance(source_strategy.get("selection_notes", {}), Mapping) else {},
        "why_skipped": dict(source_strategy.get("skip_notes", {})) if isinstance(source_strategy.get("skip_notes", {}), Mapping) else {},
    }


def _build_quality_report(
    *,
    section_outputs: list[dict[str, Any]],
    evidence_ledger: list[dict[str, Any]],
    coverage_report: Mapping[str, Any],
    consolidated_references: list[dict[str, Any]],
) -> dict[str, Any]:
    weak_sections: list[dict[str, Any]] = []
    unsupported_sections: list[dict[str, Any]] = []
    for section in section_outputs or []:
        refs = section.get("references", []) if isinstance(section.get("references", []), list) else []
        title = str(section.get("title", "")).strip() or f"Section {section.get('index', '?')}"
        if not refs:
            unsupported_sections.append({"section": title, "issue": "No supporting references were attached."})
        elif len(refs) < 2:
            weak_sections.append({"section": title, "issue": "Only one supporting reference was attached."})

    duplicate_keys: set[str] = set()
    seen_keys: set[str] = set()
    for ref in consolidated_references or []:
        if not isinstance(ref, dict):
            continue
        key = str(ref.get("url", "")).strip().lower() or str(ref.get("label", "")).strip().lower()
        if not key:
            continue
        if key in seen_keys:
            duplicate_keys.add(key)
        seen_keys.add(key)
    flags: list[str] = []
    if unsupported_sections:
        flags.append("unsupported_claim_risk")
    if weak_sections:
        flags.append("weak_section_grounding")
    if duplicate_keys:
        flags.append("duplicate_sources")
    if coverage_report.get("gaps"):
        flags.append("coverage_gaps")
    if int(coverage_report.get("failed_selected_files", 0) or 0) > 0:
        flags.append("failed_extractions")

    status = "pass"
    if unsupported_sections or coverage_report.get("status") == "blocked":
        status = "fail"
    elif weak_sections or duplicate_keys or coverage_report.get("gaps") or int(coverage_report.get("failed_selected_files", 0) or 0) > 0:
        status = "warn"

    return {
        "status": status,
        "flags": flags,
        "unsupported_sections": unsupported_sections,
        "weak_sections": weak_sections,
        "duplicate_source_count": len(duplicate_keys),
        "coverage_gap_count": len(coverage_report.get("gaps", []) or []),
        "failed_extraction_count": int(coverage_report.get("failed_selected_files", 0) or 0),
        "ledger_source_count": len(evidence_ledger),
        "summary": (
            "Quality gate passed."
            if status == "pass"
            else "Quality gate found grounding or coverage risks that should be reviewed."
        ),
    }


def _tier_budget(tier: int) -> dict[str, Any]:
    budgets = {
        1: {"max_tokens": 5000, "max_sources": 3, "max_duration_minutes": 1},
        2: {"max_tokens": 30000, "max_sources": 20, "max_duration_minutes": 5},
        3: {"max_tokens": 150000, "max_sources": 60, "max_duration_minutes": 20},
        4: {"max_tokens": 500000, "max_sources": 150, "max_duration_minutes": 90},
        5: {"max_tokens": 0, "max_sources": 0, "max_duration_minutes": 0},
    }
    return budgets.get(_safe_int(tier, 2, 1, 5), budgets[2])


def _default_subtopics(objective: str) -> list[str]:
    seeds = [
        "Market structure and baseline context",
        "Key players and competitive dynamics",
        "Evidence, benchmarks, and current data points",
        "Risks, contradictions, and areas of uncertainty",
        "Implications, outlook, and conclusion",
    ]
    text = str(objective or "").strip()
    if not text:
        return seeds[:3]
    normalized = re.sub(r"\s+", " ", text)

    def _clean_topic(raw: str) -> str:
        value = re.sub(
            r"^(?:so\s+i\s+am\s+looking\s+to\s+research|now\s+what\s+i['’]?m\s+trying\s+to\s+figure\s+out\s+is|what\s+i['’]?m\s+trying\s+to\s+figure\s+out\s+is)\s+",
            "",
            str(raw or "").strip(),
            flags=re.IGNORECASE,
        )
        value = re.sub(r"^(?:and\s+also|also)\s+", "", value, flags=re.IGNORECASE)
        value = value.strip(" -:;,.?")
        if not value:
            return ""
        return value[0].upper() + value[1:]

    topical: list[str] = []
    intro_match = re.split(
        r"\b(?:now\s+what\s+i['’]?m\s+trying\s+to\s+figure\s+out\s+is|what\s+i['’]?m\s+trying\s+to\s+figure\s+out\s+is)\b",
        normalized,
        maxsplit=1,
        flags=re.IGNORECASE,
    )
    if intro_match:
        intro = _clean_topic(intro_match[0])
        if len(intro.split()) >= 6:
            topical.append(intro)

    numbered_parts = [part for part in re.split(r"\b\d+\.\s*", normalized) if part.strip()]
    if len(numbered_parts) > 1:
        numbered_parts = numbered_parts[1:]
    for part in numbered_parts[:6]:
        cleaned = _clean_topic(part)
        if len(cleaned.split()) >= 3:
            topical.append(cleaned)

    if not topical:
        fragments = re.split(r"[?.!;\n]", normalized)
        topical.extend(_clean_topic(frag) for frag in fragments if len(str(frag).strip().split()) >= 4)

    deduped: list[str] = []
    seen: set[str] = set()
    for item in topical:
        key = item.lower()
        if not item or key in seen:
            continue
        seen.add(key)
        deduped.append(item)

    if len(deduped) < 3:
        for seed in seeds:
            if seed.lower() in seen:
                continue
            deduped.append(seed)
            seen.add(seed.lower())
            if len(deduped) >= 4:
                break

    if seeds[-1].lower() not in seen:
        deduped.append(seeds[-1])
    return deduped[:6]


def _research_depth_analysis(
    objective: str,
    *,
    target_pages: int,
    depth_mode: str,
    requested_sources: list[str],
    date_range: str,
    max_sources: int = 0,
) -> dict[str, Any]:
    text = str(objective or "").strip()
    lowered = text.lower()
    depth_config = _research_depth_config(depth_mode, target_pages)
    heuristic_tier = 2
    reason_bits: list[str] = []

    if len(text.split()) > 50:
        heuristic_tier += 1
        reason_bits.append("query length suggests broad scope")
    if any(token in lowered for token in ("compare", "analyse", "analyze", "comprehensive", "thorough", "deep", "full report", "detailed", "exhaustive")):
        heuristic_tier += 1
        reason_bits.append("explicit deep-analysis wording")
    if target_pages >= 100:
        heuristic_tier += 2
        reason_bits.append("very large page target")
    elif target_pages >= 50:
        heuristic_tier += 2
        reason_bits.append("large page target")
    elif target_pages >= 20:
        heuristic_tier += 1
        reason_bits.append("multi-section page target")
    if any(token in lowered for token in ("academic", "peer-reviewed", "scholarly", "patent", "scientific", "literature review")):
        heuristic_tier += 1
        reason_bits.append("academic / patent scope")
    if date_range and str(date_range).strip().lower() not in {"", "all_time", "all time"}:
        heuristic_tier += 1
        reason_bits.append("explicit temporal scope")
    if any(token in lowered for token in ("quick", "brief", "summary", "tldr", "short")):
        heuristic_tier -= 1
        reason_bits.append("short-form wording lowers depth")
    if len(re.findall(r"\b[A-Z][a-zA-Z0-9&.-]+\b", text)) >= 4:
        heuristic_tier += 1
        reason_bits.append("multiple named entities detected")
    heuristic_tier = max(1, min(5, heuristic_tier))

    estimated_pages = target_pages if target_pages > 0 else {1: 1, 2: 4, 3: 18, 4: 50, 5: 120}[heuristic_tier]
    estimated_sources = min(200, max(3, {1: 3, 2: 15, 3: 45, 4: 110, 5: 200}[heuristic_tier] + (10 if requested_sources else 0)))
    estimated_duration = {1: 1, 2: 5, 3: 15, 4: 45, 5: 120}[heuristic_tier]
    subtopics = _default_subtopics(text)

    fallback = {
        "tier": heuristic_tier,
        "reason": "; ".join(reason_bits) or "heuristic estimate",
        "estimated_sources": estimated_sources,
        "estimated_pages": estimated_pages,
        "depth_mode": depth_config["mode"],
        "depth_label": depth_config["label"],
        "depth_description": depth_config["description"],
        "subtopics": subtopics[:6],
        "requires_deep_research": heuristic_tier >= 3,
        "estimated_duration_minutes": estimated_duration,
    }

    prompt = f"""
You are a research complexity analyser.

Given the query and heuristic estimate below, determine the final research tier (1-5).
Respond with JSON only:
{{
  "tier": 3,
  "reason": "string",
  "estimated_sources": 40,
  "estimated_pages": 15,
  "depth_mode": "standard",
  "depth_label": "Standard Report",
  "depth_description": "Balanced multi-section research with enough room for evidence, tradeoffs, and conclusions.",
  "subtopics": ["...", "..."],
  "requires_deep_research": true,
  "estimated_duration_minutes": 15
}}

Query:
{text}

Heuristic estimate:
{json.dumps(fallback, ensure_ascii=False)}
"""
    try:
        data = llm_json(prompt, fallback)
    except Exception:
        data = fallback

    if not isinstance(data, dict):
        data = fallback
    tier = _safe_int(data.get("tier"), heuristic_tier, 1, 5)
    estimated_pages = _safe_int(data.get("estimated_pages"), fallback["estimated_pages"], 1, 500)
    estimated_sources = _safe_int(data.get("estimated_sources"), fallback["estimated_sources"], 1, 400)
    estimated_duration = _safe_int(data.get("estimated_duration_minutes"), fallback["estimated_duration_minutes"], 1, 1440)
    subtopics_value = data.get("subtopics", [])
    if not isinstance(subtopics_value, list):
        subtopics_value = fallback["subtopics"]
    subtopics = [str(item).strip() for item in subtopics_value if str(item).strip()][:8] or fallback["subtopics"]
    budget = dict(_tier_budget(tier))
    execution_budget = dict(budget)
    if budget["max_sources"]:
        estimated_sources = min(estimated_sources, int(budget["max_sources"]))
    if max_sources > 0:
        execution_budget["max_sources"] = min(int(execution_budget.get("max_sources", 0) or 0), max_sources) if execution_budget.get("max_sources", 0) else max_sources
        estimated_sources = min(estimated_sources, execution_budget["max_sources"])
    return {
        "tier": tier,
        "reason": str(data.get("reason", "")).strip() or fallback["reason"],
        "estimated_sources": estimated_sources,
        "estimated_pages": estimated_pages,
        "depth_mode": _normalize_research_depth_mode(data.get("depth_mode"), target_pages),
        "depth_label": str(data.get("depth_label", "")).strip() or depth_config["label"],
        "depth_description": str(data.get("depth_description", "")).strip() or depth_config["description"],
        "subtopics": subtopics,
        "requires_deep_research": bool(data.get("requires_deep_research", tier >= 3)),
        "estimated_duration_minutes": estimated_duration,
        "budget": budget,
        "execution_budget": execution_budget,
        "requested_target_pages": target_pages,
        "requested_depth_mode": depth_config["mode"],
        "requested_max_sources": max_sources,
        "requested_sources": requested_sources,
        "date_range": date_range or "all_time",
        "request_signature": {
            "objective": text,
            "target_pages": target_pages,
            "depth_mode": depth_config["mode"],
            "requested_sources": list(requested_sources),
            "date_range": date_range or "all_time",
            "max_sources": max_sources,
        },
    }


def _coherence_base_context(state: dict, objective: str) -> str:
    memory_blocks = [
        ("Agent.md", _read_text_file(state.get("memory_agent_file"), "")),
        ("soul.md", _read_text_file(state.get("memory_soul_file"), "")),
        ("memory.md", _read_text_file(state.get("memory_long_term_file"), "")),
        ("session.md", _read_text_file(state.get("memory_session_file"), "")),
        ("planning.md", _read_text_file(state.get("memory_planning_file"), "")),
    ]
    parts = [f"[Objective]\n{objective}"]
    for name, content in memory_blocks:
        if not str(content).strip():
            continue
        parts.append(f"[{name}]\n{_trim_text(content, 3000)}")
    fallback_context = str(state.get("file_memory_context", "")).strip()
    if fallback_context:
        parts.append(f"[file_memory_context]\n{_trim_text(fallback_context, 4000)}")
    return "\n\n".join(parts)


def _artifact_file(run_dir: str, filename: str) -> str:
    return f"{run_dir.rstrip('/')}/{filename.lstrip('/')}"


def _normalize_output_relative_path(path_value: str) -> str:
    value = str(path_value or "").strip()
    prefix = f"{OUTPUT_DIR}/"
    if value.startswith(prefix):
        return value[len(prefix):]
    return value


def _output_file_path(path_value: str) -> str:
    value = str(path_value or "").strip()
    if not value:
        return ""
    candidate = Path(value)
    if candidate.is_absolute():
        return str(candidate)
    normalized = _normalize_output_relative_path(value)
    return resolve_output_path(normalized)


def _resolve_existing_output_path(path_value: str) -> str:
    if not str(path_value or "").strip():
        return ""
    resolved = _output_file_path(str(path_value))
    if Path(resolved).exists():
        return resolved
    resolved_fallback = resolve_output_path(str(path_value))
    if Path(resolved_fallback).exists():
        return resolved_fallback
    return resolved


def _read_json_file(path_value: str, fallback: Any) -> Any:
    path = str(path_value or "").strip()
    if not path:
        return fallback
    try:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    except Exception:
        return fallback


def _normalize_objective_key(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").strip()).lower()


def _discover_deep_research_run_dirs() -> list[str]:
    base_dir = Path(resolve_output_path("deep_research_runs"))
    if not base_dir.exists() or not base_dir.is_dir():
        return []
    discovered: list[tuple[int, str]] = []
    for entry in base_dir.iterdir():
        if not entry.is_dir():
            continue
        match = re.fullmatch(r"deep_research_run_(\d+)", entry.name)
        if not match:
            continue
        discovered.append((int(match.group(1)), f"deep_research_runs/{entry.name}"))
    discovered.sort(key=lambda item: item[0], reverse=True)
    return [item[1] for item in discovered]


def _artifact_dir_has_resume_data(artifact_dir: str) -> bool:
    rel = str(artifact_dir or "").strip().rstrip("/")
    if not rel:
        return False
    base = Path(resolve_output_path(rel))
    if not base.exists() or not base.is_dir():
        return False
    probe_files = (
        "deep_research_progress.json",
        "checkpoint_after_research.json",
        "checkpoint_after_writing.json",
        "evidence_bank.json",
        "evidence_bank.md",
        "source_manifest.json",
        "source_manifest.md",
        "deep_research_analysis.md",
        "deep_research_subplan.md",
    )
    if any((base / name).exists() for name in probe_files):
        return True
    for section_dir in base.glob("section_*"):
        if not section_dir.is_dir():
            continue
        if (section_dir / "research.json").exists() or (section_dir / "section.md").exists():
            return True
    return False


def _artifact_dir_objective_key(artifact_dir: str) -> str:
    for rel_path in ("evidence_bank.json", "deep_research_manifest.json", "source_manifest.json"):
        payload = _read_json_file(resolve_output_path(_artifact_file(artifact_dir, rel_path)), {})
        if isinstance(payload, dict):
            value = _normalize_objective_key(payload.get("objective", ""))
            if value:
                return value
    return ""


def _resolve_resume_artifact_dir(
    state: dict,
    *,
    call_number: int,
    default_artifact_dir: str,
    objective: str,
) -> str:
    seen: set[str] = set()
    candidates: list[str] = []

    def _add(path_value: str) -> None:
        value = _normalize_output_relative_path(str(path_value or "").strip()).rstrip("/")
        if not value or value in seen:
            return
        seen.add(value)
        candidates.append(value)

    _add(state.get("long_document_artifact_dir", ""))
    if call_number > 1:
        _add(f"deep_research_runs/deep_research_run_{call_number - 1}")
    _add(default_artifact_dir)
    for discovered in _discover_deep_research_run_dirs():
        _add(discovered)

    objective_key = _normalize_objective_key(objective)
    for candidate in candidates:
        if not _artifact_dir_has_resume_data(candidate):
            continue
        candidate_objective_key = _artifact_dir_objective_key(candidate)
        if candidate_objective_key and objective_key and candidate_objective_key != objective_key:
            continue
        return candidate
    return default_artifact_dir


def _load_cached_evidence_bank(
    *,
    artifact_dir: str,
    objective: str,
    max_sources: int,
) -> dict[str, Any]:
    evidence_json_abs = resolve_output_path(_artifact_file(artifact_dir, "evidence_bank.json"))
    evidence_md_abs = resolve_output_path(_artifact_file(artifact_dir, "evidence_bank.md"))
    if not Path(evidence_json_abs).exists() and not Path(evidence_md_abs).exists():
        return {}
    payload = _read_json_file(evidence_json_abs, {})
    if not isinstance(payload, dict):
        payload = {}
    cached_objective = _normalize_objective_key(payload.get("objective", ""))
    if cached_objective and cached_objective != _normalize_objective_key(objective):
        return {}
    evidence_md = _read_text_file(evidence_md_abs, "")
    source_ledger = payload.get("source_ledger", [])
    if not isinstance(source_ledger, list):
        source_ledger = []
    evidence_sources = [item for item in source_ledger if isinstance(item, dict)]
    if max_sources > 0:
        evidence_sources = evidence_sources[:max_sources]
    if not evidence_md and not evidence_sources:
        return {}
    return {
        "artifact_dir": artifact_dir,
        "evidence_markdown": evidence_md,
        "evidence_sources": evidence_sources,
        "evidence_path": _output_file_path(_artifact_file(artifact_dir, "evidence_bank.md")),
        "evidence_json_path": _output_file_path(_artifact_file(artifact_dir, "evidence_bank.json")),
    }


def _path_within_root(path_value: str, root_value: str) -> bool:
    try:
        normalized_path = os.path.normcase(os.path.abspath(str(path_value or "").strip()))
        normalized_root = os.path.normcase(os.path.abspath(str(root_value or "").strip()))
        if not normalized_path or not normalized_root:
            return False
        return os.path.commonpath([normalized_path, normalized_root]) == normalized_root
    except Exception:
        return False


def _collect_local_source_roots(state: dict[str, Any]) -> list[str]:
    base_directory = (
        state.get("local_drive_working_directory")
        or state.get("document_working_directory")
        or state.get("working_directory")
    )
    raw_roots = (
        state.get("local_drive_paths")
        or state.get("knowledge_drive_paths")
        or state.get("drive_paths")
        or state.get("document_root_paths")
        or []
    )
    return _resolve_local_drive_paths(raw_roots, base_directory)


def _build_local_source_tree(local_manifest: Mapping[str, Any]) -> list[dict[str, Any]]:
    roots = [str(item).strip() for item in (local_manifest.get("roots", []) or []) if str(item).strip()]
    folder_entries = [item for item in (local_manifest.get("folders", []) or []) if isinstance(item, dict)]
    file_entries = [item for item in (local_manifest.get("files", []) or []) if isinstance(item, dict)]
    folder_paths = {str(item.get("path", "")).strip() for item in folder_entries}
    file_by_path = {
        str(item.get("path", "")).strip(): item
        for item in file_entries
        if str(item.get("path", "")).strip()
    }
    tree: list[dict[str, Any]] = []

    def _node_from_entry(entry: Mapping[str, Any], *, include_children: bool = True) -> dict[str, Any]:
        node = {
            "name": str(entry.get("name", "")).strip() or Path(str(entry.get("path", "")).strip()).name or str(entry.get("path", "")).strip(),
            "path": str(entry.get("path", "")).strip(),
            "type": "directory" if str(entry.get("entry_type", "")).strip() == "directory" else "file",
            "relative_path": str(entry.get("relative_path", "")).strip(),
            "depth": int(entry.get("depth", 0) or 0),
        }
        if node["type"] == "file":
            node["selected_for_processing"] = bool(entry.get("selected_for_processing"))
            node["exclusion_reason"] = str(entry.get("exclusion_reason", "")).strip()
            node["selection_reason"] = str(entry.get("selection_reason", "")).strip()
            node["extension"] = str(entry.get("extension", "")).strip()
        elif include_children:
            node["children"] = []
        return node

    def _sort_tree(nodes: list[dict[str, Any]]) -> list[dict[str, Any]]:
        sorted_nodes = sorted(
            nodes,
            key=lambda item: (0 if item.get("type") == "directory" else 1, str(item.get("name", "")).lower()),
        )
        for node in sorted_nodes:
            children = node.get("children")
            if isinstance(children, list):
                node["children"] = _sort_tree(children)
        return sorted_nodes

    for root in roots:
        if root in file_by_path and root not in folder_paths:
            tree.append(_node_from_entry(file_by_path[root], include_children=False))
            continue
        root_node = {
            "name": Path(root).name or root,
            "path": root,
            "type": "directory",
            "relative_path": ".",
            "depth": 0,
            "children": [],
        }
        children_by_relative: dict[str, dict[str, Any]] = {"": root_node}
        scoped_entries: list[dict[str, Any]] = []
        for entry in folder_entries + file_entries:
            path_value = str(entry.get("path", "")).strip()
            if not path_value or path_value == root or not _path_within_root(path_value, root):
                continue
            scoped_entries.append(entry)
        scoped_entries.sort(
            key=lambda item: (
                int(item.get("depth", 0) or 0),
                0 if str(item.get("entry_type", "")).strip() == "directory" else 1,
                str(item.get("relative_path", "")).lower(),
                str(item.get("name", "")).lower(),
            )
        )
        for entry in scoped_entries:
            relative_path = str(entry.get("relative_path", "")).strip().replace("\\", "/")
            if not relative_path or relative_path == ".":
                continue
            parent_relative = str(Path(relative_path).parent).replace("\\", "/")
            if parent_relative == ".":
                parent_relative = ""
            parent = children_by_relative.get(parent_relative, root_node)
            node = _node_from_entry(entry)
            parent.setdefault("children", []).append(node)
            if node.get("type") == "directory":
                children_by_relative[relative_path] = node
        tree.append(root_node)
    return _sort_tree(tree)


def _tree_status_label(node: Mapping[str, Any]) -> str:
    if str(node.get("type", "")).strip() == "directory":
        return ""
    if bool(node.get("selected_for_processing", False)):
        return " [selected]"
    exclusion_reason = str(node.get("exclusion_reason", "")).strip()
    return f" [skipped: {exclusion_reason or 'not_selected'}]"


def _render_local_source_tree_markdown(nodes: list[dict[str, Any]], *, depth: int = 0) -> list[str]:
    lines: list[str] = []
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        indent = "  " * depth
        label = str(node.get("name", "")).strip() or Path(str(node.get("path", "")).strip()).name or str(node.get("path", "")).strip()
        if depth == 0:
            label = str(node.get("path", "")).strip() or label
        if str(node.get("type", "")).strip() == "directory":
            lines.append(f"{indent}- {label}/")
            children = node.get("children", [])
            if isinstance(children, list):
                lines.extend(_render_local_source_tree_markdown(children, depth=depth + 1))
            continue
        lines.append(f"{indent}- {label}{_tree_status_label(node)}")
    return lines


def _format_source_manifest_markdown(
    *,
    objective: str,
    source_strategy: Mapping[str, Any],
    local_manifest: Mapping[str, Any],
    source_tree: list[dict[str, Any]],
    document_summaries: list[dict[str, Any]],
    rollup_summary: str,
) -> str:
    roots = [str(item).strip() for item in (local_manifest.get("roots", []) or []) if str(item).strip()]
    lines = [
        "# Source Manifest",
        "",
        "## Objective",
        str(objective or "").strip() or "No objective provided.",
        "",
        "## Discovery Summary",
        f"- Roots scanned: {len(roots)}",
        f"- Folders discovered: {int(local_manifest.get('folder_count', 0) or 0)}",
        f"- Files discovered: {int(local_manifest.get('file_count', 0) or 0)}",
        f"- Files selected for processing: {int(local_manifest.get('selected_file_count', 0) or 0)}",
        f"- Files excluded from processing: {int(local_manifest.get('excluded_file_count', 0) or 0)}",
        "",
        "## Source Strategy",
        str(source_strategy.get("summary", "")).strip() or "No source strategy recorded.",
        "",
        "## File Tree",
    ]
    tree_lines = _render_local_source_tree_markdown(source_tree)
    lines.extend(tree_lines or ["- No local files or folders were discovered."])
    lines.extend(["", "## Selected File Summaries"])
    if document_summaries:
        for item in document_summaries:
            if not isinstance(item, dict):
                continue
            lines.append(f"### {str(item.get('file_name', '')).strip() or Path(str(item.get('path', '')).strip()).name or 'Unnamed file'}")
            if str(item.get("path", "")).strip():
                lines.append(f"- Path: {item['path']}")
            if str(item.get("type", "")).strip():
                lines.append(f"- Type: {item['type']}")
            if int(item.get("char_count", 0) or 0) > 0:
                lines.append(f"- Characters: {int(item.get('char_count', 0) or 0)}")
            if str(item.get("reader", "")).strip():
                lines.append(f"- Reader: {item['reader']}")
            if str(item.get("selection_reason", "")).strip():
                lines.append(f"- Selection reason: {item['selection_reason']}")
            if str(item.get("error", "")).strip():
                lines.append(f"- Extraction error: {item['error']}")
            summary_text = str(item.get("summary", "")).strip()
            if summary_text:
                lines.append("- Summary:")
                lines.append(textwrap.indent(summary_text, "  "))
            lines.append("")
    else:
        lines.extend(["- No selected files were summarized.", ""])
    if str(rollup_summary or "").strip():
        lines.extend(["## Local Source Rollup", str(rollup_summary).strip(), ""])
    return "\n".join(lines).rstrip() + "\n"


def _summarize_local_documents(
    *,
    objective: str,
    documents: list[dict[str, Any]],
    local_manifest: Mapping[str, Any],
) -> list[dict[str, Any]]:
    manifest_files = local_manifest.get("files", []) if isinstance(local_manifest.get("files", []), list) else []
    manifest_by_path = {
        str(item.get("path", "")).strip(): item
        for item in manifest_files
        if isinstance(item, dict) and str(item.get("path", "")).strip()
    }
    summaries: list[dict[str, Any]] = []
    for index, parsed in enumerate(documents or [], start=1):
        if not isinstance(parsed, dict):
            continue
        file_path = str(parsed.get("path", "")).strip()
        parsed_text = str(parsed.get("text", "") or "").strip()
        metadata = parsed.get("metadata", {}) if isinstance(parsed.get("metadata"), dict) else {}
        document_type = str(metadata.get("type", Path(file_path).suffix.lstrip(".").lower() or "unknown"))
        manifest_entry = manifest_by_path.get(file_path, {})
        if parsed_text:
            try:
                summary = llm_text(
                    f"""
You are a document-reading sub-agent for a deep research workflow.

Task objective:
{objective}

Document path:
{file_path}

Document type:
{document_type}

Extracted content:
{parsed_text[:16000]}

Write a concise summary with:
- what this document is about
- critical facts, numbers, dates, and entities
- decisions or action items
- data quality concerns or missing pieces
"""
                ).strip()
            except Exception:
                summary = _trim_text(parsed_text, 1400)
        else:
            summary = f"No readable text extracted. Reason: {metadata.get('error', 'empty document after extraction')}"
        summaries.append(
            {
                "index": index,
                "path": file_path,
                "file_name": Path(file_path).name or file_path or f"Local file {index}",
                "type": document_type,
                "summary": summary,
                "char_count": len(parsed_text),
                "error": str(metadata.get("error", "") or ""),
                "error_kind": str(metadata.get("error_kind", "") or ""),
                "recoverable": bool(metadata.get("recoverable", False)),
                "reader": str(metadata.get("reader", "") or ""),
                "selection_reason": str(manifest_entry.get("selection_reason", "") or ""),
                "metadata": dict(metadata),
            }
        )
    return summaries


def _roll_up_local_document_summaries(
    *,
    objective: str,
    document_summaries: list[dict[str, Any]],
) -> str:
    if not document_summaries:
        return ""
    rollup_input = {
        "objective": objective,
        "document_count": len(document_summaries),
        "documents": [
            {
                "index": item.get("index"),
                "path": item.get("path", ""),
                "type": item.get("type", ""),
                "summary": item.get("summary", ""),
                "error": item.get("error", ""),
            }
            for item in document_summaries
        ],
    }
    try:
        return llm_text(
            f"""
You are a knowledge-synthesis agent. Build one actionable summary from these document summaries.

Return:
- 5-10 bullet key findings
- contradictions or missing data
- recommended next tasks for report generation
- list of highest-priority source files to inspect deeper

Input:
{json.dumps(rollup_input, indent=2, ensure_ascii=False)[:28000]}
"""
        ).strip()
    except Exception:
        strong_files = [
            str(item.get("file_name", "")).strip()
            for item in document_summaries
            if str(item.get("summary", "")).strip() and not str(item.get("error", "")).strip()
        ]
        return (
            "Local-source rollup fallback:\n"
            f"- Objective: {objective or 'not provided'}\n"
            f"- Files summarized: {len(document_summaries)}\n"
            f"- Strongest sources: {', '.join(strong_files[:5]) or 'none'}"
        )


def _load_cached_local_source_manifest(
    *,
    artifact_dir: str,
    objective: str,
) -> dict[str, Any]:
    manifest_json_abs = resolve_output_path(_artifact_file(artifact_dir, "source_manifest.json"))
    manifest_md_abs = resolve_output_path(_artifact_file(artifact_dir, "source_manifest.md"))
    if not Path(manifest_json_abs).exists() and not Path(manifest_md_abs).exists():
        return {}
    payload = _read_json_file(manifest_json_abs, {})
    if not isinstance(payload, dict):
        payload = {}
    cached_objective = _normalize_objective_key(payload.get("objective", ""))
    if cached_objective and cached_objective != _normalize_objective_key(objective):
        return {}
    local_manifest = payload.get("manifest", {})
    document_summaries = payload.get("document_summaries", [])
    if not isinstance(local_manifest, dict):
        local_manifest = {}
    if not isinstance(document_summaries, list):
        document_summaries = []
    if not local_manifest and not document_summaries and not Path(manifest_md_abs).exists():
        return {}
    return {
        "artifact_dir": artifact_dir,
        "manifest": local_manifest,
        "source_tree": payload.get("tree", []) if isinstance(payload.get("tree", []), list) else [],
        "document_summaries": [item for item in document_summaries if isinstance(item, dict)],
        "rollup_summary": str(payload.get("rollup_summary", "") or ""),
        "manifest_path": _output_file_path(_artifact_file(artifact_dir, "source_manifest.md")),
        "manifest_json_path": _output_file_path(_artifact_file(artifact_dir, "source_manifest.json")),
        "manifest_markdown": _read_text_file(manifest_md_abs, ""),
    }


def _ensure_local_source_manifest(
    state: dict[str, Any],
    *,
    objective: str,
    artifact_dir: str,
    source_strategy: Mapping[str, Any],
) -> dict[str, Any]:
    local_manifest = state.get("local_drive_manifest", {}) if isinstance(state.get("local_drive_manifest"), dict) else {}
    document_summaries = state.get("local_drive_document_summaries", []) if isinstance(state.get("local_drive_document_summaries"), list) else []
    rollup_summary = str(state.get("local_drive_rollup_summary", "") or "")
    roots = _collect_local_source_roots(state)

    if not local_manifest and not document_summaries:
        cached_manifest = _load_cached_local_source_manifest(artifact_dir=artifact_dir, objective=objective)
        if cached_manifest:
            local_manifest = cached_manifest.get("manifest", {}) if isinstance(cached_manifest.get("manifest", {}), dict) else {}
            document_summaries = list(cached_manifest.get("document_summaries", []) or [])
            rollup_summary = str(cached_manifest.get("rollup_summary", "") or "")
            state["local_drive_manifest"] = local_manifest
            state["local_drive_document_summaries"] = document_summaries
            state["local_drive_files"] = list(local_manifest.get("selected_files", []) or [])
            state["local_drive_summary_bank"] = {
                str(item.get("path", "")).strip(): str(item.get("summary", "")).strip()
                for item in document_summaries
                if isinstance(item, dict) and str(item.get("path", "")).strip()
            }
            state["local_drive_rollup_summary"] = rollup_summary
            state["long_document_source_manifest_path"] = str(cached_manifest.get("manifest_path", "") or "")
            state["long_document_source_manifest_json_path"] = str(cached_manifest.get("manifest_json_path", "") or "")
            return {
                "from_cache": True,
                "manifest": local_manifest,
                "document_summaries": document_summaries,
                "rollup_summary": rollup_summary,
                "source_tree": list(cached_manifest.get("source_tree", []) or []),
            }

    if not local_manifest and not roots:
        return {}

    if not local_manifest:
        local_drive_started_at = _trace_now()
        _trace_research_event(
            state,
            title="Discovering local source tree",
            detail="Scanning attached local files and folders before research synthesis begins.",
            command="\n".join(roots),
            status="running",
            started_at=local_drive_started_at,
            metadata={"phase": "local_source_manifest", "roots": roots},
            subtask="Build local source manifest",
        )
        local_manifest = _scan_local_drive_tree(
            roots,
            recursive=bool(state.get("local_drive_recursive", True)),
            include_hidden=bool(state.get("local_drive_include_hidden", False)),
            max_files=max(1, min(int(state.get("local_drive_max_files", 200) or 200), 1000)),
            allowed_extensions=_normalize_extension_set(state.get("local_drive_extensions")),
            objective=objective,
            source_strategy=source_strategy,
        )
        state["local_drive_manifest"] = local_manifest
        state["local_drive_files"] = list(local_manifest.get("selected_files", []) or [])
        _trace_research_event(
            state,
            title="Discovering local source tree",
            detail=(
                f"Discovered {int(local_manifest.get('file_count', 0) or 0)} file(s) and selected "
                f"{int(local_manifest.get('selected_file_count', 0) or 0)} for processing."
            ),
            command="\n".join(roots),
            status="completed",
            started_at=local_drive_started_at,
            completed_at=_trace_now(),
            metadata={"phase": "local_source_manifest", "roots": roots, "manifest": local_manifest},
            subtask="Build local source manifest",
        )

    selected_files = list(local_manifest.get("selected_files", []) or state.get("local_drive_files") or [])
    if not document_summaries and selected_files:
        manifest_files = local_manifest.get("files", []) if isinstance(local_manifest.get("files", []), list) else []
        manifest_by_path = {
            str(item.get("path", "")).strip(): item
            for item in manifest_files
            if isinstance(item, dict) and str(item.get("path", "")).strip()
        }

        def _document_progress_callback(path_value: str, status: str, payload: Mapping[str, Any] | None, position: int, total: int) -> None:
            _log_local_file_review(
                path_value,
                status=status,
                position=position,
                total=total,
                manifest_entry=manifest_by_path.get(str(path_value or "").strip(), {}),
                payload=payload,
            )

        documents = parse_documents(
            selected_files,
            continue_on_error=True,
            ocr_images=bool(state.get("local_drive_enable_image_ocr", True)),
            ocr_instruction=state.get("local_drive_ocr_instruction"),
            progress_callback=_document_progress_callback,
        )
        state["local_drive_documents"] = documents
        document_summaries = _summarize_local_documents(
            objective=objective,
            documents=documents,
            local_manifest=local_manifest,
        )
        state["local_drive_document_summaries"] = document_summaries
        state["local_drive_summary_bank"] = {
            str(item.get("path", "")).strip(): str(item.get("summary", "")).strip()
            for item in document_summaries
            if isinstance(item, dict) and str(item.get("path", "")).strip()
        }

    if not rollup_summary and document_summaries:
        rollup_summary = _roll_up_local_document_summaries(
            objective=objective,
            document_summaries=document_summaries,
        )
        state["local_drive_rollup_summary"] = rollup_summary

    source_tree = _build_local_source_tree(local_manifest)
    payload = {
        "objective": objective,
        "roots": list(local_manifest.get("roots", []) or roots),
        "manifest": local_manifest,
        "tree": source_tree,
        "document_summaries": document_summaries,
        "rollup_summary": rollup_summary,
        "source_strategy": dict(source_strategy) if isinstance(source_strategy, Mapping) else {},
    }
    manifest_md = _format_source_manifest_markdown(
        objective=objective,
        source_strategy=source_strategy,
        local_manifest=local_manifest,
        source_tree=source_tree,
        document_summaries=document_summaries,
        rollup_summary=rollup_summary,
    )
    manifest_filename = _artifact_file(artifact_dir, "source_manifest.md")
    manifest_json_filename = _artifact_file(artifact_dir, "source_manifest.json")
    write_text_file(manifest_filename, manifest_md)
    write_text_file(manifest_json_filename, json.dumps(payload, indent=2, ensure_ascii=False))
    state["long_document_source_manifest_path"] = _output_file_path(manifest_filename)
    state["long_document_source_manifest_json_path"] = _output_file_path(manifest_json_filename)
    return {
        "from_cache": False,
        "manifest": local_manifest,
        "document_summaries": document_summaries,
        "rollup_summary": rollup_summary,
        "source_tree": source_tree,
    }


def _load_cached_section_research_package(
    *,
    artifact_dir: str,
    objective: str,
    section: dict[str, Any],
    section_index: int,
    section_pages: int,
    collect_sources_first: bool,
    evidence_excerpt: str,
    evidence_sources: list[dict],
    explicit_source_entries: list[dict],
    max_sources: int,
) -> dict[str, Any]:
    section_title = str(section.get("title", f"Section {section_index}")).strip() or f"Section {section_index}"
    section_objective = str(section.get("objective", objective)).strip() or objective
    section_questions = section.get("key_questions", [])
    if not isinstance(section_questions, list):
        section_questions = []
    target_section_pages = _safe_int(section.get("target_pages"), section_pages, 1, 30)

    research_abs = resolve_output_path(_artifact_file(artifact_dir, f"section_{section_index:02d}/research.json"))
    sources_abs = resolve_output_path(_artifact_file(artifact_dir, f"section_{section_index:02d}/sources.json"))
    if not Path(research_abs).exists() and not Path(sources_abs).exists():
        return {}

    research_pass = _read_json_file(research_abs, {})
    if not isinstance(research_pass, dict):
        research_pass = {}
    status = str(research_pass.get("status", "")).strip() or "completed"
    if status not in {"completed", "local_only", "evidence_bank"}:
        return {}

    section_sources_raw = _read_json_file(sources_abs, [])
    section_sources = [item for item in section_sources_raw if isinstance(item, dict)] if isinstance(section_sources_raw, list) else []
    if not section_sources and collect_sources_first and evidence_sources:
        section_sources = list(evidence_sources)
    if not section_sources:
        section_sources = _merge_sources(explicit_source_entries, _extract_source_entries(research_pass))
    if max_sources > 0:
        section_sources = section_sources[:max_sources]

    research_output = str(research_pass.get("output_text", "")).strip()
    if not research_output and collect_sources_first and evidence_excerpt:
        research_output = evidence_excerpt
    if not research_output:
        research_output = "Reused cached research package did not include output text."

    return {
        "index": section_index,
        "title": section_title,
        "objective": section_objective,
        "key_questions": section_questions,
        "target_pages": target_section_pages,
        "research_pass": research_pass,
        "research_text": research_output,
        "sources": section_sources,
        "source_ledger_md": _source_ledger_markdown(section_sources),
        "search_query": "",
        "section_search_results": {},
        "from_cache": True,
        "cache_artifact_dir": artifact_dir,
    }


def _fallback_continuity_note(section_text: str) -> str:
    lines: list[str] = []
    for raw in str(section_text or "").splitlines():
        cleaned = re.sub(r"^#{1,6}\s*", "", raw).strip()
        if not cleaned or cleaned.startswith("!["):
            continue
        lines.append(cleaned)
        if len(lines) >= 6:
            break
    if not lines:
        return "- Continuity note unavailable in cached draft."
    return "\n".join(f"- {item}" for item in lines)


def _load_cached_section_draft(
    *,
    artifact_dir: str,
    section_index: int,
    section_title: str,
) -> dict[str, Any]:
    section_md_abs = resolve_output_path(_artifact_file(artifact_dir, f"section_{section_index:02d}/section.md"))
    if not Path(section_md_abs).exists():
        return {}
    section_text = _read_text_file(section_md_abs, "").strip()
    if not section_text:
        return {}

    visual_assets_abs = resolve_output_path(_artifact_file(artifact_dir, f"section_{section_index:02d}/visual_assets.json"))
    continuity_abs = resolve_output_path(_artifact_file(artifact_dir, f"section_{section_index:02d}/continuity.txt"))
    metadata_abs = resolve_output_path(_artifact_file(artifact_dir, f"section_{section_index:02d}/section_metadata.json"))

    visual_assets = _read_json_file(visual_assets_abs, {})
    if not isinstance(visual_assets, dict):
        visual_assets = {"tables": [], "flowcharts": [], "notes": ""}
    continuity_note = _read_text_file(continuity_abs, "").strip() or _fallback_continuity_note(section_text)

    metadata = _read_json_file(metadata_abs, {})
    if not isinstance(metadata, dict):
        metadata = {}
    cached_title = str(metadata.get("section_title", "")).strip()
    if cached_title and cached_title != section_title:
        return {}

    flowchart_files = metadata.get("flowchart_files", [])
    if not isinstance(flowchart_files, list):
        flowchart_files = []
    section_images = metadata.get("section_images", [])
    if not isinstance(section_images, list):
        section_images = []

    return {
        "section_text": section_text,
        "continuity_note": continuity_note,
        "visual_assets": visual_assets,
        "flowchart_files": flowchart_files,
        "section_images": section_images,
        "from_cache": True,
        "cache_artifact_dir": artifact_dir,
    }


def _source_label(url: str) -> str:
    parsed = urlparse(url)
    if parsed.scheme == "file":
        file_name = Path(parsed.path or "").name
        return f"file:{file_name or 'local-source'}"
    host = parsed.netloc or "source"
    path = (parsed.path or "/").strip("/")
    if not path:
        return host
    parts = [item for item in path.split("/") if item][:2]
    return f"{host}/{'/'.join(parts)}"


def _collect_local_drive_evidence(state: dict, *, max_items: int = 25) -> list[dict]:
    summaries = state.get("local_drive_document_summaries") or []
    if not isinstance(summaries, list):
        return []
    entries = []
    for item in summaries[:max_items]:
        if not isinstance(item, dict):
            continue
        entries.append(
            {
                "source_type": "local_drive",
                "path": str(item.get("path", "")).strip(),
                "file_name": str(item.get("file_name", "")).strip(),
                "type": str(item.get("type", "")).strip(),
                "char_count": int(item.get("char_count", 0) or 0),
                "summary": str(item.get("summary", "")).strip(),
                "error": str(item.get("error", "")).strip(),
            }
        )
    return entries


def _file_source_url(path_value: str) -> str:
    try:
        return Path(str(path_value or "")).resolve().as_uri()
    except Exception:
        return f"file://{str(path_value or '').strip()}"


def _sources_from_local_entries(entries: list[dict]) -> list[dict]:
    sources = []
    for index, entry in enumerate(entries or [], start=1):
        path_value = str(entry.get("path", "")).strip()
        if not path_value:
            continue
        label = str(entry.get("file_name", "")).strip() or Path(path_value).name or f"Local file {index}"
        sources.append({"id": f"L{index}", "url": _file_source_url(path_value), "label": label})
    return sources


def _normalize_research_urls(raw_value: Any) -> list[str]:
    if isinstance(raw_value, str):
        raw_items = re.split(r"[\n,]", raw_value)
    elif isinstance(raw_value, list):
        raw_items = raw_value
    else:
        raw_items = []
    urls: list[str] = []
    seen: set[str] = set()
    for item in raw_items:
        candidate = _normalize_url(str(item or "").strip())
        if not candidate or candidate in seen:
            continue
        seen.add(candidate)
        urls.append(candidate)
    return urls


def _sources_from_url_entries(entries: list[dict]) -> list[dict]:
    sources = []
    for index, entry in enumerate(entries or [], start=1):
        url = str(entry.get("url", "")).strip()
        if not url:
            continue
        label = str(entry.get("label", "")).strip() or _source_label(url)
        sources.append({"id": f"U{index}", "url": url, "label": label})
    return sources


def _sources_from_kb_grounding(grounding: Mapping[str, Any], *, prefix: str = "K") -> list[dict]:
    entries: list[dict] = []
    citations = grounding.get("citations", []) if isinstance(grounding.get("citations", []), list) else []
    for index, item in enumerate(citations, start=1):
        if not isinstance(item, Mapping):
            continue
        source_id = str(item.get("source_id", "")).strip()
        url = str(item.get("url", "")).strip()
        path_value = str(item.get("path", "")).strip()
        ref = url or path_value or source_id
        if not ref:
            continue
        label = str(item.get("label", "")).strip() or Path(path_value or source_id).name or ref
        entries.append(
            {
                "id": f"{prefix}{index}",
                "url": ref,
                "label": label,
                "path": path_value,
                "source_id": source_id,
                "type": str(item.get("source_type", "")).strip() or "knowledge_base",
                "source_type": str(item.get("source_type", "")).strip(),
                "kb_provenance": dict(item.get("kb_provenance", {}) or {}),
                "chunk_index": item.get("chunk_index"),
                "score": item.get("score"),
            }
        )
    return entries


def _kb_grounding_summary(grounding: Mapping[str, Any]) -> str:
    context = str(grounding.get("prompt_context", "") or "").strip()
    if context:
        return _trim_text(context, 12000)
    kb_name = str(grounding.get("kb_name", "") or "").strip()
    if not kb_name:
        return ""
    return f"Knowledge Base Grounding:\n- KB: {kb_name}\n- Hits: {int(grounding.get('hit_count', 0) or 0)}"


def _collect_single_user_url_entry(objective: str, url: str) -> dict:
    try:
        page = fetch_url_content(url, timeout=20)
        page_text = _trim_text(str(page.get("text", "")).strip(), 8000)
        if page_text:
            summary = llm_text(
                f"""
You are summarizing a user-provided URL for a deep research report.

Primary objective:
{objective}

URL:
{url}

Extracted content:
{page_text}

Write a concise evidence summary covering:
- what this page is about
- concrete facts, numbers, dates, and named entities
- relevance to the objective
- any credibility or freshness concerns
"""
            ).strip()
        else:
            summary = "No readable text was extracted from this URL."
        return {
            "source_type": "provided_url",
            "url": url,
            "label": _source_label(url),
            "content_type": str(page.get("content_type", "")).strip(),
            "char_count": len(str(page.get("text", "")).strip()),
            "summary": summary,
            "excerpt": _trim_text(page_text, 2500),
            "error": "",
        }
    except Exception as exc:
        return {
            "source_type": "provided_url",
            "url": url,
            "label": _source_label(url),
            "content_type": "",
            "char_count": 0,
            "summary": "",
            "excerpt": "",
            "error": str(exc),
        }


def _collect_user_url_evidence(objective: str, state: dict, *, max_items: int = 12) -> list[dict]:
    urls = _normalize_research_urls(state.get("deep_research_source_urls", []))
    if not urls:
        return []
    _raise_if_cancelled(state, phase="provided_urls")
    bounded_urls = urls[:max_items]
    fetch_parallelism = min(
        len(bounded_urls),
        _parallelism(
            state.get("research_url_fetch_concurrency"),
            env_key="KENDR_RESEARCH_URL_FETCH_CONCURRENCY",
            default=4,
            minimum=1,
            maximum=12,
        ),
    )
    fetch_started_at = _trace_now()
    _trace_research_event(
        state,
        title="Extracting provided URLs",
        detail=(
            f"Fetching and summarizing {len(bounded_urls)} user-provided URLs"
            + (f" with {fetch_parallelism} parallel workers." if fetch_parallelism > 1 else ".")
        ),
        command="\n".join(bounded_urls),
        status="running",
        started_at=fetch_started_at,
        metadata={"phase": "provided_urls", "urls": _trace_url_list(bounded_urls), "parallelism": fetch_parallelism},
        subtask="Fetch provided URLs",
    )
    indexed_entries: dict[int, dict] = {}
    if fetch_parallelism <= 1:
        for index, url in enumerate(bounded_urls):
            _raise_if_cancelled(state, phase="provided_urls")
            _log_web_review(url, status="started", position=index + 1, total=len(bounded_urls), context="website")
            entry = _collect_single_user_url_entry(objective, url)
            indexed_entries[index] = entry
            _log_web_review(
                url,
                status="completed" if not str(entry.get("error", "")).strip() else "failed",
                position=index + 1,
                total=len(bounded_urls),
                payload=entry,
                context="website",
            )
    else:
        with ThreadPoolExecutor(max_workers=fetch_parallelism, thread_name_prefix="kendr-url") as executor:
            future_map = {
                executor.submit(_collect_single_user_url_entry, objective, url): (index, url)
                for index, url in enumerate(bounded_urls)
            }
            for index, url in enumerate(bounded_urls):
                _log_web_review(url, status="started", position=index + 1, total=len(bounded_urls), context="website")
            for future in as_completed(future_map):
                _raise_if_cancelled(state, phase="provided_urls")
                index, url = future_map[future]
                entry = future.result()
                indexed_entries[index] = entry
                _log_web_review(
                    url,
                    status="completed" if not str(entry.get("error", "")).strip() else "failed",
                    position=index + 1,
                    total=len(bounded_urls),
                    payload=entry,
                    context="website",
                )
    entries = [indexed_entries[index] for index in sorted(indexed_entries)]
    ok_urls = [str(item.get("url", "")).strip() for item in entries if not str(item.get("error", "")).strip()]
    failed_urls = [str(item.get("url", "")).strip() for item in entries if str(item.get("error", "")).strip()]
    _trace_research_event(
        state,
        title="Extracting provided URLs",
        detail=(
            f"Extracted {len(ok_urls)} URL(s)"
            + (f"; {len(failed_urls)} failed." if failed_urls else ".")
        ),
        command="\n".join(bounded_urls),
        status="completed" if not failed_urls else ("failed" if not ok_urls else "completed"),
        started_at=fetch_started_at,
        completed_at=_trace_now(),
        metadata={
            "phase": "provided_urls",
            "urls": _trace_url_list(ok_urls or urls[:max_items]),
            "viewed_urls": _trace_url_list(ok_urls),
            "failed_urls": _trace_url_list(failed_urls),
            "parallelism": fetch_parallelism,
        },
        subtask="Fetch provided URLs",
    )
    return entries


def _build_local_only_research_notes(
    *,
    objective: str,
    focus: str,
    local_entries: list[dict],
    url_entries: list[dict],
    continuity_notes: list[str] | None = None,
) -> str:
    corpus = {
        "objective": objective,
        "focus": focus,
        "continuity_notes": list(continuity_notes or [])[-8:],
        "local_files": [
            {
                "file_name": item.get("file_name", ""),
                "path": item.get("path", ""),
                "type": item.get("type", ""),
                "summary": item.get("summary", ""),
                "error": item.get("error", ""),
            }
            for item in (local_entries or [])[:25]
        ],
        "provided_urls": [
            {
                "url": item.get("url", ""),
                "label": item.get("label", ""),
                "summary": item.get("summary", ""),
                "error": item.get("error", ""),
                "excerpt": item.get("excerpt", ""),
            }
            for item in (url_entries or [])[:12]
        ],
    }
    return llm_text(
        f"""
You are synthesizing a deep research corpus without open web search.

Use only the supplied local files and explicitly provided URL extracts.
Do not invent external facts or imply that web search was performed.

Corpus:
{json.dumps(corpus, indent=2, ensure_ascii=False)[:24000]}

Write a structured evidence memo that includes:
- the strongest supported findings
- important numbers, dates, entities, and claims
- contradictions, gaps, or low-confidence areas
- which files or URLs are most relevant to this focus
"""
    ).strip()


def _collect_google_search_evidence(query: str, *, num: int = 10, progress_callback=None, search_backend: str = "auto") -> dict:
    normalized_backend = normalize_research_search_backend(search_backend)
    payload = fetch_search_results(
        query,
        num=num,
        fetch_pages=min(max(num, 1), 3),
        progress_callback=progress_callback,
        provider_hint="" if normalized_backend == "auto" else normalized_backend,
        focused_brief=query,
    )
    results = list(payload.get("results", []) or [])
    viewed_pages = list(payload.get("viewed_pages", []) or [])
    viewed_map = {
        str(item.get("url", "")).strip(): item
        for item in viewed_pages
        if str(item.get("url", "")).strip()
    }
    for item in results:
        url = str(item.get("url", "")).strip()
        if url and url in viewed_map:
            evidence = viewed_map[url]
            item["evidence_excerpt"] = str(evidence.get("excerpt", "")).strip()
            item["content_type"] = str(evidence.get("content_type", "")).strip()
            if evidence.get("error"):
                item["view_error"] = str(evidence.get("error", "")).strip()
    response_payload = {
        "results": results,
        "viewed_pages": viewed_pages,
        "provider": str(payload.get("provider", "")).strip(),
        "providers_tried": list(payload.get("providers_tried", []) or []),
        "raw": payload.get("raw", {}),
        "instant_answer": payload.get("instant_answer", {}),
        "query_plan": payload.get("query_plan", []),
        "error": str(payload.get("error", "")).strip(),
    }
    _log_search_collection(query, response_payload)
    return response_payload


# ---------------------------------------------------------------------------
# Image collection pipeline
# ---------------------------------------------------------------------------

_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg", ".bmp"}
_IMAGE_MAX_BYTES = int(os.getenv("KENDR_IMAGE_MAX_BYTES", str(4 * 1024 * 1024)))  # 4 MB default
_IMAGE_DOWNLOAD_TIMEOUT = int(os.getenv("KENDR_IMAGE_DOWNLOAD_TIMEOUT", "15"))
_IMAGE_SOURCE_PAGE_LIMIT = max(1, min(int(os.getenv("KENDR_IMAGE_SOURCE_PAGE_LIMIT", "5") or 5), 10))
_IMAGE_CANDIDATES_PER_PAGE = max(1, min(int(os.getenv("KENDR_IMAGE_CANDIDATES_PER_PAGE", "12") or 12), 40))
_IMAGE_MIN_DIMENSION = max(64, int(os.getenv("KENDR_IMAGE_MIN_DIMENSION", "120") or 120))
_IMAGE_RELEVANCE_THRESHOLD = float(os.getenv("KENDR_IMAGE_RELEVANCE_THRESHOLD", "5.0") or 5.0)
_IMAGE_CONTEXT_STOPWORDS = {
    "about", "after", "also", "among", "and", "are", "because", "been", "being", "between", "both",
    "but", "chart", "could", "data", "does", "figure", "from", "have", "into", "just", "more", "most",
    "over", "page", "report", "section", "should", "than", "that", "their", "them", "there", "these",
    "they", "this", "those", "through", "under", "using", "very", "what", "when", "where", "which",
    "with", "your",
}
_IMAGE_DECORATIVE_MARKERS = {
    "logo", "logos", "icon", "icons", "avatar", "avatars", "banner", "banners", "hero", "heroes", "thumbnail",
    "thumbnails", "thumb", "social", "footer", "header", "sprite", "badge", "placeholder", "promo", "advert",
    "advertisement", "marketing", "tracking", "pixel", "brand", "navbar", "button",
}
_IMAGE_PRIORITY_MARKERS = {
    "architecture", "diagram", "workflow", "process", "flow", "flowchart", "chart", "graph", "dashboard",
    "screenshot", "figure", "infographic", "timeline", "matrix", "heatmap", "pipeline", "network", "topology",
    "system", "interface", "console", "reference", "overview",
}
_IMAGE_CLASS_REJECTION_MARKERS = {"logo", "icon", "avatar", "badge", "hero", "banner", "social", "promo"}
_IMAGE_REQUIRED_LOCAL_SIGNAL_FIELDS = ("alt_text", "figcaption", "nearby_text", "heading")


def _normalize_image_term(token: str) -> str:
    value = re.sub(r"[^a-z0-9]+", "", str(token or "").lower())
    if len(value) > 4 and value.endswith("ies"):
        value = value[:-3] + "y"
    elif len(value) > 4 and value.endswith("s") and not value.endswith("ss"):
        value = value[:-1]
    return value


def _image_terms(*texts: Any, limit: int = 120) -> set[str]:
    terms: list[str] = []
    seen: set[str] = set()
    for raw in texts:
        for match in re.findall(r"[A-Za-z0-9][A-Za-z0-9_-]{2,}", str(raw or "").lower()):
            normalized = _normalize_image_term(match)
            if (
                not normalized
                or normalized in seen
                or normalized in _IMAGE_CONTEXT_STOPWORDS
                or normalized.isdigit()
            ):
                continue
            seen.add(normalized)
            terms.append(normalized)
            if len(terms) >= limit:
                return set(terms)
    return set(terms)


def _extract_image_src(tag: Any) -> str:
    for key in ("src", "data-src", "data-original", "data-lazy-src", "data-image", "data-url"):
        value = str(tag.get(key, "") or "").strip()
        if value:
            return value
    srcset = str(tag.get("srcset", "") or "").strip()
    if srcset:
        first_candidate = srcset.split(",")[0].strip().split(" ")[0].strip()
        if first_candidate:
            return first_candidate
    return ""


def _image_attr_int(value: Any) -> int:
    text = str(value or "").strip()
    if not text:
        return 0
    match = re.search(r"\d+", text)
    return int(match.group()) if match else 0


def _truncate_context_text(value: Any, limit: int = 280) -> str:
    text = re.sub(r"\s+", " ", str(value or "").strip())
    return _trim_text(text, limit) if text else ""


def _source_image_urls(section_sources: list[dict], *, max_pages: int = _IMAGE_SOURCE_PAGE_LIMIT) -> list[str]:
    urls: list[str] = []
    seen: set[str] = set()
    for item in section_sources or []:
        if not isinstance(item, dict):
            continue
        candidate = str(item.get("source_page", "") or item.get("url", "") or "").strip()
        parsed = urlparse(candidate)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc or candidate in seen:
            continue
        seen.add(candidate)
        urls.append(candidate)
        if len(urls) >= max_pages:
            break
    return urls


def _extract_source_page_image_candidates(
    page_url: str,
    payload: Mapping[str, Any],
    *,
    limit: int = _IMAGE_CANDIDATES_PER_PAGE,
) -> list[dict]:
    content_type = str(payload.get("content_type", "") or "").lower()
    raw_html = str(payload.get("raw_text", "") or "")
    if "html" not in content_type or not raw_html.strip():
        return []

    try:
        from bs4 import BeautifulSoup
    except Exception:
        BeautifulSoup = None

    candidates: list[dict] = []
    if BeautifulSoup is not None:
        soup = BeautifulSoup(raw_html, "html.parser")
        page_title = _truncate_context_text(soup.title.get_text(" ", strip=True) if soup.title else "", 180)
        for img in soup.find_all("img"):
            src = _extract_image_src(img)
            if not src:
                continue
            resolved_url = urljoin(page_url, src)
            parsed = urlparse(resolved_url)
            if parsed.scheme not in {"http", "https"}:
                continue
            figure = img.find_parent("figure")
            figcaption = ""
            if figure is not None:
                fig_node = figure.find("figcaption")
                if fig_node is not None:
                    figcaption = _truncate_context_text(fig_node.get_text(" ", strip=True), 220)
            parent = figure or img.parent
            nearby_text = ""
            if parent is not None and hasattr(parent, "get_text"):
                nearby_text = _truncate_context_text(parent.get_text(" ", strip=True), 260)
            heading = ""
            prev_heading = img.find_previous(["h1", "h2", "h3", "h4"])
            if prev_heading is not None:
                heading = _truncate_context_text(prev_heading.get_text(" ", strip=True), 160)
            img_classes = img.get("class", [])
            if isinstance(img_classes, str):
                img_classes = [img_classes]
            candidates.append(
                {
                    "url": resolved_url,
                    "source_page": page_url,
                    "source": urlparse(page_url).netloc,
                    "page_title": page_title,
                    "alt_text": _truncate_context_text(img.get("alt", ""), 180),
                    "title": _truncate_context_text(img.get("title", ""), 180),
                    "figcaption": figcaption,
                    "nearby_text": nearby_text,
                    "heading": heading,
                    "img_class": " ".join(str(item).strip() for item in img_classes if str(item).strip()),
                    "img_id": str(img.get("id", "") or "").strip(),
                    "width": _image_attr_int(img.get("width")),
                    "height": _image_attr_int(img.get("height")),
                }
            )
            if len(candidates) >= limit:
                break
        return candidates

    page_title_match = re.search(r"<title[^>]*>(.*?)</title>", raw_html, flags=re.IGNORECASE | re.DOTALL)
    page_title = _truncate_context_text(re.sub(r"<[^>]+>", " ", page_title_match.group(1) if page_title_match else ""), 180)
    for match in re.finditer(r"<img\b([^>]+)>", raw_html, flags=re.IGNORECASE | re.DOTALL):
        attrs = match.group(1)
        src_match = re.search(r'(?:src|data-src|data-original|data-lazy-src)=["\']([^"\']+)["\']', attrs, flags=re.IGNORECASE)
        if not src_match:
            continue
        resolved_url = urljoin(page_url, src_match.group(1).strip())
        parsed = urlparse(resolved_url)
        if parsed.scheme not in {"http", "https"}:
            continue
        alt_match = re.search(r'alt=["\']([^"\']*)["\']', attrs, flags=re.IGNORECASE)
        title_match = re.search(r'title=["\']([^"\']*)["\']', attrs, flags=re.IGNORECASE)
        width_match = re.search(r'width=["\']?(\d+)', attrs, flags=re.IGNORECASE)
        height_match = re.search(r'height=["\']?(\d+)', attrs, flags=re.IGNORECASE)
        class_match = re.search(r'class=["\']([^"\']*)["\']', attrs, flags=re.IGNORECASE)
        id_match = re.search(r'id=["\']([^"\']*)["\']', attrs, flags=re.IGNORECASE)
        candidates.append(
            {
                "url": resolved_url,
                "source_page": page_url,
                "source": urlparse(page_url).netloc,
                "page_title": page_title,
                "alt_text": _truncate_context_text(alt_match.group(1) if alt_match else "", 180),
                "title": _truncate_context_text(title_match.group(1) if title_match else "", 180),
                "figcaption": "",
                "nearby_text": "",
                "heading": "",
                "img_class": _truncate_context_text(class_match.group(1) if class_match else "", 120),
                "img_id": _truncate_context_text(id_match.group(1) if id_match else "", 120),
                "width": int(width_match.group(1)) if width_match else 0,
                "height": int(height_match.group(1)) if height_match else 0,
            }
        )
        if len(candidates) >= limit:
            break
    return candidates


def _score_source_image_candidate(
    candidate: Mapping[str, Any],
    *,
    section_title: str,
    section_objective: str,
    section_text: str,
) -> tuple[float, list[str], str]:
    topic_terms = _image_terms(section_title, section_objective, _trim_text(section_text, 1800))
    signal_sources = {
        "alt": str(candidate.get("alt_text", "") or ""),
        "caption": str(candidate.get("figcaption", "") or ""),
        "heading": str(candidate.get("heading", "") or ""),
        "nearby": str(candidate.get("nearby_text", "") or ""),
        "page_title": str(candidate.get("page_title", "") or ""),
        "title": str(candidate.get("title", "") or ""),
        "url": str(urlparse(str(candidate.get("url", "") or "")).path or ""),
    }
    weights = {"alt": 5.0, "caption": 6.0, "heading": 4.0, "nearby": 2.0, "page_title": 2.0, "title": 3.0, "url": 1.0}
    score = 0.0
    reasons: list[str] = []
    matched_terms: set[str] = set()
    local_match_terms: set[str] = set()
    for label, text in signal_sources.items():
        terms = _image_terms(text, limit=40)
        overlap = sorted(topic_terms & terms)
        if not overlap:
            continue
        matched_terms.update(overlap)
        if label in {"alt", "caption", "heading", "nearby"}:
            local_match_terms.update(overlap)
        score += weights.get(label, 1.0) * min(len(overlap), 3)
        reasons.append(f"{label}:{', '.join(overlap[:2])}")

    combined_terms = _image_terms(
        candidate.get("alt_text", ""),
        candidate.get("title", ""),
        candidate.get("figcaption", ""),
        candidate.get("heading", ""),
        candidate.get("nearby_text", ""),
        candidate.get("img_class", ""),
        candidate.get("img_id", ""),
        urlparse(str(candidate.get("url", "") or "")).path,
    )
    priority_hits = sorted(_IMAGE_PRIORITY_MARKERS & combined_terms)
    if priority_hits:
        score += 3.0 + min(len(priority_hits), 2)
        reasons.append(f"visual:{', '.join(priority_hits[:2])}")

    decorative_hits = sorted((_IMAGE_DECORATIVE_MARKERS | _IMAGE_CLASS_REJECTION_MARKERS) & combined_terms)
    if decorative_hits:
        score -= 8.0 + min(len(decorative_hits), 2)
        reasons.append(f"decorative:{', '.join(decorative_hits[:2])}")

    width = int(candidate.get("width", 0) or 0)
    height = int(candidate.get("height", 0) or 0)
    min_dimension = min(value for value in (width, height) if value > 0) if (width > 0 or height > 0) else 0
    if min_dimension and min_dimension < _IMAGE_MIN_DIMENSION:
        score -= 6.0
        reasons.append(f"size:{width}x{height}")
    elif min_dimension >= 200:
        score += 1.0

    alt_text = str(candidate.get("alt_text", "") or "").strip().lower()
    if alt_text in {"image", "photo", "graphic", "illustration", "banner", "hero", "logo"}:
        score -= 4.0
        reasons.append("generic_alt")

    rejection = ""
    if decorative_hits:
        rejection = "decorative"
    elif min_dimension and min_dimension < 96:
        rejection = "too_small"
    elif not local_match_terms:
        rejection = "weak_context"
    elif score < _IMAGE_RELEVANCE_THRESHOLD:
        rejection = "low_relevance"
    return score, reasons, rejection


def _parse_image_vision_verdict(text: str) -> tuple[bool | None, str]:
    raw = str(text or "").strip()
    if not raw:
        return None, ""
    try:
        payload = json.loads(strip_code_fences(raw))
        if isinstance(payload, dict):
            relevant_value = payload.get("relevant")
            decorative_value = payload.get("decorative")
            if decorative_value is True:
                return False, str(payload.get("reason", "") or "vision_marked_decorative").strip()
            if isinstance(relevant_value, bool):
                return relevant_value, str(payload.get("reason", "") or "").strip()
    except Exception:
        pass
    lowered = raw.lower()
    if '"relevant": false' in lowered or "relevant: false" in lowered or '"decorative": true' in lowered:
        return False, raw
    if '"relevant": true' in lowered or "relevant: true" in lowered:
        return True, raw
    return None, raw


def _verify_downloaded_image_with_vision(
    abs_path: str,
    *,
    section_title: str,
    section_objective: str,
    section_text: str,
) -> tuple[bool, str]:
    try:
        verdict = openai_analyze_image(
            abs_path,
            instruction=(
                "You are validating whether an image belongs in a professional deep-research report section.\n"
                f"Section title: {section_title}\n"
                f"Section objective: {section_objective}\n"
                f"Section preview: {_trim_text(section_text, 900)}\n\n"
                "Reject decorative assets, logos, hero banners, generic stock photos, portraits, and unrelated branding.\n"
                "Accept diagrams, dashboards, screenshots, charts, process visuals, system architecture, and figures that clearly support the section.\n"
                'Return ONLY JSON like {"relevant": true|false, "decorative": true|false, "reason": "short reason"}.'
            ),
        )
        relevant, reason = _parse_image_vision_verdict(str(verdict.get("analysis", "") or ""))
        if relevant is None:
            return True, "vision_inconclusive"
        return bool(relevant), reason or ("vision_verified" if relevant else "vision_rejected")
    except Exception as exc:
        return True, f"vision_unavailable:{exc}"


def _collect_image_search_results(query: str, *, num: int = 10) -> list[dict]:
    """Search Google Images via SerpAPI and return a list of image candidate dicts."""
    api_key = os.getenv("SERP_API_KEY", "").strip()
    if not api_key:
        log_task_update(DEEP_RESEARCH_LABEL, "Image search skipped: SERP_API_KEY not set.")
        return []
    try:
        from urllib.parse import urlencode
        from urllib.request import urlopen as _urlopen
        params = {
            "engine": "google_images",
            "q": query,
            "api_key": api_key,
            "num": num,
            "hl": "en",
            "gl": "us",
            "safe": "active",
        }
        url = f"https://serpapi.com/search.json?{urlencode(params)}"
        with _urlopen(url, timeout=20) as resp:
            payload = json.loads(resp.read().decode("utf-8"))
        candidates = []
        for item in payload.get("images_results", [])[:num]:
            original = str(item.get("original", "")).strip()
            thumbnail = str(item.get("thumbnail", "")).strip()
            img_url = original or thumbnail
            if not img_url:
                continue
            candidates.append({
                "url": img_url,
                "thumbnail": thumbnail,
                "title": str(item.get("title", "")).strip(),
                "source": str(item.get("source", "")).strip(),
                "source_page": str(item.get("link", "")).strip(),
                "width": int(item.get("original_width", 0) or 0),
                "height": int(item.get("original_height", 0) or 0),
            })
        log_task_update(DEEP_RESEARCH_LABEL, f"Image search '{query[:60]}': {len(candidates)} candidates found.")
        return candidates
    except Exception as exc:
        log_task_update(DEEP_RESEARCH_LABEL, f"Image search failed for '{query[:60]}': {exc}")
        return []


def _download_image(url: str, save_path: str, *, timeout: int = _IMAGE_DOWNLOAD_TIMEOUT, max_bytes: int = _IMAGE_MAX_BYTES) -> bool:
    """Download an image to save_path. Returns True on success."""
    try:
        from urllib.request import Request as _Request, urlopen as _urlopen
        req = _Request(
            url,
            headers={"User-Agent": os.getenv("RESEARCH_USER_AGENT", "multi-agent-research-bot/1.0")},
        )
        with _urlopen(req, timeout=timeout) as resp:
            content_type = resp.headers.get("Content-Type", "")
            if content_type and "image" not in content_type and "octet-stream" not in content_type:
                log_task_update(DEEP_RESEARCH_LABEL, f"Skipping non-image content-type '{content_type}' for {url[:80]}")
                return False
            data = b""
            chunk_size = 65536
            while True:
                chunk = resp.read(chunk_size)
                if not chunk:
                    break
                data += chunk
                if len(data) > max_bytes:
                    log_task_update(DEEP_RESEARCH_LABEL, f"Image too large (>{max_bytes // 1024}KB), skipping: {url[:80]}")
                    return False
        if len(data) < 512:
            log_task_update(DEEP_RESEARCH_LABEL, f"Image too small ({len(data)} bytes), skipping: {url[:80]}")
            return False
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        Path(save_path).write_bytes(data)
        log_task_update(DEEP_RESEARCH_LABEL, f"Downloaded image ({len(data) // 1024}KB) → {save_path}")
        return True
    except Exception as exc:
        log_task_update(DEEP_RESEARCH_LABEL, f"Image download failed for {url[:80]}: {exc}")
        return False


def _filter_relevant_images(candidates: list[dict], *, section_title: str, section_text: str, max_images: int = 3) -> list[dict]:
    """Use an LLM to select the most relevant diagrams/charts/visuals from image candidates."""
    if not candidates:
        return []
    candidates_summary = "\n".join(
        f"{i + 1}. title={item.get('title', '')!r} source={item.get('source', '')!r} "
        f"size={item.get('width', 0)}x{item.get('height', 0)} url={item.get('url', '')[:80]}"
        for i, item in enumerate(candidates)
    )
    prompt = f"""You are selecting relevant images for a deep research report section.

Section title: {section_title}
Section preview (first 800 chars):
{section_text[:800]}

Image candidates (from a Google Images search):
{candidates_summary}

Select up to {max_images} images that would genuinely enhance this section. Prioritize:
- Technical diagrams, architecture charts, system flows
- Data visualizations, graphs, charts, tables-as-images
- Infographics, process diagrams, comparison visuals
- Screenshots of relevant technical interfaces or metrics

REJECT:
- Stock photos, portraits, headshots
- Generic decorative images, logos, icons unrelated to the topic
- Low-resolution thumbnails (width < 200 or height < 200)
- Images with no clear relation to the section topic

Return ONLY a JSON array of the selected candidate numbers (1-based), e.g. [2, 5] or [] if none qualify.
Return an empty array if no images are clearly relevant."""
    try:
        raw = llm_text(prompt).strip()
        # Extract JSON array from response
        match = re.search(r"\[[\d,\s]*\]", raw)
        if not match:
            return []
        indices = json.loads(match.group())
        selected = []
        for idx in indices:
            if isinstance(idx, int) and 1 <= idx <= len(candidates):
                selected.append(candidates[idx - 1])
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"Image filter for '{section_title[:50]}': {len(selected)}/{len(candidates)} candidates selected.",
        )
        return selected[:max_images]
    except Exception as exc:
        log_task_update(DEEP_RESEARCH_LABEL, f"Image relevance filter failed: {exc}")
        return []


def _collect_section_images(
    *,
    section_title: str,
    section_objective: str,
    section_text: str,
    section_sources: list[dict],
    artifact_dir: str,
    section_index: int,
    max_images: int = 3,
    search_num: int = 12,
    fallback_to_search: bool = False,
    verify_with_vision: bool = True,
) -> list[dict]:
    """Collect, rank, and download images for a report section.

    Returns a list of dicts with keys:
        relative_path, abs_path, title, source, source_page, alt_text, section_index, section_title
    """
    source_urls = _source_image_urls(section_sources, max_pages=_IMAGE_SOURCE_PAGE_LIMIT)
    ranked_candidates: list[dict] = []
    rejection_counts = {"decorative": 0, "too_small": 0, "weak_context": 0, "low_relevance": 0, "vision_rejected": 0}
    if source_urls:
        for page_index, page_url in enumerate(source_urls, start=1):
            page_label = _web_source_log_label(page_url)
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Reviewing source images [{page_index}/{len(source_urls)}]: {page_label}",
            )
            try:
                payload = fetch_url_content(page_url, timeout=20)
            except Exception as exc:
                log_task_update(
                    DEEP_RESEARCH_LABEL,
                    f"Source image review failed [{page_index}/{len(source_urls)}]: {page_label} ({exc})",
                )
                continue
            page_candidates = _extract_source_page_image_candidates(page_url, payload, limit=_IMAGE_CANDIDATES_PER_PAGE)
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Found {len(page_candidates)} source image candidate(s) on {page_label}",
            )
            accepted_on_page = 0
            for candidate in page_candidates:
                score, reasons, rejection = _score_source_image_candidate(
                    candidate,
                    section_title=section_title,
                    section_objective=section_objective,
                    section_text=section_text,
                )
                if rejection:
                    rejection_counts[rejection] = rejection_counts.get(rejection, 0) + 1
                    continue
                accepted_on_page += 1
                ranked_candidates.append(
                    {
                        **candidate,
                        "candidate_kind": "source_page",
                        "relevance_score": score,
                        "selection_reason": "; ".join(reasons[:3]) if reasons else "context match",
                    }
                )
            if accepted_on_page:
                log_task_update(
                    DEEP_RESEARCH_LABEL,
                    f"Accepted {accepted_on_page} grounded image candidate(s) from {page_label}",
                )
        if any(rejection_counts.values()):
            rejection_summary = ", ".join(
                f"{label}={count}" for label, count in rejection_counts.items() if count
            )
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Rejected grounded image candidates for section {section_index}: {rejection_summary}",
            )
    else:
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"No eligible web source pages for grounded image collection in section {section_index}: '{section_title[:50]}'",
        )

    ranked_candidates.sort(key=lambda item: (-float(item.get("relevance_score", 0.0) or 0.0), str(item.get("url", ""))))
    selected = ranked_candidates[:max_images]

    if not selected and fallback_to_search:
        query = f"{section_title} diagram chart architecture infographic"
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"No grounded source images qualified for section {section_index}; falling back to external image search.",
        )
        candidates = _collect_image_search_results(query, num=search_num)
        selected = _filter_relevant_images(
            candidates,
            section_title=section_title,
            section_text=section_text,
            max_images=max_images,
        )
        selected = [
            {
                **item,
                "candidate_kind": "external_search",
                "selection_reason": "external fallback search",
            }
            for item in selected
        ]

    if not selected:
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"No grounded images selected for section {section_index}: '{section_title[:50]}'",
        )
        return []

    images_dir = _artifact_file(artifact_dir, f"section_{section_index:02d}/images")
    abs_images_dir = resolve_output_path(images_dir)
    os.makedirs(abs_images_dir, exist_ok=True)

    saved: list[dict] = []
    for img_index, img in enumerate(selected, start=1):
        img_url = img.get("url", "")
        if not img_url:
            continue
        # Derive a file extension from the URL or default to .jpg
        parsed_path = urlparse(img_url).path
        ext = Path(parsed_path).suffix.lower()
        if ext not in _IMAGE_EXTENSIONS:
            ext = ".jpg"
        filename = f"img_{section_index:02d}_{img_index:02d}{ext}"
        abs_path = os.path.join(abs_images_dir, filename)
        relative_path = f"section_{section_index:02d}/images/{filename}"

        if _download_image(img_url, abs_path):
            alt_text = (
                str(img.get("alt_text", "") or "").strip()
                or str(img.get("title", "") or "").strip()
                or str(img.get("figcaption", "") or "").strip()
                or f"{section_title} — visual {img_index}"
            )
            if verify_with_vision:
                is_relevant, vision_reason = _verify_downloaded_image_with_vision(
                    abs_path,
                    section_title=section_title,
                    section_objective=section_objective,
                    section_text=section_text,
                )
                if not is_relevant:
                    rejection_counts["vision_rejected"] = rejection_counts.get("vision_rejected", 0) + 1
                    try:
                        Path(abs_path).unlink(missing_ok=True)
                    except Exception:
                        pass
                    log_task_update(
                        DEEP_RESEARCH_LABEL,
                        f"Rejected downloaded image for section {section_index} via vision check: "
                        f"{_web_source_log_label(str(img.get('source_page', '') or img_url))} ({vision_reason or 'vision_rejected'})",
                    )
                    continue
            saved.append({
                "relative_path": relative_path,
                "abs_path": abs_path,
                "filename": filename,
                "title": str(img.get("title", "") or "").strip() or alt_text,
                "source": str(img.get("source", "") or "").strip() or urlparse(str(img.get("source_page", "") or img_url)).netloc,
                "source_page": str(img.get("source_page", "") or "").strip(),
                "url": img_url,
                "alt_text": alt_text,
                "section_index": section_index,
                "section_title": section_title,
                "width": int(img.get("width", 0) or 0),
                "height": int(img.get("height", 0) or 0),
                "source_page_title": str(img.get("page_title", "") or "").strip(),
                "figcaption": str(img.get("figcaption", "") or "").strip(),
                "selection_reason": str(img.get("selection_reason", "") or "").strip(),
                "relevance_score": float(img.get("relevance_score", 0.0) or 0.0),
                "candidate_kind": str(img.get("candidate_kind", "") or "source_page").strip(),
            })
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Selected image [{len(saved)}/{len(selected)}] for section {section_index}: "
                f"{_web_source_log_label(str(img.get('source_page', '') or img_url))} "
                f"({str(img.get('selection_reason', '') or 'context match')})",
            )

    log_task_update(
        DEEP_RESEARCH_LABEL,
        f"Section {section_index} images: {len(saved)}/{len(selected)} downloaded to {abs_images_dir}",
    )
    return saved


def _image_appendix_markdown(all_image_entries: list[dict]) -> str:
    """Build Appendix D — Image Gallery markdown from all collected section images."""
    lines = [_heading_with_anchor(2, "Appendix D - Image Gallery", "appendix-d-image-gallery"), ""]
    if not all_image_entries:
        lines.append("_No images were collected during research._")
        return "\n".join(lines)

    lines.append(
        f"This appendix contains {len(all_image_entries)} image(s) collected from the web during research. "
        "Images are grouped by the report section they support."
    )
    lines.append("")

    # Group by section
    by_section: dict[int, list[dict]] = {}
    for entry in all_image_entries:
        idx = int(entry.get("section_index", 0))
        by_section.setdefault(idx, []).append(entry)

    for section_idx in sorted(by_section):
        entries = by_section[section_idx]
        section_title = entries[0].get("section_title", f"Section {section_idx}")
        lines.append(f"### Section {section_idx}: {section_title}")
        lines.append("")
        for i, entry in enumerate(entries, start=1):
            rel_path = entry.get("relative_path", "")
            alt = entry.get("alt_text", "") or entry.get("title", f"Image {i}")
            source = entry.get("source", "")
            source_page = entry.get("source_page", "")
            lines.append(f"![{alt}]({rel_path})")
            caption_parts = [f"**Figure {section_idx}.{i}** — {alt}"]
            if source:
                caption_parts.append(f"Source: {source}")
            if source_page:
                caption_parts.append(f"([link]({source_page}))")
            lines.append(" ".join(caption_parts))
            lines.append("")

    return "\n".join(lines)


def _embed_images_in_section(section_text: str, images: list[dict]) -> str:
    """Append downloaded image references at the end of a section's markdown."""
    if not images:
        return section_text
    lines = [section_text.rstrip(), "", "**Visual References**", ""]
    for i, entry in enumerate(images, start=1):
        rel_path = entry.get("relative_path", "")
        alt = entry.get("alt_text", "") or entry.get("title", f"Figure {i}")
        source = entry.get("source", "")
        lines.append(f"![{alt}]({rel_path})")
        if source:
            lines.append(f"*Source: {source}*")
        lines.append("")
    return "\n".join(lines)


def _sources_from_search_results(results: list[dict], *, prefix: str = "S") -> list[dict]:
    entries = []
    for index, item in enumerate(results or [], start=1):
        url = str(item.get("url", "")).strip()
        if not url:
            continue
        label = str(item.get("title", "")).strip() or _source_label(url)
        entries.append({"id": f"{prefix}{index}", "url": url, "label": label})
    return entries


def _search_results_markdown(results: list[dict]) -> str:
    if not results:
        return "No web search results were available."
    lines = ["Web search results:"]
    for item in results:
        title = str(item.get("title", "")).strip() or "Untitled"
        url = str(item.get("url", "")).strip()
        snippet = str(item.get("snippet", "")).strip()
        evidence_excerpt = str(item.get("evidence_excerpt", "")).strip()
        line = f"- {title}"
        if snippet:
            line += f": {snippet}"
        if url:
            line += f" ({url})"
        lines.append(line)
        if evidence_excerpt:
            lines.append(f"  - Viewed evidence: {evidence_excerpt}")
    return "\n".join(lines)


def _merge_sources(primary: list[dict], secondary: list[dict], *, max_sources: int = 60) -> list[dict]:
    seen: set[str] = set()
    merged: list[dict] = []
    for item in (primary or []) + (secondary or []):
        url = str(item.get("url", "")).strip()
        path_value = str(item.get("path", "")).strip()
        source_id = str(item.get("source_id", "")).strip()
        dedupe_key = source_id or path_value or url
        if not dedupe_key or dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        merged.append(item)
        if len(merged) >= max_sources:
            break
    # Re-number sequentially
    renumbered = []
    for index, item in enumerate(merged, start=1):
        ref = str(item.get("url", "")).strip() or str(item.get("path", "")).strip()
        renumbered.append(
            {
                **item,
                "id": f"S{index}",
                "url": ref,
                "label": item.get("label", "") or _source_label(ref),
            }
        )
    return renumbered


def _format_evidence_bank_markdown(
    *,
    objective: str,
    local_entries: list[dict],
    url_entries: list[dict],
    kb_grounding: Mapping[str, Any],
    kb_warning: str,
    search_results: dict,
    research_notes_text: str,
    source_ledger: list[dict],
    insufficiency_note: str,
) -> str:
    lines = ["# Evidence Bank", "", "## Objective", objective.strip(), ""]
    if insufficiency_note:
        lines.extend(["## Data Sufficiency Note", insufficiency_note.strip(), ""])
    lines.append("## Local File Sources")
    if local_entries:
        for entry in local_entries:
            label = entry.get("file_name") or entry.get("path") or "unknown"
            lines.append(f"- File: {label}")
            if entry.get("path"):
                lines.append(f"  - Path: {entry['path']}")
            if entry.get("type"):
                lines.append(f"  - Type: {entry['type']}")
            if entry.get("char_count"):
                lines.append(f"  - Characters: {entry['char_count']}")
            if entry.get("error"):
                lines.append(f"  - Extraction error: {entry['error']}")
            if entry.get("summary"):
                lines.append(f"  - Summary: {entry['summary']}")
    else:
        lines.append("- No local-drive sources were found.")
    lines.append("")

    lines.append("## User-Provided URL Sources")
    if url_entries:
        for entry in url_entries:
            label = entry.get("label") or entry.get("url") or "URL"
            lines.append(f"- URL: {label}")
            if entry.get("url"):
                lines.append(f"  - Link: {entry['url']}")
            if entry.get("content_type"):
                lines.append(f"  - Content type: {entry['content_type']}")
            if entry.get("char_count"):
                lines.append(f"  - Characters: {entry['char_count']}")
            if entry.get("error"):
                lines.append(f"  - Extraction error: {entry['error']}")
            if entry.get("summary"):
                lines.append(f"  - Summary: {entry['summary']}")
    else:
        lines.append("- No explicit URL sources were provided.")
    lines.append("")

    lines.append("## Knowledge Base Sources")
    if kb_grounding and int(kb_grounding.get("hit_count", 0) or 0) > 0:
        lines.append(f"- KB: {kb_grounding.get('kb_name', 'unknown')}")
        lines.append(f"- Retrieved hits: {int(kb_grounding.get('hit_count', 0) or 0)}")
        for item in _sources_from_kb_grounding(kb_grounding)[:12]:
            lines.append(f"  - {item.get('label', '')} ({item.get('url', '')})")
    elif kb_warning:
        lines.append(f"- KB warning: {kb_warning}")
    else:
        lines.append("- No knowledge base grounding was used.")
    lines.append("")

    lines.append("## Web Search Results (Google via SerpAPI)")
    results = search_results.get("results", []) if isinstance(search_results, dict) else []
    if results:
        for item in results:
            title = item.get("title") or "Untitled"
            url = item.get("url") or ""
            snippet = item.get("snippet") or ""
            lines.append(f"- {title}")
            if url:
                lines.append(f"  - URL: {url}")
            if snippet:
                lines.append(f"  - Snippet: {snippet}")
    else:
        error = ""
        if isinstance(search_results, dict):
            error = str(search_results.get("error", "")).strip()
        if error:
            lines.append(f"- Search unavailable: {error}")
        else:
            lines.append("- No search results were collected.")
    lines.append("")

    lines.append("## Research Notes")
    lines.append(research_notes_text.strip() or "No research notes generated.")
    lines.append("")
    lines.append("## Source Ledger")
    if source_ledger:
        for item in source_ledger:
            lines.append(f"- [{item['id']}] {item['label']} - {item['url']}")
    else:
        lines.append("- No sources were extracted.")
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def _write_absolute_text_file(path_value: str, content: str) -> None:
    path = Path(path_value)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def _run_long_document_addendum(
    state: dict,
    *,
    objective: str,
    call_number: int,
    artifact_dir: str,
) -> dict:
    compiled_path = _resolve_existing_output_path(state.get("long_document_compiled_path", ""))
    if not compiled_path or not Path(compiled_path).exists():
        summary = "Addendum requested but compiled document was not found. Skipping addendum."
        state["draft_response"] = summary
        state["_skip_review_once"] = True
        return publish_agent_output(
            state,
            "long_document_agent",
            summary,
            f"long_document_addendum_{call_number}",
            recipients=["orchestrator_agent", "reviewer_agent", "report_agent"],
        )

    attempts = int(state.get("long_document_addendum_attempts", 0) or 0) + 1
    max_attempts = int(state.get("long_document_addendum_max_attempts", 1) or 1)
    state["long_document_addendum_attempts"] = attempts
    if attempts > max_attempts:
        summary = f"Addendum requested but max attempts ({max_attempts}) reached. Skipping."
        state["draft_response"] = summary
        state["long_document_addendum_requested"] = False
        state["long_document_addendum_completed"] = True
        state["_skip_review_once"] = True
        return publish_agent_output(
            state,
            "long_document_agent",
            summary,
            f"long_document_addendum_{call_number}",
            recipients=["orchestrator_agent", "reviewer_agent", "report_agent"],
        )

    instructions = str(state.get("long_document_addendum_instructions") or objective).strip()
    evidence_excerpt = (
        str(state.get("long_document_evidence_bank_excerpt") or "").strip()
        or _read_text_file(state.get("long_document_evidence_bank_path", ""), "")
    )
    compiled_excerpt = _trim_text(_read_text_file(compiled_path, 12000), 12000)
    words_target = _safe_int(state.get("long_document_addendum_words"), 1200, 400, 5000)

    log_task_update(
        DEEP_RESEARCH_LABEL,
        f"Generating addendum (attempt {attempts + 1}) targeting ~{words_target} words.",
    )

    prompt = f"""
You are drafting a focused addendum to a deep research financial advisory report.
Use the reviewer feedback and evidence bank to add missing analysis without rewriting the full report.

Primary objective:
{objective}

Reviewer feedback / required additions:
{instructions}

Evidence bank excerpt:
{_trim_text(evidence_excerpt, 6000)}

Excerpt from the existing report:
{compiled_excerpt}

Write an addendum of about {words_target} words that:
- Directly addresses the missing market dynamics and geographic insights.
- Cites gaps explicitly where data is unavailable.
- Includes a short "Addendum Takeaways" list.
"""
    addendum_text = llm_text(prompt).strip()
    if not addendum_text:
        addendum_text = "Addendum could not be generated from the available evidence."
    log_task_update(
        DEEP_RESEARCH_LABEL,
        f"Addendum draft complete: {len(addendum_text.split())} words, {len(addendum_text)} chars.",
    )

    addendum_title = "## Addendum: Reviewer Follow-ups"
    addendum_body = f"{addendum_title}\n\n{addendum_text}\n"
    compiled_existing = _read_text_file(compiled_path, "")
    compiled_new = f"{compiled_existing.rstrip()}\n\n{addendum_body}"
    _write_absolute_text_file(compiled_path, compiled_new)

    addendum_filename = _artifact_file(artifact_dir, f"long_document_addendum_{attempts}.md")
    write_text_file(addendum_filename, addendum_body)

    manifest_path = _resolve_existing_output_path(state.get("long_document_manifest_path", ""))
    if manifest_path and Path(manifest_path).exists():
        try:
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
        except Exception:
            manifest = {}
        if isinstance(manifest, dict):
            manifest["addendum_file"] = _output_file_path(addendum_filename)
            manifest["addendum_attempts"] = attempts
            _write_absolute_text_file(manifest_path, json.dumps(manifest, indent=2, ensure_ascii=False))

    state["long_document_addendum_requested"] = False
    state["long_document_addendum_completed"] = True
    state["long_document_addendum_path"] = _output_file_path(addendum_filename)
    state["_skip_review_once"] = True
    if bool(state.get("long_document_addendum_force_no_review", True)):
        state["skip_reviews"] = True
    summary = (
        "Addendum generated and appended to compiled report.\n"
        f"- Compiled markdown: {compiled_path}\n"
        f"- Addendum file: {_output_file_path(addendum_filename)}\n"
    )
    state["draft_response"] = summary
    log_task_update(DEEP_RESEARCH_LABEL, "Addendum appended to compiled report.", summary)
    return publish_agent_output(
        state,
        "long_document_agent",
        summary,
        f"long_document_addendum_{call_number}",
        recipients=["orchestrator_agent", "reviewer_agent", "report_agent"],
    )


def _normalize_url(raw: str) -> str:
    value = str(raw or "").strip()
    while value and value[-1] in {".", ",", ";", ":", ")", "]", "}", "\"", "'"}:
        value = value[:-1]
    return value


def _extract_urls_from_text(text: str) -> list[str]:
    if not str(text or "").strip():
        return []
    matches = re.findall(r"https?://[^\s<>\"]+", text)
    urls: list[str] = []
    for item in matches:
        normalized = _normalize_url(item)
        if normalized and normalized not in urls:
            urls.append(normalized)
    return urls


def _extract_urls_from_obj(value: Any, urls: list[str], limit: int = 120) -> None:
    if len(urls) >= limit:
        return
    if isinstance(value, dict):
        for item in value.values():
            _extract_urls_from_obj(item, urls, limit=limit)
            if len(urls) >= limit:
                return
        return
    if isinstance(value, list):
        for item in value:
            _extract_urls_from_obj(item, urls, limit=limit)
            if len(urls) >= limit:
                return
        return
    if isinstance(value, str):
        for url in _extract_urls_from_text(value):
            if url not in urls:
                urls.append(url)
                if len(urls) >= limit:
                    return


def _extract_source_entries(research_pass: dict, *, max_sources: int = 30) -> list[dict]:
    urls: list[str] = []
    _extract_urls_from_obj(research_pass.get("raw"), urls)
    for item in _extract_urls_from_text(str(research_pass.get("output_text", "") or "")):
        if item not in urls:
            urls.append(item)
    entries = []
    for index, url in enumerate(urls[:max_sources], start=1):
        entries.append(
            {
                "id": f"S{index}",
                "url": url,
                "label": _source_label(url),
            }
        )
    return entries


def _source_ledger_markdown(entries: list[dict]) -> str:
    if not entries:
        return "- No externally verifiable URLs were extracted from this section research pass."
    return "\n".join(f"- [{item['id']}] {item['label']} - {item['url']}" for item in entries)


def _references_markdown(entries: list[dict], *, heading: str = "### Verified References", limit: int = 20) -> str:
    lines = [heading, ""]
    if not entries:
        lines.append("- No external references were extracted for this section.")
        return "\n".join(lines).strip()
    for item in entries[:limit]:
        lines.append(f"- [{item['id']}] {item['label']} - {item['url']}")
    return "\n".join(lines).strip()


def _append_verified_references(section_text: str, entries: list[dict]) -> str:
    text = (section_text or "").rstrip()
    references_block = _references_markdown(entries)
    return f"{text}\n\n{references_block}\n"

def _strip_section_references(section_text: str) -> str:
    text = str(section_text or "")
    match = re.search(r"\n#{2,3}\\s+References\\b", text, flags=re.IGNORECASE)
    if match:
        return text[: match.start()].rstrip()
    return text.rstrip()


def _remap_section_citations(section_text: str, section_sources: list[dict], consolidated: list[dict]) -> str:
    url_to_global = {str(item.get("url", "")).strip(): str(item.get("id", "")).strip() for item in consolidated}
    local_map: dict[str, str] = {}
    for item in section_sources:
        local_id = str(item.get("id", "")).strip()
        url = str(item.get("url", "")).strip()
        global_id = url_to_global.get(url)
        if local_id and global_id:
            local_map[local_id] = global_id

    if not local_map:
        return section_text

    def _swap(match: re.Match) -> str:
        local_id = match.group(1)
        return f"[{local_map.get(local_id, local_id)}]"

    # Match any source citation tag of the form [S1], [S12], [P1], etc.
    return re.sub(r"\[([A-Z]\d+)\]", _swap, section_text)


def _strip_leading_section_heading(section_text: str, *, section_title: str, section_index: int) -> str:
    lines = str(section_text or "").splitlines()
    first_content_index = 0
    while first_content_index < len(lines) and not lines[first_content_index].strip():
        first_content_index += 1
    if first_content_index >= len(lines):
        return ""

    first_line = lines[first_content_index].strip()
    normalized_first = _normalized_heading_candidate(first_line)
    heading_variants = {
        _normalized_heading_candidate(section_title),
        _normalized_heading_candidate(f"{section_index}. {section_title}"),
        _normalized_heading_candidate(f"Section {section_index}: {section_title}"),
    }
    if normalized_first and normalized_first in heading_variants:
        next_index = first_content_index + 1
        while next_index < len(lines) and not lines[next_index].strip():
            next_index += 1
        return "\n".join(lines[next_index:]).strip()
    return str(section_text or "").strip()


def _consolidate_references(section_outputs: list[dict]) -> list[dict]:
    combined: list[dict] = []
    seen: set[str] = set()
    for section in section_outputs:
        for item in section.get("references", []):
            url = str(item.get("url", "")).strip()
            if not url or url in seen:
                continue
            seen.add(url)
            combined.append({"url": url, "label": str(item.get("label", "")).strip() or _source_label(url)})
    consolidated = []
    for index, item in enumerate(combined, start=1):
        consolidated.append({"id": f"R{index}", "url": item["url"], "label": item["label"]})
    return consolidated


def _extract_mermaid_blocks(text: str) -> list[str]:
    blocks = []
    pattern = re.compile(r"```mermaid\s*(.*?)```", flags=re.IGNORECASE | re.DOTALL)
    for match in pattern.finditer(text or ""):
        block = str(match.group(1) or "").strip()
        if block and block not in blocks:
            blocks.append(block)
    return blocks


def _mermaid_to_dot(mermaid_text: str) -> str:
    """Convert a Mermaid flowchart to Graphviz DOT notation."""
    lines = [l.strip() for l in (mermaid_text or "").strip().splitlines()]
    rankdir = "TB"
    nodes: dict[str, dict] = {}
    edges: list[tuple] = []
    data_lines: list[str] = []

    for line in lines:
        if not line or line.startswith("%%"):
            continue
        lower = line.lower()
        if lower.startswith("flowchart") or lower.startswith("graph"):
            parts = line.split()
            if len(parts) > 1:
                rankdir = {"TD": "TB", "TB": "TB", "LR": "LR", "RL": "RL", "BT": "BT"}.get(parts[1].upper(), "TB")
        else:
            data_lines.append(line)

    def _parse_node(token: str) -> tuple[str, str, str]:
        token = token.strip()
        for pat, shape in [
            (r'^(\w+)\{([^}]*)\}$', "diamond"),
            (r'^(\w+)\(\(([^)]*)\)\)$', "ellipse"),
            (r'^(\w+)\[([^\]]*)\]$', "box"),
            (r'^(\w+)>([^\]]*)\]$', "box"),
        ]:
            m = re.match(pat, token)
            if m:
                return m.group(1), m.group(2), shape
        m = re.match(r'^(\w+)$', token)
        if m:
            return m.group(1), m.group(1), "box"
        return token, token, "box"

    edge_re = re.compile(r'(-->|-\.->|==>|---)')
    label_re = re.compile(r'^\|([^|]*)\|')

    for line in data_lines:
        parts = edge_re.split(line, maxsplit=1)
        if len(parts) == 3:
            left = parts[0].strip()
            right = parts[2].strip()
            edge_label = ""
            lm = label_re.match(right)
            if lm:
                edge_label = lm.group(1).strip()
                right = right[lm.end():].strip()
            fid, flabel, fshape = _parse_node(left)
            tid, tlabel, tshape = _parse_node(right)
            if fid not in nodes:
                nodes[fid] = {"label": flabel, "shape": fshape}
            elif flabel != fid:
                nodes[fid].update({"label": flabel, "shape": fshape})
            if tid not in nodes:
                nodes[tid] = {"label": tlabel, "shape": tshape}
            elif tlabel != tid:
                nodes[tid].update({"label": tlabel, "shape": tshape})
            edges.append((fid, tid, edge_label))
        else:
            nid, nlabel, nshape = _parse_node(line)
            if nid not in nodes:
                nodes[nid] = {"label": nlabel, "shape": nshape}

    if not nodes and not edges:
        return ""

    def _esc(s: str) -> str:
        return s.replace("\\", "\\\\").replace('"', '\\"')

    dot = [
        "digraph G {",
        f"    rankdir={rankdir};",
        '    node [fontname="Helvetica", fontsize=11, style=filled, fillcolor="#f0f4ff", color="#4a6fa5"];',
        '    edge [fontname="Helvetica", fontsize=10, color="#555555"];',
        "",
    ]
    for nid, attrs in nodes.items():
        dot.append(f'    {nid} [label="{_esc(attrs.get("label", nid))}", shape={attrs.get("shape", "box")}];')
    dot.append("")
    for fid, tid, label in edges:
        label_attr = f' [label="{_esc(label)}"]' if label else ""
        dot.append(f"    {fid} -> {tid}{label_attr};")
    dot.append("}")
    return "\n".join(dot)


def _render_mermaid_png(mermaid_text: str, output_path: str) -> bool:
    """Render a Mermaid diagram to a PNG file.
    Tries mmdc (Mermaid CLI) first, then graphviz as fallback.
    Returns True if a PNG was successfully written."""
    output_path = str(output_path)

    # --- attempt 1: mmdc (npm install -g @mermaid-js/mermaid-cli) ---
    try:
        mmd_tmp = output_path.replace(".png", ".mmd")
        Path(mmd_tmp).write_text(mermaid_text, encoding="utf-8")
        result = subprocess.run(
            ["mmdc", "-i", mmd_tmp, "-o", output_path, "--backgroundColor", "white"],
            capture_output=True, timeout=30,
        )
        if result.returncode == 0 and Path(output_path).exists():
            log_task_update(DEEP_RESEARCH_LABEL, f"Rendered flowchart via mmdc → {output_path}")
            return True
    except Exception:
        pass

    # --- attempt 2: graphviz Python package ---
    try:
        import graphviz  # type: ignore
        dot_src = _mermaid_to_dot(mermaid_text)
        if not dot_src:
            return False
        base = output_path[:-4] if output_path.endswith(".png") else output_path
        src = graphviz.Source(dot_src, format="png")
        rendered = src.render(filename=base, cleanup=True)
        # graphviz appends format extension; rename if needed
        rendered_path = Path(rendered)
        target_path = Path(output_path)
        if rendered_path != target_path and rendered_path.exists():
            rendered_path.rename(target_path)
        if target_path.exists():
            log_task_update(DEEP_RESEARCH_LABEL, f"Rendered flowchart via graphviz → {output_path}")
            return True
    except Exception as exc:
        log_task_update(DEEP_RESEARCH_LABEL, f"Flowchart PNG render failed: {exc}")

    return False


def _replace_mermaid_with_png(text: str, artifact_dir: str, section_index: int) -> tuple[str, list[str]]:
    """Replace all ```mermaid...``` blocks in *text* with PNG image embeds.
    Returns (updated_text, list_of_png_absolute_paths)."""
    pattern = re.compile(r"```mermaid\s*(.*?)```", flags=re.IGNORECASE | re.DOTALL)
    png_files: list[str] = []
    counter = [0]

    def _replace(match: re.Match) -> str:
        mermaid_src = match.group(1).strip()
        if not mermaid_src:
            return match.group(0)
        counter[0] += 1
        png_rel = _artifact_file(artifact_dir, f"section_{section_index:02d}/flowchart_{counter[0]:02d}.png")
        png_abs = resolve_output_path(png_rel)
        os.makedirs(os.path.dirname(png_abs), exist_ok=True)
        if _render_mermaid_png(mermaid_src, png_abs):
            png_files.append(png_abs)
            # Use a relative path from the artifact_dir root so it works in both MD and HTML
            rel_from_report = f"section_{section_index:02d}/flowchart_{counter[0]:02d}.png"
            return f"![Flowchart {counter[0]}]({rel_from_report})"
        # Fallback: keep original Mermaid block
        return match.group(0)

    updated = pattern.sub(_replace, text)
    return updated, png_files


def _extract_markdown_tables(text: str) -> list[str]:
    lines = (text or "").splitlines()
    tables = []
    index = 0
    while index < len(lines):
        line = lines[index]
        if "|" not in line or line.strip().startswith("```"):
            index += 1
            continue
        start = index
        block = []
        while index < len(lines):
            current = lines[index]
            if "|" not in current or current.strip().startswith("```"):
                break
            block.append(current.rstrip())
            index += 1
        if len(block) >= 2 and re.match(r"^\s*\|?[\s:\-\|]+\|?\s*$", block[1]):
            table = "\n".join(block).strip()
            if table not in tables:
                tables.append(table)
        if index == start:
            index += 1
    return tables


def _generate_visual_assets(section_title: str, section_text: str, research_text: str) -> dict:
    prompt = f"""
You are a technical writing visual-design assistant.

From the section below, decide whether a markdown table and/or mermaid flowchart would improve clarity.
Generate visuals only if they add real value.

Section title:
{section_title}

Section text:
{_trim_text(section_text, 10000)}

Research notes:
{_trim_text(research_text, 9000)}

Return ONLY valid JSON:
{{
  "tables": [
    {{
      "title": "short title",
      "markdown": "| Col A | Col B |\\n|---|---|\\n| ... | ... |"
    }}
  ],
  "flowcharts": [
    {{
      "title": "short title",
      "mermaid": "flowchart TD\\nA --> B"
    }}
  ],
  "notes": "brief note"
}}
"""
    return llm_json(prompt, {"tables": [], "flowcharts": [], "notes": ""})


def _normalize_visual_assets(existing_tables: list[str], existing_flowcharts: list[str], generated: dict) -> dict:
    tables = []
    flowcharts = []

    for index, table_md in enumerate(existing_tables, start=1):
        cleaned = str(table_md or "").strip()
        if not cleaned:
            continue
        tables.append({"title": f"Existing Table {index}", "markdown": cleaned, "source": "existing"})

    for index, mermaid in enumerate(existing_flowcharts, start=1):
        cleaned = str(mermaid or "").strip()
        if not cleaned:
            continue
        flowcharts.append({"title": f"Existing Flowchart {index}", "mermaid": cleaned, "source": "existing"})

    if isinstance(generated, dict):
        for index, item in enumerate(generated.get("tables", []) if isinstance(generated.get("tables", []), list) else [], start=1):
            if not isinstance(item, dict):
                continue
            markdown = str(item.get("markdown", "")).strip()
            if "|" not in markdown or markdown in {entry["markdown"] for entry in tables}:
                continue
            title = str(item.get("title", "")).strip() or f"Generated Table {index}"
            tables.append({"title": title, "markdown": markdown, "source": "generated"})

        for index, item in enumerate(generated.get("flowcharts", []) if isinstance(generated.get("flowcharts", []), list) else [], start=1):
            if not isinstance(item, dict):
                continue
            mermaid = str(item.get("mermaid", "")).strip()
            if not mermaid:
                continue
            lower = mermaid.lower()
            if not (lower.startswith("flowchart") or lower.startswith("graph")):
                continue
            if mermaid in {entry["mermaid"] for entry in flowcharts}:
                continue
            title = str(item.get("title", "")).strip() or f"Generated Flowchart {index}"
            flowcharts.append({"title": title, "mermaid": mermaid, "source": "generated"})

    return {"tables": tables[:4], "flowcharts": flowcharts[:4], "notes": str((generated or {}).get("notes", "")).strip()}


def _render_visual_assets_md(visual_assets: dict) -> str:
    lines = ["### Visual Assets", ""]
    tables = visual_assets.get("tables", [])
    flowcharts = visual_assets.get("flowcharts", [])
    if not tables and not flowcharts:
        lines.append("- No additional visual assets were generated for this section.")
        return "\n".join(lines).strip()

    if tables:
        lines.append("#### Tables")
        lines.append("")
        for item in tables:
            lines.append(f"##### {item.get('title', 'Table')}")
            lines.append("")
            lines.append(str(item.get("markdown", "")).strip())
            lines.append("")

    if flowcharts:
        lines.append("#### Flowcharts")
        lines.append("")
        for item in flowcharts:
            lines.append(f"##### {item.get('title', 'Flowchart')}")
            lines.append("")
            png_path = str(item.get("png_path", "")).strip()
            if png_path:
                lines.append(f"![{item.get('title', 'Flowchart')}]({png_path})")
            else:
                lines.append("```mermaid")
                lines.append(str(item.get("mermaid", "")).strip())
                lines.append("```")
            lines.append("")

    return "\n".join(lines).strip()


def _append_generated_visuals(section_text: str, visual_assets: dict) -> str:
    generated_tables = [item for item in visual_assets.get("tables", []) if item.get("source") == "generated"]
    generated_flowcharts = [item for item in visual_assets.get("flowcharts", []) if item.get("source") == "generated"]

    if not generated_tables and not generated_flowcharts:
        return section_text

    generated_assets = {"tables": generated_tables, "flowcharts": generated_flowcharts}
    rendered = _render_visual_assets_md(generated_assets)
    return f"{section_text.rstrip()}\n\n{rendered}\n"


def _build_visual_index(section_outputs: list[dict]) -> dict:
    entries = []
    for section in section_outputs:
        visual_assets = section.get("visual_assets", {}) if isinstance(section.get("visual_assets"), dict) else {}
        table_count = len(visual_assets.get("tables", []))
        flowchart_count = len(visual_assets.get("flowcharts", []))
        entries.append(
            {
                "section_index": section.get("index"),
                "section_title": section.get("title", ""),
                "tables": table_count,
                "flowcharts": flowchart_count,
                "visual_assets_file": section.get("visual_assets_file", ""),
            }
        )
    return {"entries": entries}


def _visual_index_markdown(visual_index: dict) -> str:
    entries = visual_index.get("entries", []) if isinstance(visual_index, dict) else []
    lines = ["## Visual Asset Index", ""]
    if not entries:
        lines.append("- No visual assets were captured.")
        return "\n".join(lines).strip() + "\n"

    lines.extend(
        [
            "| Section | Tables | Flowcharts | Asset File |",
            "|---|---:|---:|---|",
        ]
    )
    for item in entries:
        section_label = f"{item.get('section_index', '?')}. {str(item.get('section_title', '')).strip() or 'Section'}"
        tables = int(item.get("tables", 0) or 0)
        flowcharts = int(item.get("flowcharts", 0) or 0)
        asset_file = str(item.get("visual_assets_file", "")).strip() or "-"
        lines.append(f"| {section_label} | {tables} | {flowcharts} | {asset_file} |")
    return "\n".join(lines).strip() + "\n"


def _build_outline(objective: str, *, title: str, section_count: int, section_pages: int, subtopics: list[str] | None = None) -> dict:
    subtopic_list = [str(item).strip() for item in (subtopics or []) if str(item).strip()]
    fallback = {
        "title": title,
        "sections": [
            {
                "id": index + 1,
                "title": f"Section {index + 1}",
                "objective": objective,
                "key_questions": [],
                "target_pages": section_pages,
            }
            for index in range(section_count)
        ],
    }
    prompt = f"""
You are planning a deep research report.

Create a coherent section-by-section outline for this objective:
{objective}

Constraints:
- report title: {title}
- sections needed: {section_count}
- target pages per section: about {section_pages}
- preferred subtopics: {subtopic_list or ['Use the objective to derive subtopics']}
- each section must contribute to one coherent final narrative
- avoid overlap; ensure logical progression
- include a methodology / data sources section near the beginning
- place the conclusion last

Return ONLY valid JSON with schema:
{{
  "title": "string",
  "sections": [
    {{
      "id": 1,
      "title": "string",
      "objective": "string",
      "key_questions": ["string"],
      "target_pages": {section_pages}
    }}
  ]
}}
"""
    data = llm_json(prompt, fallback)
    if not isinstance(data, dict):
        return fallback
    sections = data.get("sections")
    if not isinstance(sections, list) or not sections:
        return fallback
    normalized_sections = []
    for index, section in enumerate(sections[:section_count]):
        if not isinstance(section, dict):
            continue
        questions = section.get("key_questions", [])
        if not isinstance(questions, list):
            questions = []
        normalized_sections.append(
            {
                "id": _safe_int(section.get("id", index + 1), index + 1, 1, 500),
                "title": str(section.get("title", f"Section {index + 1}")).strip() or f"Section {index + 1}",
                "objective": str(section.get("objective", objective)).strip() or objective,
                "key_questions": [str(item).strip() for item in questions if str(item).strip()][:12],
                "target_pages": _safe_int(section.get("target_pages", section_pages), section_pages, 1, 40),
            }
        )
    if not normalized_sections:
        return fallback
    methodology_title = "Methodology and Data Sources"
    conclusion_title = "Conclusion and Implications"
    titles = [str(item.get("title", "")).strip().lower() for item in normalized_sections]
    if methodology_title.lower() not in titles:
        normalized_sections.insert(
            min(1, len(normalized_sections)),
            {
                "id": len(normalized_sections) + 1,
                "title": methodology_title,
                "objective": "Explain the research method, source mix, date range, and evidence quality controls.",
                "key_questions": [],
                "target_pages": section_pages,
            },
        )
    titles = [str(item.get("title", "")).strip().lower() for item in normalized_sections]
    if conclusion_title.lower() not in titles:
        normalized_sections.append(
            {
                "id": len(normalized_sections) + 1,
                "title": conclusion_title,
                "objective": "Synthesize the key findings, limitations, and next-step implications.",
                "key_questions": [],
                "target_pages": section_pages,
            }
        )
    for index, item in enumerate(normalized_sections, start=1):
        item["id"] = index
    while len(normalized_sections) < section_count:
        next_index = len(normalized_sections) + 1
        normalized_sections.append(
            {
                "id": next_index,
                "title": f"Section {next_index}",
                "objective": objective,
                "key_questions": [],
                "target_pages": section_pages,
            }
        )
    return {"title": _normalize_title(data.get("title", ""), title), "sections": normalized_sections}


def _outline_markdown(outline: dict, *, fallback_title: str) -> str:
    lines = [f"# {_normalize_title(outline.get('title', ''), fallback_title)}", "", "## Section Outline"]
    sections = outline.get("sections", []) if isinstance(outline, dict) else []
    for section in sections:
        if not isinstance(section, dict):
            continue
        questions = section.get("key_questions", [])
        question_text = "; ".join(str(item).strip() for item in questions[:6] if str(item).strip()) or "n/a"
        lines.append(
            f"- {section.get('id')}. {section.get('title')}: {section.get('objective')} "
            f"| target_pages={section.get('target_pages')} | questions={question_text}"
        )
    return "\n".join(lines).strip() + "\n"


def _deep_research_analysis_markdown(
    analysis: dict,
    *,
    formats: list[str],
    citation_style: str,
    plagiarism_enabled: bool,
    web_search_enabled: bool,
    local_source_count: int,
    provided_url_count: int,
) -> str:
    budget = analysis.get("execution_budget", analysis.get("budget", {})) if isinstance(analysis, dict) else {}
    depth_label = str(analysis.get("depth_label", "") or "Standard Report").strip()
    depth_description = str(analysis.get("depth_description", "") or "").strip()

    def _budget_value(key: str, *, suffix: str = "") -> str:
        raw = 0
        if isinstance(budget, dict):
            raw = int(budget.get(key, 0) or 0)
        if raw > 0:
            return f"{raw}{suffix}"
        return "not explicitly capped"

    lines = [
        "# Deep Research Analysis",
        "",
        f"- Tier: {analysis.get('tier', '?')}",
        f"- Research depth: {depth_label}",
        f"- Estimated sources: {analysis.get('estimated_sources', '?')}",
        f"- Estimated duration: {analysis.get('estimated_duration_minutes', '?')} minutes",
        f"- Citation style: {citation_style.upper()}",
        f"- Output formats: {', '.join(fmt.upper() for fmt in formats)}",
        f"- Plagiarism check: {'enabled' if plagiarism_enabled else 'disabled'}",
        f"- Web search: {'enabled' if web_search_enabled else 'disabled'}",
        f"- Local file sources detected: {local_source_count}",
        f"- User-provided URLs: {provided_url_count}",
        f"- Date range: {analysis.get('date_range', 'all_time')}",
    ]
    if depth_description:
        lines.append(f"- Scope profile: {depth_description}")
    lines.extend(["", "## Detected Subtopics"])
    for idx, topic in enumerate(analysis.get("subtopics", []) if isinstance(analysis.get("subtopics", []), list) else [], start=1):
        lines.append(f"- {idx}. {topic}")
    if budget:
        lines.extend(
            [
                "",
                "## Session Budget",
                f"- Max tokens: {_budget_value('max_tokens')}",
                f"- Max sources: {_budget_value('max_sources')}",
                f"- Max duration: {_budget_value('max_duration_minutes', suffix=' minutes')}",
            ]
        )
    if str(analysis.get("reason", "")).strip():
        lines.extend(["", "## Why this tier", f"- {str(analysis.get('reason', '')).strip()}"])
    return "\n".join(lines).strip() + "\n"


def _compact_text(value: str, *, width: int = 120) -> str:
    text = re.sub(r"\s+", " ", str(value or "").strip())
    if not text:
        return ""
    return textwrap.shorten(text, width=width, placeholder="...")


def _deep_research_analysis_summary_prompt(
    *,
    title: str,
    analysis: dict,
    formats: list[str],
    citation_style: str,
    plagiarism_enabled: bool,
    web_search_enabled: bool,
    local_source_count: int,
    provided_url_count: int,
    analysis_storage_path: str,
    version: int,
) -> str:
    depth_label = str(analysis.get("depth_label", "") or "Standard Report").strip()
    depth_description = str(analysis.get("depth_description", "") or "").strip()
    lines = [
        f"The deep research confirmation v{version} is ready for approval.",
        "Review this compact summary before the expensive research phases begin.",
        "",
        "Summary",
        f"- Title: {title}",
        f"- Tier: {analysis.get('tier', '?')}",
        f"- Research depth: {depth_label}",
        f"- Estimated sources: {analysis.get('estimated_sources', '?')}",
        f"- Estimated duration: {analysis.get('estimated_duration_minutes', '?')} minutes",
        f"- Citation style: {citation_style.upper()}",
        f"- Output formats: {', '.join(fmt.upper() for fmt in formats)}",
        f"- Plagiarism check: {'enabled' if plagiarism_enabled else 'disabled'}",
        f"- Web search: {'enabled' if web_search_enabled else 'disabled'}",
        f"- Local file sources detected: {local_source_count}",
        f"- User-provided URLs: {provided_url_count}",
    ]
    if depth_description:
        lines.append(f"- Scope profile: {_compact_text(depth_description, width=140)}")
    subtopics = analysis.get("subtopics", []) if isinstance(analysis.get("subtopics", []), list) else []
    if subtopics:
        lines.extend(["", "Key focus areas"])
        for idx, topic in enumerate(subtopics[:6], start=1):
            lines.append(f"- {idx}. {_compact_text(str(topic), width=110)}")
    reason = _compact_text(str(analysis.get("reason", "")).strip(), width=160)
    if reason:
        lines.extend(["", "Why this tier", f"- {reason}"])
    if analysis_storage_path:
        lines.extend(["", f"Full analysis saved in {analysis_storage_path}."])
    lines.extend(
        [
            "",
            "Reply `approve` to continue, or describe the changes you want and I will regenerate the research setup before execution.",
        ]
    )
    return "\n".join(lines).strip()


def _deep_research_subplan_summary_prompt(
    *,
    title: str,
    analysis: dict,
    outline: dict,
    target_pages: int,
    section_pages: int,
    formats: list[str],
    web_search_enabled: bool,
    outline_storage_path: str,
    subplan_storage_path: str,
    version: int,
) -> str:
    sections = outline.get("sections", []) if isinstance(outline, dict) else []
    depth_label = str(analysis.get("depth_label", "") or "Standard Report").strip()
    lines = [
        f"The deep research section plan v{version} is ready for approval.",
        "Review this compact section outline before the long-running research and drafting phases begin.",
        "",
        "Plan summary",
        f"- Title: {title}",
        f"- Tier: {analysis.get('tier', '?')}",
        f"- Research depth: {depth_label}",
        f"- Sections: {len(sections)}",
        f"- Web search: {'enabled' if web_search_enabled else 'disabled'}",
        f"- Output formats: {', '.join(fmt.upper() for fmt in formats)}",
    ]
    if sections:
        lines.extend(["", "Section outline"])
        for section in sections[:8]:
            lines.append(
                f"- {section.get('id')}. {_compact_text(str(section.get('title', 'Section')), width=70)}"
                f": {_compact_text(str(section.get('objective', '')), width=130)}"
            )
    if outline_storage_path or subplan_storage_path:
        lines.extend(["", "Full details"])
        if outline_storage_path:
            lines.append(f"- Outline: {outline_storage_path}")
        if subplan_storage_path:
            lines.append(f"- Step-by-step plan: {subplan_storage_path}")
    lines.extend(
        [
            "",
            "Reply `approve` to continue, or describe the changes you want and I will regenerate the section plan before execution.",
        ]
    )
    return "\n".join(lines).strip()


def _deep_research_subplan_auto_approved_message(
    *,
    title: str,
    outline: dict,
    analysis: dict,
    target_pages: int,
    section_pages: int,
    web_search_enabled: bool,
    outline_storage_path: str,
    subplan_storage_path: str,
    version: int,
) -> str:
    sections = outline.get("sections", []) if isinstance(outline, dict) else []
    depth_label = str(analysis.get("depth_label", "") or "Standard Report").strip()
    lines = [
        f"Deep research section plan v{version} auto-approved.",
        "",
        "Execution handoff",
        f"- Title: {title}",
        f"- Research depth: {depth_label}",
        f"- Sections: {len(sections)}",
        f"- Web search: {'enabled' if web_search_enabled else 'disabled'}",
        "- Next phase: evidence collection and report drafting",
    ]
    if outline_storage_path or subplan_storage_path:
        lines.extend(["", "Saved artifacts"])
        if outline_storage_path:
            lines.append(f"- Outline: {outline_storage_path}")
        if subplan_storage_path:
            lines.append(f"- Step-by-step plan: {subplan_storage_path}")
    return "\n".join(lines).strip()


def _deep_research_handoff_only_enabled(state: Mapping[str, Any]) -> bool:
    value = state.get("_phase0_handoff_only", os.getenv("KENDR_PHASE0_HANDOFF_ONLY", ""))
    if isinstance(value, bool):
        return value
    return str(value or "").strip().lower() in {"1", "true", "yes", "on"}


def _build_deep_research_confirmation_prompt(analysis_md: str, *, version: int) -> str:
    return build_plan_approval_prompt(
        analysis_md,
        scope_title=f"deep research confirmation v{version}",
        storage_note="Review the tier, subtopics, and output settings before the expensive research phases begin.",
    )


def _build_deep_research_analysis_request(
    *,
    title: str,
    analysis: dict[str, Any],
    formats: list[str],
    citation_style: str,
    plagiarism_enabled: bool,
    web_search_enabled: bool,
    local_source_count: int,
    provided_url_count: int,
    analysis_storage_path: str,
    version: int,
) -> dict[str, Any]:
    budget = analysis.get("execution_budget", analysis.get("budget", {})) if isinstance(analysis, dict) else {}
    depth_label = str(analysis.get("depth_label", "") or "Standard Report").strip()
    depth_description = str(analysis.get("depth_description", "") or "").strip()

    def _budget_value(key: str, *, suffix: str = "") -> str:
        raw = 0
        if isinstance(budget, dict):
            raw = int(budget.get(key, 0) or 0)
        if raw > 0:
            return f"{raw}{suffix}"
        return "not explicitly capped"

    overview_items = [
        f"Tier {int(analysis.get('tier', 0) or 0)} deep research run.",
        f"Research depth: {depth_label}.",
        f"Estimated sources: {int(analysis.get('estimated_sources', 0) or 0)}.",
        f"Estimated duration: {int(analysis.get('estimated_duration_minutes', 0) or 0)} minutes.",
        f"Citation style: {citation_style.upper()}.",
        f"Output formats: {', '.join(fmt.upper() for fmt in formats) or 'MD'}.",
        f"Plagiarism check: {'enabled' if plagiarism_enabled else 'disabled'}.",
        f"Web search: {'enabled' if web_search_enabled else 'disabled'}.",
        f"Local file sources detected: {local_source_count}.",
        f"User-provided URLs: {provided_url_count}.",
        f"Date range: {str(analysis.get('date_range', 'all_time') or 'all_time')}.",
    ]
    if depth_description:
        overview_items.insert(2, f"Scope profile: {depth_description}.")
    detected_subtopics = [str(topic).strip() for topic in (analysis.get("subtopics", []) or []) if str(topic).strip()]
    budget_items = [
        f"Max tokens: {_budget_value('max_tokens')}.",
        f"Max sources: {_budget_value('max_sources')}.",
        f"Max duration: {_budget_value('max_duration_minutes', suffix=' minutes')}.",
    ]
    rationale_items = [part.strip() for part in str(analysis.get("reason", "") or "").split(";") if part.strip()]
    return build_approval_request(
        scope="deep_research_confirmation",
        title="Deep Research Analysis",
        summary="The deep research confirmation is ready for approval. Review the tier, subtopics, and output settings before the expensive research phases begin.",
        sections=[
            {"title": "Overview", "items": overview_items},
            {"title": "Detected Subtopics", "items": detected_subtopics},
            {"title": "Session Budget", "items": budget_items},
            {"title": "Why This Tier", "items": rationale_items},
        ],
        accept_label="Start Deep Research",
        reject_label="Revise Scope",
        suggest_label="Suggestion",
        help_text="Approve starts the expensive research pipeline. Reject or Suggestion sends changes back into the paused run.",
        artifact_paths=[analysis_storage_path] if analysis_storage_path else [],
        metadata={"version": version, "report_title": title},
    )


def _build_deep_research_subplan_request(
    *,
    title: str,
    analysis: dict[str, Any],
    outline: dict[str, Any],
    target_pages: int,
    section_pages: int,
    formats: list[str],
    web_search_enabled: bool,
    outline_storage_path: str,
    subplan_storage_path: str,
    version: int,
) -> dict[str, Any]:
    sections = outline.get("sections", []) if isinstance(outline.get("sections"), list) else []
    depth_label = str(analysis.get("depth_label", "") or "Standard Report").strip()
    section_titles = []
    for index, section in enumerate(sections, start=1):
        section_title = str((section or {}).get("title", f"Section {index}") or f"Section {index}").strip()
        questions = (section or {}).get("key_questions", []) if isinstance(section, dict) else []
        q_count = len(questions) if isinstance(questions, list) else 0
        section_titles.append(f"{index}. {section_title} | questions={q_count or 'n/a'}")
    execution_items = [
        f"Execution scope: {depth_label}.",
        f"Section count: {len(sections)}.",
        f"Web search: {'enabled' if web_search_enabled else 'disabled'}.",
        f"Output formats: {', '.join(fmt.upper() for fmt in formats) or 'MD'}.",
        f"Research tier: {int(analysis.get('tier', 0) or 0)}.",
    ]
    return build_approval_request(
        scope="long_document_plan",
        title=title or "Deep Research Report",
        summary="The deep research section plan is ready for approval.",
        sections=[
            {"title": "Section Outline", "items": section_titles},
            {"title": "Execution Plan", "items": execution_items},
        ],
        accept_label="Approve Plan",
        reject_label="Revise Plan",
        suggest_label="Suggestion",
        help_text="Approve keeps the same workflow and starts the section research and drafting stages. Reject or Suggestion regenerates the plan before execution.",
        artifact_paths=[path for path in [outline_storage_path, subplan_storage_path] if path],
        metadata={"version": version, "section_count": len(sections)},
    )


def _section_text_blocks(text: str) -> list[str]:
    blocks = []
    for part in re.split(r"\n\s*\n", str(text or "").strip()):
        cleaned = part.strip()
        if cleaned:
            blocks.append(cleaned)
    return blocks


def _build_correlation_package(objective: str, section_packages: list[dict]) -> dict:
    section_titles = [str(item.get("title", f"Section {idx}")).strip() or f"Section {idx}" for idx, item in enumerate(section_packages, start=1)]
    fallback = {
        "section_order": section_titles,
        "cross_cutting_themes": [],
        "contradictions": [],
        "briefing": "Use the outline order, avoid overlap, and explicitly cross-reference related sections where evidence intersects.",
    }
    prompt = f"""
You are the subtopic correlation engine for a deep research report.

Objective:
{objective}

Section research packages:
{json.dumps([
    {
        "title": item.get("title", ""),
        "objective": item.get("objective", ""),
        "key_questions": item.get("key_questions", []),
        "research_preview": _trim_text(item.get("research_text", ""), 1800),
        "source_count": len(item.get("sources", []) if isinstance(item.get("sources", []), list) else []),
    }
    for item in section_packages
], indent=2, ensure_ascii=False)}

Return JSON only:
{{
  "section_order": ["title 1", "title 2"],
  "cross_cutting_themes": ["theme 1", "theme 2"],
  "contradictions": ["contradiction 1"],
  "briefing": "500-1000 word editorial briefing for the writer"
}}
"""
    try:
        data = llm_json(prompt, fallback)
    except Exception:
        data = fallback
    if not isinstance(data, dict):
        data = fallback
    raw_order = data.get("section_order", [])
    if not isinstance(raw_order, list):
        raw_order = section_titles
    order: list[str] = []
    seen: set[str] = set()
    title_lookup = {title.lower(): title for title in section_titles}
    for item in raw_order:
        key = str(item).strip().lower()
        resolved = title_lookup.get(key)
        if not resolved or resolved in seen:
            continue
        seen.add(resolved)
        order.append(resolved)
    for title in section_titles:
        if title not in seen:
            order.append(title)
    themes = [str(item).strip() for item in data.get("cross_cutting_themes", []) if str(item).strip()] if isinstance(data.get("cross_cutting_themes", []), list) else []
    contradictions = [str(item).strip() for item in data.get("contradictions", []) if str(item).strip()] if isinstance(data.get("contradictions", []), list) else []
    briefing = str(data.get("briefing", "")).strip() or fallback["briefing"]
    nodes = [{"id": item.get("title", ""), "label": item.get("title", ""), "source_count": len(item.get("sources", []))} for item in section_packages]
    edges = []
    for left, right in zip(order, order[1:]):
        edges.append({"from": left, "to": right, "type": "depends_on", "weight": 0.6})
    if contradictions and len(order) >= 2:
        edges.append({"from": order[0], "to": order[-1], "type": "contradicts", "weight": 0.4, "note": contradictions[0]})
    return {
        "section_order": order,
        "cross_cutting_themes": themes,
        "contradictions": contradictions,
        "briefing": briefing,
        "knowledge_graph": {"nodes": nodes, "edges": edges},
    }


def _ensure_section_citations(section_text: str, section_sources: list[dict]) -> str:
    source_ids = [str(item.get("id", "")).strip() for item in section_sources if str(item.get("id", "")).strip()]
    if not source_ids:
        return section_text.strip()
    default_source = source_ids[0]
    lines = []
    citation_re = re.compile(r"\[[A-Z]\d+\]")
    for raw in str(section_text or "").splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or stripped.startswith("```") or stripped.startswith("|") or stripped.startswith("- ") or stripped.startswith("* ") or re.match(r"^\d+\.\s", stripped):
            lines.append(line)
            continue
        if citation_re.search(stripped):
            lines.append(line)
            continue
        lines.append(f"{line} [{default_source}]")
    return "\n".join(lines).strip()


def _build_plagiarism_source_texts(
    *,
    evidence_bank_md: str,
    section_packages: list[dict],
    local_entries: list[dict],
    url_entries: list[dict],
) -> list[dict]:
    entries: list[dict[str, str]] = []
    seen: set[str] = set()

    def _add_entry(*, label: str, text: str, url: str = "", kind: str = "external_source") -> None:
        cleaned_text = str(text or "").strip()
        if not cleaned_text:
            return
        normalized = re.sub(r"\s+", " ", cleaned_text).strip().lower()
        if normalized in seen:
            return
        seen.add(normalized)
        entries.append({"label": label, "url": url, "text": cleaned_text, "kind": kind})

    _add_entry(label="Evidence bank", text=evidence_bank_md, kind="evidence_bank")
    for package in section_packages or []:
        title = str(package.get("title", "")).strip() or "Section research"
        _add_entry(
            label=f"{title} research memo",
            text=str(package.get("research_text", "")).strip(),
            kind="research_memo",
        )
    for entry in local_entries or []:
        if str(entry.get("error", "")).strip():
            continue
        path_value = str(entry.get("path", "")).strip()
        label = str(entry.get("file_name", "")).strip() or Path(path_value).name or "Local file"
        source_text = "\n\n".join(
            part
            for part in (
                f"Summary:\n{str(entry.get('summary', '')).strip()}" if str(entry.get("summary", "")).strip() else "",
                f"Excerpt:\n{str(entry.get('excerpt', '')).strip()}" if str(entry.get("excerpt", "")).strip() else "",
            )
            if part
        )
        _add_entry(
            label=label,
            text=source_text,
            url=_file_source_url(path_value) if path_value else "",
            kind="local_file",
        )
    for entry in url_entries or []:
        if str(entry.get("error", "")).strip():
            continue
        url = str(entry.get("url", "")).strip()
        label = str(entry.get("label", "")).strip() or _source_label(url)
        source_text = "\n\n".join(
            part
            for part in (
                f"Summary:\n{str(entry.get('summary', '')).strip()}" if str(entry.get("summary", "")).strip() else "",
                f"Excerpt:\n{str(entry.get('excerpt', '')).strip()}" if str(entry.get("excerpt", "")).strip() else "",
            )
            if part
        )
        _add_entry(label=label or "Provided URL", text=source_text, url=url, kind="provided_url")
    return entries


def _build_plagiarism_report(section_outputs: list[dict], source_texts: list[dict]) -> dict:
    return _dedicated_build_plagiarism_report(section_outputs, source_texts)


def _format_citation(entry: dict, *, style: str, index: int) -> str:
    label = str(entry.get("label", "Untitled source")).strip() or "Untitled source"
    url = str(entry.get("url", "")).strip()
    access_date = dt.datetime.now(dt.UTC).strftime("%Y-%m-%d")
    site = _source_label(url) if url else label
    year_match = re.search(r"(20\d{2}|19\d{2})", label)
    year = year_match.group(1) if year_match else dt.datetime.now(dt.UTC).strftime("%Y")
    if style == "mla":
        return f'"{label}." {site}, {year}, {url}. Accessed {access_date}.'.strip()
    if style == "chicago":
        return f'{site}. "{label}." Accessed {access_date}. {url}.'.strip()
    if style == "ieee":
        return f'[{index}] "{label}," {site}, {year}. [Online]. Available: {url}.'.strip()
    if style == "vancouver":
        return f"{index}. {label}. {site}. {year}. Available from: {url}".strip()
    if style == "harvard":
        return f"{site} ({year}) {label}. Available at: {url} (Accessed: {access_date}).".strip()
    return f"{site}. ({year}). {label}. Retrieved {access_date}, from {url}".strip()


def _bibliography_markdown(entries: list[dict], *, style: str) -> str:
    lines = [_heading_with_anchor(2, f"Bibliography ({style.upper()})", "bibliography"), ""]
    if not entries:
        lines.append("- No sources were available for bibliography generation.")
        return "\n".join(lines).strip() + "\n"
    for index, entry in enumerate(entries, start=1):
        citation = _format_citation(entry, style=style, index=index)
        if style in {"ieee", "vancouver"}:
            lines.append(f"- {citation}")
        else:
            lines.append(f"- [R{index}] {citation}")
    return "\n".join(lines).strip() + "\n"


def _plagiarism_report_markdown(report: dict) -> str:
    lines = [_heading_with_anchor(2, "Appendix B - Plagiarism Report", "appendix-b-plagiarism-report"), ""]
    lines.append(f"- Overall similarity score: {report.get('overall_score', 0)}%")
    lines.append(f"- AI-writing risk score: {report.get('ai_content_score', 0)}%")
    lines.append(f"- Status: {report.get('status', 'PASS')}")
    summary = str(report.get("summary", "")).strip()
    if summary:
        lines.append(f"- Summary: {summary}")
    source_block_count = int(report.get("source_block_count", 0) or 0)
    if source_block_count:
        lines.append(f"- Source passages scanned: {source_block_count}")
    method = report.get("method", {}) if isinstance(report, dict) else {}
    matching = method.get("matching", []) if isinstance(method, dict) else []
    ai_scoring = method.get("ai_scoring", []) if isinstance(method, dict) else []
    if matching:
        lines.append(f"- Matching signals: {', '.join(str(item) for item in matching)}")
    if ai_scoring:
        lines.append(f"- AI signals: {', '.join(str(item) for item in ai_scoring)}")
    sections = report.get("sections", []) if isinstance(report, dict) else []
    if not sections:
        lines.append("- No section-level findings.")
        return "\n".join(lines).strip() + "\n"
    for section in sections:
        lines.extend(
            [
                "",
                f"### {section.get('section_title', 'Section')}",
                f"- Similarity score: {section.get('plagiarism_score', 0)}%",
                f"- AI-writing risk: {section.get('ai_score', 0)}%",
                f"- Passages scanned: {section.get('passages_scanned', 0)}",
                f"- Words scanned: {section.get('words_scanned', 0)}",
            ]
        )
        ai_components = section.get("ai_components", {}) if isinstance(section, dict) else {}
        ranked_components = [
            (str(key).replace("_", " "), float(value))
            for key, value in ai_components.items()
            if float(value or 0) > 0
        ]
        ranked_components.sort(key=lambda item: item[1], reverse=True)
        if ranked_components:
            top_components = ", ".join(f"{label}={round(value, 1)}" for label, value in ranked_components[:4])
            lines.append(f"- Strongest AI signals: {top_components}")
        flagged = []
        if isinstance(section, dict):
            flagged = section.get("flagged_passages", []) or section.get("flagged_paragraphs", [])
        if not flagged:
            lines.append("- No flagged passages.")
            continue
        lines.append("- Matched passages:")
        for item in flagged[:5]:
            match_type = str(item.get("type", "flagged")).replace("_", " ")
            source_label = str(item.get("source_label", "")).strip() or str(item.get("source_url", "")).strip() or "unknown source"
            lines.append(
                f"  - {match_type} | severity={item.get('severity', 'n/a')} | "
                f"similarity={item.get('similarity', 0)} | source={source_label}"
            )
            excerpt = str(item.get("text_excerpt", "")).strip()
            if excerpt:
                lines.append(f"    Excerpt: {excerpt}")
            matched_phrase = str(item.get("matched_phrase", "")).strip()
            if matched_phrase:
                lines.append(f"    Matched phrase: {matched_phrase}")
            recommendation = str(item.get("recommendation", "")).strip()
            if recommendation:
                lines.append(f"    Recommendation: {recommendation}")
    return "\n".join(lines).strip() + "\n"


def _long_document_subplan(outline: dict, *, objective: str, target_pages: int, research_model: str) -> dict:
    draft_selection = model_selection_for_agent("long_document_agent")
    raw_steps: list[dict[str, Any]] = []
    for index, section in enumerate(outline.get("sections", []) if isinstance(outline, dict) else [], start=1):
        if not isinstance(section, dict):
            continue
        section_title = str(section.get("title", f"Section {index}")).strip() or f"Section {index}"
        section_objective = str(section.get("objective", objective)).strip() or objective
        section_questions = section.get("key_questions", [])
        if not isinstance(section_questions, list):
            section_questions = []
        research_step_id = f"section-{index}-research"
        draft_step_id = f"section-{index}-draft"
        raw_steps.append(
            {
                "id": research_step_id,
                "title": f"Research evidence for {section_title}",
                "agent": "deep_research_agent",
                "task": (
                    f"Collect authoritative evidence, citations, and contradictions for {section_title}. "
                    f"Focus on: {section_objective}. Questions: {section_questions[:6]}"
                ),
                "depends_on": [],
                "parallel_group": "",
                "success_criteria": "Research package is strong enough to support a sourced section draft.",
                "rationale": "Front-load evidence before drafting.",
                "llm_model": research_model,
                "model_source": "state.research_model/KENDR_DEEP_RESEARCH_MODEL",
                "substeps": [],
            }
        )
        raw_steps.append(
            {
                "id": draft_step_id,
                "title": f"Draft and enrich {section_title}",
                "agent": "long_document_agent",
                "task": (
                    f"Draft {section_title} using the approved outline, recent continuity notes, and gathered research. "
                    "Include citations and visuals when they add clarity."
                ),
                "depends_on": [research_step_id],
                "parallel_group": "",
                "success_criteria": "Section is coherent, cited, and ready to merge into the full document.",
                "rationale": "Convert research into the actual deliverable section.",
                "llm_model": draft_selection.get("model", ""),
                "model_source": draft_selection.get("source", ""),
                "substeps": [],
            }
        )
    raw_steps.append(
        {
            "id": "final-merge",
            "title": "Merge sections and produce deep research artifacts",
            "agent": "long_document_agent",
            "task": "Merge all approved sections into one compiled deep research report with executive summary, references, appendices, and export artifacts.",
            "depends_on": [step["id"] for step in raw_steps if str(step.get("id", "")).endswith("-draft")] or [step["id"] for step in raw_steps],
            "parallel_group": "",
            "success_criteria": "Compiled deep research report, references, plagiarism appendix, and manifest are written to disk.",
            "rationale": "Finish the document only after the section plan has been executed.",
            "llm_model": draft_selection.get("model", ""),
            "model_source": draft_selection.get("source", ""),
            "substeps": [],
        }
    )
    return normalize_plan_data(
        {
            "needs_clarification": False,
            "clarification_questions": [],
            "summary": (
                f"Deliver a {target_pages}-page deep research report through approved section-by-section research, "
                "drafting, and final merge."
            ),
            "steps": raw_steps,
        },
        objective,
    )


def _run_research_pass(
    client: OpenAI,
    *,
    query: str,
    model: str,
    instructions: str,
    max_tool_calls: int,
    max_output_tokens: int | None,
    poll_interval_seconds: int,
    max_wait_seconds: int,
    heartbeat_interval_seconds: int | None = None,
    heartbeat_label: str = "Research pass in progress",
    heartbeat_callback: Any | None = None,
    cancel_check: Any | None = None,
) -> dict:
    create_kwargs = {
        "model": model,
        "input": query,
        "instructions": instructions,
        "background": True,
        "max_tool_calls": max_tool_calls,
        "reasoning": {"summary": "auto"},
        "tools": [{"type": OPENAI_WEB_SEARCH_TOOL}],
    }
    if max_output_tokens is not None and max_output_tokens > 0:
        create_kwargs["max_output_tokens"] = max_output_tokens

    response = client.responses.create(**create_kwargs)
    response_id = response.id
    status = str(getattr(response, "status", "unknown"))
    elapsed_seconds = 0
    terminal = {"completed", "failed", "cancelled", "incomplete"}
    heartbeat_every = int(heartbeat_interval_seconds or 0)
    last_heartbeat = 0

    while status not in terminal:
        if callable(cancel_check):
            cancel_check()
        if elapsed_seconds >= max_wait_seconds:
            return {
                "response_id": response_id,
                "status": "timeout",
                "elapsed_seconds": elapsed_seconds,
                "output_text": getattr(response, "output_text", "") or "",
                "raw": response.model_dump() if hasattr(response, "model_dump") else {"status": status},
            }
        time.sleep(poll_interval_seconds)
        elapsed_seconds += poll_interval_seconds
        if heartbeat_every > 0 and (elapsed_seconds - last_heartbeat) >= heartbeat_every:
            last_heartbeat = elapsed_seconds
            if heartbeat_label:
                log_task_update(DEEP_RESEARCH_LABEL, f"{heartbeat_label} ({elapsed_seconds}s elapsed).")
            if callable(heartbeat_callback):
                try:
                    heartbeat_callback(elapsed_seconds)
                except Exception:
                    pass
        if callable(cancel_check):
            cancel_check()
        response = client.responses.retrieve(response_id)
        status = str(getattr(response, "status", "unknown"))

    payload = response.model_dump() if hasattr(response, "model_dump") else {"status": status}
    output_text = getattr(response, "output_text", None) or _extract_output_text(payload)
    return {
        "response_id": response_id,
        "status": status,
        "elapsed_seconds": elapsed_seconds,
        "output_text": output_text or "",
        "raw": payload,
    }


def _collect_section_research_package(
    *,
    api_key: str,
    objective: str,
    section: dict,
    section_index: int,
    total_sections: int,
    section_pages: int,
    use_section_search: bool,
    section_search_results_count: int,
    collect_sources_first: bool,
    evidence_excerpt: str,
    evidence_sources: list[dict],
    explicit_source_entries: list[dict],
    local_entries: list[dict],
    url_entries: list[dict],
    continuity_notes: list[str],
    coherence_context_md: str,
    web_search_enabled: bool,
    native_web_search_enabled: bool,
    research_model: str,
    research_instructions: str,
    max_tool_calls: int,
    max_output_tokens_int: int | None,
    poll_interval_seconds: int,
    max_wait_seconds: int,
    heartbeat_seconds: int,
    max_sources: int,
    research_kb_enabled: bool,
    research_kb_id: str,
    research_kb_top_k: int,
    cancel_check: Any | None = None,
) -> dict[str, Any]:
    section_title = str(section.get("title", f"Section {section_index}")).strip() or f"Section {section_index}"
    section_objective = str(section.get("objective", objective)).strip() or objective
    section_questions = section.get("key_questions", [])
    if not isinstance(section_questions, list):
        section_questions = []
    target_section_pages = _safe_int(section.get("target_pages"), section_pages, 1, 30)

    search_query = ""
    section_search_results: dict[str, Any] = {}
    section_search_sources: list[dict] = []
    section_search_text = ""
    if use_section_search:
        if callable(cancel_check):
            cancel_check()
        search_query = f"{section_title} {section_objective}"
        section_search_results = _collect_google_search_evidence(search_query, num=section_search_results_count)
        section_search_sources = _sources_from_search_results(section_search_results.get("results", []))
        section_search_text = _search_results_markdown(section_search_results.get("results", []))

    section_kb_query = "\n".join(
        [
            section_title,
            section_objective,
            *[str(item).strip() for item in section_questions[:6] if str(item).strip()],
        ]
    ).strip()
    section_kb_grounding: dict[str, Any] = {"requested": bool(research_kb_enabled)}
    section_kb_warning = ""
    if research_kb_enabled:
        try:
            section_kb_grounding = build_research_grounding(
                section_kb_query or section_objective or objective,
                kb_ref=research_kb_id,
                top_k=research_kb_top_k,
                use_active_if_empty=True,
                require_indexed=True,
            )
            section_kb_grounding["requested"] = True
            if int(section_kb_grounding.get("hit_count", 0) or 0) <= 0:
                section_kb_warning = (
                    f"Knowledge base '{section_kb_grounding.get('kb_name', 'unknown')}' returned no relevant results for section {section_index}."
                )
        except Exception as exc:
            section_kb_warning = str(exc)
            section_kb_grounding = {"requested": True, "hit_count": 0, "citations": [], "prompt_context": ""}

    local_section_notes = (
        _build_local_only_research_notes(
            objective=objective,
            focus=f"section {section_index}: {section_title} - {section_objective}",
            local_entries=local_entries,
            url_entries=url_entries,
            continuity_notes=continuity_notes,
        )
        if explicit_source_entries
        else ""
    )

    if collect_sources_first and evidence_excerpt:
        research_pass = {
            "response_id": "evidence_bank" if web_search_enabled else f"local_only_section_{section_index}",
            "status": "evidence_bank" if web_search_enabled else "local_only",
            "elapsed_seconds": 0,
            "output_text": "\n\n".join(
                [
                    item
                    for item in (
                        section_search_text,
                        _kb_grounding_summary(section_kb_grounding),
                        local_section_notes or evidence_excerpt,
                    )
                    if str(item).strip()
                ]
            ).strip(),
            "raw": {"evidence_bank_path": ""},
        }
        research_pass["research_kb"] = section_kb_grounding
        research_pass["research_kb_warning"] = section_kb_warning
        research_output = str(research_pass.get("output_text", "")).strip()
        section_sources = _merge_sources(list(evidence_sources), _sources_from_kb_grounding(section_kb_grounding))
        section_sources = _merge_sources(section_sources, section_search_sources)
    else:
        if web_search_enabled:
            if native_web_search_enabled:
                if callable(cancel_check):
                    cancel_check()
                query_lines = [
                    f"Global objective: {objective}",
                    f"Section {section_index} title: {section_title}",
                    f"Section objective: {section_objective}",
                    "Key questions:",
                    *[f"- {item}" for item in section_questions[:12]],
                    "Continuity notes from previous sections:",
                    *[f"- {item}" for item in continuity_notes[-10:]],
                    "Markdown coherence anchors:",
                    _trim_text(coherence_context_md, 4000),
                ]
                research_pass = _run_research_pass(
                    OpenAI(api_key=api_key),
                    query="\n".join(query_lines),
                    model=research_model,
                    instructions=research_instructions,
                    max_tool_calls=max_tool_calls,
                    max_output_tokens=max_output_tokens_int,
                    poll_interval_seconds=poll_interval_seconds,
                    max_wait_seconds=max_wait_seconds,
                    heartbeat_interval_seconds=heartbeat_seconds,
                    heartbeat_label=f"Section {section_index} research in progress",
                    cancel_check=cancel_check,
                )
                research_output = str(research_pass.get("output_text", "")).strip()
            else:
                if not search_query:
                    search_query = f"{section_title} {section_objective}"
                if not section_search_text:
                    if callable(cancel_check):
                        cancel_check()
                    section_search_results = _collect_google_search_evidence(search_query, num=section_search_results_count)
                    section_search_sources = _sources_from_search_results(section_search_results.get("results", []))
                    section_search_text = _search_results_markdown(section_search_results.get("results", []))
                research_pass = {
                    "response_id": f"kendr_web_search_section_{section_index}",
                    "status": "fallback_web_search",
                    "elapsed_seconds": 0,
                    "output_text": local_section_notes,
                    "raw": {
                        "reason": "native_web_search_unavailable",
                        "search_provider": str((section_search_results or {}).get("provider", "")).strip(),
                        "search_providers_tried": list((section_search_results or {}).get("providers_tried", []) or []),
                    },
                }
                research_output = local_section_notes
        else:
            research_pass = {
                "response_id": f"local_only_section_{section_index}",
                "status": "local_only",
                "elapsed_seconds": 0,
                "output_text": "\n\n".join(
                    item
                    for item in (_kb_grounding_summary(section_kb_grounding), local_section_notes)
                    if str(item).strip()
                ).strip(),
                "raw": {"reason": "web_search_disabled"},
            }
            research_output = local_section_notes
        research_pass["research_kb"] = section_kb_grounding
        research_pass["research_kb_warning"] = section_kb_warning
        if section_kb_grounding.get("prompt_context"):
            research_output = (
                f"{_kb_grounding_summary(section_kb_grounding)}\n\n{research_output}".strip()
                if research_output
                else _kb_grounding_summary(section_kb_grounding)
            )
        if section_search_text:
            research_output = f"{section_search_text}\n\n{research_output}".strip()
        if not research_output:
            research_output = (
                section_search_text
                or "Research output was empty. Use only explicitly supported claims and call out uncertainty."
            )

        section_sources = _merge_sources(explicit_source_entries, _extract_source_entries(research_pass))
        section_sources = _merge_sources(section_sources, _sources_from_kb_grounding(section_kb_grounding))
        section_sources = _merge_sources(section_sources, section_search_sources)

    if max_sources > 0:
        section_sources = section_sources[:max_sources]
    source_ledger_md = _source_ledger_markdown(section_sources)
    return {
        "index": section_index,
        "title": section_title,
        "objective": section_objective,
        "key_questions": section_questions,
        "target_pages": target_section_pages,
        "research_pass": research_pass,
        "research_text": research_output,
        "sources": section_sources,
        "source_ledger_md": source_ledger_md,
        "search_query": search_query,
        "section_search_results": section_search_results,
        "research_kb": section_kb_grounding,
        "research_kb_warning": section_kb_warning,
    }


def _record_section_research_package(
    state: dict,
    *,
    package: dict[str, Any],
    total_sections: int,
    artifact_dir: str,
    research_log_lines: list[str],
    section_research_started_at: str,
) -> dict[str, Any]:
    index = int(package.get("index", 0) or 0)
    section_title = str(package.get("title", f"Section {index}")).strip() or f"Section {index}"
    from_cache = bool(package.get("from_cache", False))
    research_pass = package.get("research_pass", {}) if isinstance(package.get("research_pass"), dict) else {}
    section_sources = package.get("sources", []) if isinstance(package.get("sources"), list) else []
    search_query = str(package.get("search_query", "")).strip()
    section_search_results = (
        package.get("section_search_results", {}) if isinstance(package.get("section_search_results"), dict) else {}
    )

    if search_query:
        _trace_research_event(
            state,
            title=f"Google search results for section {index}",
            detail=f"{section_title} returned {len((section_search_results or {}).get('results', []))} result(s).",
            command=search_query,
            status="completed" if not str((section_search_results or {}).get("error", "")).strip() else "failed",
            metadata={
                "phase": "section_search",
                "section_index": index,
                "section_title": section_title,
                "search_query": search_query,
                "search_provider": str((section_search_results or {}).get("provider", "")).strip(),
                "search_providers_tried": list((section_search_results or {}).get("providers_tried", []) or []),
                "candidate_urls": _trace_url_list(
                    [str(item.get("url", "")).strip() for item in (section_search_results or {}).get("results", [])]
                ),
                "viewed_urls": _trace_url_list(
                    [str(item.get("url", "")).strip() for item in (section_search_results or {}).get("viewed_pages", [])]
                ),
                "urls": _trace_url_list(
                    [str(item.get("url", "")).strip() for item in (section_search_results or {}).get("results", [])]
                ),
            },
            subtask=f"Search for {section_title}",
        )
        if str((section_search_results or {}).get("error", "")).strip():
            log_task_update(DEEP_RESEARCH_LABEL, f"Section {index} web search error: {section_search_results.get('error')}")

    section_research_status = str(research_pass.get("status", "")).strip() or "completed"
    if section_research_status not in {"completed", "", "local_only", "evidence_bank", "fallback_web_search"}:
        log_task_update(DEEP_RESEARCH_LABEL, f"Section {index} research status: {research_pass.get('status')}")

    _trace_research_event(
        state,
        title=f"Researching section {index}/{total_sections}",
        detail=(
            (
                f"{section_title} reused from resume artifacts with {len(section_sources)} sources."
                if from_cache
                else f"{section_title} completed with {len(section_sources)} sources."
            )
            if section_research_status in {"completed", "local_only", "evidence_bank", "fallback_web_search"}
            else f"{section_title} finished with status '{section_research_status}'. Using partial research output."
        ),
        command=str(package.get("objective", "")).strip(),
        status="completed" if section_research_status in {"completed", "local_only", "evidence_bank", "fallback_web_search"} else "failed",
        started_at=section_research_started_at,
        completed_at=_trace_now(),
        metadata={
            "phase": "section_research",
            "section_index": index,
            "section_title": section_title,
            "response_status": section_research_status,
            "sources": len(section_sources),
            "elapsed_seconds": int(research_pass.get("elapsed_seconds", 0) or 0),
            "search_query": search_query,
            "urls": _trace_url_list([str(item.get("url", "")).strip() for item in section_sources]),
            "resumed_from_cache": from_cache,
        },
        subtask=f"Gather evidence for {section_title}",
    )

    write_text_file(
        _artifact_file(artifact_dir, f"section_{index:02d}/research.json"),
        json.dumps(research_pass, indent=2, ensure_ascii=False),
    )
    write_text_file(
        _artifact_file(artifact_dir, f"section_{index:02d}/sources.json"),
        json.dumps(section_sources, indent=2, ensure_ascii=False),
    )
    write_text_file(
        _artifact_file(artifact_dir, f"section_{index:02d}/sources.md"),
        _references_markdown(section_sources),
    )
    research_log_lines.append(
        (
            f"Reused section {index} research: {section_title} ({len(section_sources)} sources)"
            if from_cache
            else f"Researched section {index}: {section_title} ({len(section_sources)} sources)"
        )
    )

    package = dict(package)
    package.pop("section_search_results", None)
    return package


def _format_section_prompt(
    *,
    global_objective: str,
    section: dict,
    section_index: int,
    section_count: int,
    words_target: int,
    continuity_notes: list[str],
    coherence_context_md: str,
    correlation_briefing: str,
    source_ledger_md: str,
    research_text: str,
    evidence_bank_md: str = "",
    citation_style: str = "apa",
    date_range: str = "all_time",
) -> str:
    continuity = "\n".join(f"- {item}" for item in continuity_notes[-8:]) or "- No prior continuity notes."
    key_questions = "\n".join(f"- {item}" for item in section.get("key_questions", [])) or "- None provided."
    return f"""
You are writing Section {section_index} of {section_count} in a deep research report.

Global objective:
{global_objective}

Section title:
{section.get("title", f"Section {section_index}")}

Section objective:
{section.get("objective", global_objective)}

Key questions:
{key_questions}

Continuity constraints from prior sections:
{continuity}

Markdown coherence anchors:
{coherence_context_md}

Extracted source ledger for this section:
{source_ledger_md}

Pre-collected evidence bank (cross-section reference):
{evidence_bank_md or "No evidence bank provided."}

Write a section draft of about {words_target} words.
Requirements:
- Keep conceptual continuity with prior sections.
- Keep factual claims tied to evidence gathered in the research notes below.
- Use the editorial correlation briefing to avoid overlap and to cross-reference related sections.
- Do not repeat the section title as a heading; the compiler adds the section heading for you.
- Include a short "Section Takeaways" list at the end.
- Use source tags like "[S1]" inline for factual claims that come from the source ledger.
- Do not cite source ids that are not present in the source ledger.
- Do not add a references section; references will be consolidated and renumbered at the end of the document.
- Maintain a formal research tone and write for a deep-research report.
- Cite every factual paragraph.
- Use cross-reference phrasing like "As discussed in Section X..." where genuinely helpful.
- Citation style target for bibliography: {citation_style.upper()}.
- Date range emphasis: {date_range}.

Correlation briefing:
{correlation_briefing}

Research notes:
{research_text}
"""


def _section_continuity_note(section_title: str, section_text: str) -> str:
    prompt = f"""
Summarize this section into continuity notes for later chapters.
Return 6-10 concise bullets, each one line.

Section title: {section_title}
Section text:
{section_text}
"""
    return llm_text(prompt)


def _markdown_to_plain_text(markdown_text: str) -> str:
    input_chars = len(markdown_text or "")
    lines: list[str] = []
    in_code = False
    for raw in (markdown_text or "").splitlines():
        line = raw.rstrip()
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            lines.append(line)
            continue
        line = re.sub(r"^#{1,6}\s+", "", line)
        line = re.sub(r"^[-*]\s+", "- ", line)
        line = re.sub(r"^\d+\.\s+", "", line)
        lines.append(line)
    cleaned = "\n".join(lines).strip()
    # Best-effort ASCII clean-up for PDF rendering (latin-1 PDF stream)
    cleaned_ascii = cleaned.encode("ascii", errors="ignore").decode("ascii")
    stripped_chars = len(cleaned) - len(cleaned_ascii)
    if stripped_chars > 0:
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"_markdown_to_plain_text: stripped {stripped_chars} non-ASCII characters for PDF rendering "
            f"(input {input_chars} chars → {len(cleaned_ascii)} chars).",
        )
    else:
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"_markdown_to_plain_text: converted {input_chars} chars → {len(cleaned_ascii)} plain-text chars (no non-ASCII loss).",
        )
    return cleaned_ascii


def _markdown_to_html(markdown_text: str, *, base_path: str | Path | None = None) -> str:
    input_chars = len(markdown_text or "")
    log_task_update(DEEP_RESEARCH_LABEL, f"_markdown_to_html: converting {input_chars} chars of markdown to HTML.")
    from tasks.md_to_pdf import build_html_from_markdown

    heading_count = len(re.findall(r"^#{1,6}\s+", markdown_text or "", flags=re.MULTILINE))
    list_item_count = len(re.findall(r"^(?:[-*]|\d+\.)\s+", markdown_text or "", flags=re.MULTILINE))
    table_count = len(re.findall(r"^\|.*\|\s*$", markdown_text or "", flags=re.MULTILINE))
    html_result = build_html_from_markdown(markdown_text or "", for_pdf=False, base_path=base_path)
    log_task_update(
        DEEP_RESEARCH_LABEL,
        f"_markdown_to_html: done — {heading_count} headings, {list_item_count} list items, "
        f"{table_count} tables → {len(html_result)} chars of HTML.",
    )
    return html_result


def _markdown_to_docx(markdown_text: str, output_path: str) -> None:
    input_chars = len(markdown_text or "")
    log_task_update(DEEP_RESEARCH_LABEL, f"_markdown_to_docx: converting {input_chars} chars → {output_path}")
    if not _ensure_python_package("python-docx", "docx"):
        raise RuntimeError("python-docx not available")
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.86)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Georgia"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(9)
    normal_style.paragraph_format.line_spacing = 1.5

    title_style = doc.styles["Title"]
    title_style.font.name = "Georgia"
    title_style.font.size = Pt(25)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(15, 23, 42)

    for style_name, font_size, color in (
        ("Heading 1", 17, RGBColor(15, 23, 42)),
        ("Heading 2", 15, RGBColor(15, 118, 110)),
        ("Heading 3", 12.5, RGBColor(20, 78, 74)),
    ):
        style = doc.styles[style_name]
        style.font.name = "Georgia"
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.font.color.rgb = color

    bookmark_id = 1
    heading_count = 0
    bullet_count = 0
    numbered_count = 0
    para_count = 0

    def _strip_inline_markdown(value: str) -> str:
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", str(value or ""))
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        return text

    def _add_bookmark(paragraph, name: str) -> None:
        nonlocal bookmark_id
        anchor = _slugify_heading(name, fallback=f"bookmark-{bookmark_id}")
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), anchor)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        paragraph._p.insert(0, start)
        paragraph._p.append(end)
        bookmark_id += 1

    def _append_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), _slugify_heading(anchor, fallback="section"))
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0F766E")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_properties.append(color)
        run_properties.append(underline)
        run.append(run_properties)
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _append_inline_runs(paragraph, text: str) -> None:
        token_pattern = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|\*[^*]+\*)")
        cursor = 0
        for match in token_pattern.finditer(text):
            if match.start() > cursor:
                paragraph.add_run(text[cursor:match.start()])
            token = match.group(0)
            if token.startswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
                if link_match:
                    label = link_match.group(1)
                    target = link_match.group(2)
                    if target.startswith("#"):
                        _append_internal_hyperlink(paragraph, label, target[1:])
                    else:
                        run = paragraph.add_run(label)
                        run.font.color.rgb = RGBColor(15, 118, 110)
                        run.underline = True
            cursor = match.end()
        if cursor < len(text):
            paragraph.add_run(text[cursor:])

    lines = (markdown_text or "").splitlines()
    in_code = False
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if line.strip().startswith("```"):
            in_code = not in_code
            index += 1
            continue
        if in_code:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            index += 1
            continue
        if not line.strip():
            doc.add_paragraph("")
            index += 1
            continue
        if "|" in line and index + 1 < len(lines) and re.match(r"^\s*\|?\s*[:\-\s|]+\|?\s*$", lines[index + 1]):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                table_lines.append(lines[index])
                index += 1
            rows = []
            for row_line in table_lines:
                parts = [part.strip() for part in row_line.strip().strip("|").split("|")]
                rows.append(parts)
            if rows:
                header = rows[0]
                body_rows = rows[2:] if len(rows) > 2 else rows[1:]
                table = doc.add_table(rows=1 + len(body_rows), cols=len(header))
                table.style = "Table Grid"
                for col_index, value in enumerate(header):
                    cell = table.rows[0].cells[col_index]
                    cell.text = _strip_inline_markdown(value)
                    if cell.paragraphs:
                        cell.paragraphs[0].runs[0].bold = True
                for row_index, row_values in enumerate(body_rows, start=1):
                    padded = row_values + [""] * max(0, len(header) - len(row_values))
                    for col_index, value in enumerate(padded[:len(header)]):
                        table.rows[row_index].cells[col_index].text = _strip_inline_markdown(value)
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            raw_heading = heading_match.group(2).strip()
            anchor_match = re.match(r"^(.*?)(?:\s+\{#([^}]+)\})?$", raw_heading)
            heading_text = _strip_heading_attributes(anchor_match.group(1).strip() if anchor_match else raw_heading)
            anchor_name = anchor_match.group(2).strip() if anchor_match and anchor_match.group(2) else _slugify_heading(heading_text, fallback=f"heading-{level}")
            if level == 1 and heading_count == 0:
                paragraph = doc.add_paragraph(heading_text, style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph = doc.add_heading(heading_text, level=level)
            _add_bookmark(paragraph, anchor_name)
            heading_count += 1
            index += 1
            continue
        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        if bullet_match:
            paragraph = doc.add_paragraph(style="List Bullet")
            bullet_text = bullet_match.group(1).strip()
            link_match = re.fullmatch(r"\[([^\]]+)\]\((#[^)]+)\)", bullet_text)
            if link_match:
                _append_internal_hyperlink(paragraph, link_match.group(1), link_match.group(2)[1:])
            else:
                _append_inline_runs(paragraph, bullet_text)
            bullet_count += 1
            index += 1
            continue
        numbered_match = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered_match:
            paragraph = doc.add_paragraph(style="List Number")
            _append_inline_runs(paragraph, numbered_match.group(1).strip())
            numbered_count += 1
            index += 1
            continue
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line.strip())
        if img_match:
            img_alt = img_match.group(1).strip()
            img_path = img_match.group(2).strip()
            if not img_path.startswith("http"):
                try:
                    abs_img = resolve_output_path(img_path) if not os.path.isabs(img_path) else img_path
                    if os.path.isfile(abs_img):
                        doc.add_picture(abs_img, width=Inches(5.5))
                        caption_para = doc.add_paragraph(img_alt or "Figure")
                        caption_para.style = "Caption" if "Caption" in [s.name for s in doc.styles] else "Normal"
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        log_task_update(DEEP_RESEARCH_LABEL, f"  Embedded image in DOCX: {abs_img}")
                    else:
                        doc.add_paragraph(f"[Image not found: {img_path}]")
                        log_task_update(DEEP_RESEARCH_LABEL, f"  Image not found for DOCX embed: {abs_img}")
                except Exception as img_exc:
                    doc.add_paragraph(f"[Image: {img_alt or img_path}]")
                    log_task_update(DEEP_RESEARCH_LABEL, f"  DOCX image embed failed ({img_path}): {img_exc}")
            else:
                doc.add_paragraph(f"[Image: {img_alt or img_path}]")
            index += 1
            continue
        if re.match(r"^\s*>\s*", line):
            paragraph = doc.add_paragraph(style="Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else "Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _append_inline_runs(paragraph, re.sub(r"^\s*>\s*", "", line))
            para_count += 1
            index += 1
            continue
        if re.match(r"^\s*-{3,}\s*$", line):
            doc.add_paragraph("")
            index += 1
            continue

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _append_inline_runs(paragraph, line)
        para_count += 1
        index += 1

    doc.save(output_path)
    log_task_update(
        DEEP_RESEARCH_LABEL,
        f"_markdown_to_docx: saved — {heading_count} headings, {bullet_count} bullets, "
        f"{numbered_count} numbered items, {para_count} paragraphs → {output_path}",
    )


def _ensure_python_package(package: str, import_name: str | None = None) -> bool:
    module_name = import_name or package
    try:
        __import__(module_name)
        return True
    except Exception:
        pass

    log_task_update(DEEP_RESEARCH_LABEL, f"Attempting to install missing dependency: {package}.")
    try:
        subprocess.run(
            [
                sys.executable,
                "-m",
                "pip",
                "install",
                "--user",
                "--break-system-packages",
                package,
            ],
            check=True,
            capture_output=True,
            text=True,
        )
    except Exception as exc:
        log_task_update(DEEP_RESEARCH_LABEL, f"Dependency install failed for {package}: {exc}")
        return False

    try:
        __import__(module_name)
        return True
    except Exception as exc:
        log_task_update(DEEP_RESEARCH_LABEL, f"Dependency import failed for {package} after install: {exc}")
        return False


def _wrap_pdf_lines(text: str, width: int = 92) -> list[str]:
    lines = []
    for paragraph in text.splitlines():
        if not paragraph.strip():
            lines.append("")
            continue
        lines.extend(textwrap.wrap(paragraph, width=width) or [""])
    return lines


def _escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _render_pdf_bytes(text: str) -> bytes:
    log_task_update(DEEP_RESEARCH_LABEL, f"_render_pdf_bytes: building PDF from {len(text)} chars of plain text.")
    lines = _wrap_pdf_lines(text)
    page_height = 792
    margin_top = 742
    leading = 14
    lines_per_page = 48
    pages = [lines[i : i + lines_per_page] for i in range(0, max(len(lines), 1), lines_per_page)] or [[]]
    log_task_update(DEEP_RESEARCH_LABEL, f"_render_pdf_bytes: {len(lines)} wrapped lines → {len(pages)} pages.")

    objects: list[bytes] = []

    def add_object(payload: str | bytes) -> int:
        objects.append(payload if isinstance(payload, bytes) else payload.encode("latin-1", errors="replace"))
        return len(objects)

    font_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    page_ids = []
    content_ids = []

    for page_lines in pages:
        stream_lines = ["BT", "/F1 11 Tf", f"50 {margin_top} Td"]
        first_line = True
        for line in page_lines:
            escaped_line = _escape_pdf_text(line)
            if first_line:
                stream_lines.append(f"({escaped_line}) Tj")
                first_line = False
            else:
                stream_lines.append(f"0 -{leading} Td")
                stream_lines.append(f"({escaped_line}) Tj")
        if first_line:
            stream_lines.append("( ) Tj")
        stream_lines.append("ET")
        stream = "\n".join(stream_lines).encode("latin-1", errors="replace")
        content_id = add_object(
            b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream"
        )
        content_ids.append(content_id)
        page_ids.append(add_object(""))

    pages_id = add_object("")
    for idx, page_id in enumerate(page_ids):
        page_object = (
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 612 {page_height}] "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_ids[idx]} 0 R >>"
        )
        objects[page_id - 1] = page_object.encode("latin-1")

    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects[pages_id - 1] = f"<< /Type /Pages /Count {len(page_ids)} /Kids [{kids}] >>".encode("latin-1")
    catalog_id = add_object(f"<< /Type /Catalog /Pages {pages_id} 0 R >>")

    buffer = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, payload in enumerate(objects, start=1):
        offsets.append(len(buffer))
        buffer.extend(f"{index} 0 obj\n".encode("ascii"))
        buffer.extend(payload)
        buffer.extend(b"\nendobj\n")

    xref_offset = len(buffer)
    buffer.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    buffer.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        buffer.extend(f"{offset:010d} 00000 n \n".encode("ascii"))

    buffer.extend(b"trailer\n")
    buffer.extend(f"<< /Size {len(objects) + 1} /Root {catalog_id} 0 R >>\n".encode("ascii"))
    buffer.extend(b"startxref\n")
    buffer.extend(f"{xref_offset}\n".encode("ascii"))
    buffer.extend(b"%%EOF\n")
    result = bytes(buffer)
    log_task_update(DEEP_RESEARCH_LABEL, f"_render_pdf_bytes: PDF built — {len(pages)} pages, {len(result)} bytes.")
    return result


def _export_long_document_formats(
    compiled_markdown: str,
    compiled_filename: str,
    *,
    requested_formats: list[str] | None = None,
) -> dict[str, str]:
    """Returns dict with keys 'html', 'docx', 'pdf' (paths) and 'errors' (dict of format->error msg)."""
    compiled_path = Path(resolve_output_path(compiled_filename))
    base_path = compiled_path.with_suffix("")
    html_path = base_path.with_suffix(".html")
    docx_path = base_path.with_suffix(".docx")
    pdf_path = base_path.with_suffix(".pdf")
    formats = set(_normalize_research_formats(requested_formats or DEFAULT_RESEARCH_FORMATS))
    export_errors: dict[str, str] = {}

    html_text = _markdown_to_html(compiled_markdown, base_path=compiled_path.parent)
    log_task_update(DEEP_RESEARCH_LABEL, f"Document export starting. Requested formats: {', '.join(formats)}.")
    if "html" in formats or "pdf" in formats:
        log_task_update(DEEP_RESEARCH_LABEL, f"Rendering HTML ({len(html_text)} chars)...")
        try:
            html_path.write_text(html_text, encoding="utf-8")
            log_task_update(DEEP_RESEARCH_LABEL, f"HTML written → {html_path}")
        except Exception as exc:
            err_msg = f"HTML export failed: {exc}"
            log_task_update(DEEP_RESEARCH_LABEL, err_msg)
            export_errors["html"] = err_msg
            html_path = Path("")
    else:
        log_task_update(DEEP_RESEARCH_LABEL, "HTML export skipped (not in requested formats).")
        html_path = Path("")

    if "docx" in formats:
        log_task_update(DEEP_RESEARCH_LABEL, f"Rendering DOCX → {docx_path}...")
        try:
            _markdown_to_docx(compiled_markdown, str(docx_path))
            log_task_update(DEEP_RESEARCH_LABEL, f"DOCX written → {docx_path}")
        except Exception as exc:
            err_msg = f"DOCX export failed: {exc}"
            log_task_update(DEEP_RESEARCH_LABEL, err_msg)
            export_errors["docx"] = err_msg
            docx_path = Path("")
    else:
        log_task_update(DEEP_RESEARCH_LABEL, "DOCX export skipped (not in requested formats).")
        docx_path = Path("")

    if "pdf" in formats:
        log_task_update(DEEP_RESEARCH_LABEL, f"Rendering PDF → {pdf_path}...")
        try:
            from tasks.md_to_pdf import md_to_pdf
            compiled_path.write_text(compiled_markdown, encoding="utf-8")
            md_to_pdf(str(compiled_path), str(pdf_path))
            log_task_update(DEEP_RESEARCH_LABEL, f"PDF written → {pdf_path}")
        except Exception as exc:
            err_msg = f"PDF export failed: {exc}"
            log_task_update(DEEP_RESEARCH_LABEL, err_msg)
            export_errors["pdf"] = err_msg
            pdf_path = Path("")
    else:
        log_task_update(DEEP_RESEARCH_LABEL, "PDF export skipped (not in requested formats).")
        pdf_path = Path("")

    return {
        "html": str(html_path) if str(html_path) else "",
        "docx": str(docx_path) if str(docx_path) else "",
        "pdf": str(pdf_path) if str(pdf_path) else "",
        "errors": export_errors,
    }


def _build_compiled_markdown(
    title: str,
    objective: str,
    section_outputs: list[dict],
    executive_summary: str,
    consolidated_references: list[dict],
    *,
    citation_style: str,
    methodology_text: str,
    plagiarism_report: dict,
    source_entries: list[dict],
    research_log_lines: list[str],
    generated_at: str,
    model_name: str,
    deep_research_tier: int,
    image_entries: list[dict] | None = None,
) -> str:
    def _section_anchor(item: dict) -> str:
        index = int(item["index"])
        fallback = f"section-{index}"
        return f"section-{index}-{_slugify_heading(item['title'], fallback=fallback)}"

    toc_entries = [
        ("executive-summary", "Executive Summary"),
        *[
            (_section_anchor(item), f"{item['index']}. {item['title']}")
            for item in section_outputs
        ],
        ("bibliography", f"Bibliography ({citation_style.upper()})"),
        ("appendix-a-source-list", "Appendix A - Source List"),
        ("appendix-b-plagiarism-report", "Appendix B - Plagiarism Report"),
        ("appendix-c-research-log", "Appendix C - Research Log"),
    ]
    if image_entries:
        toc_entries.append(("appendix-d-image-gallery", "Appendix D - Image Gallery"))
    lines = [
        _heading_with_anchor(1, title, "report-title"),
        "",
        _heading_with_anchor(2, "Table of Contents", "table-of-contents"),
    ]
    for anchor, label in toc_entries:
        lines.append(f"- [{label}](#{anchor})")
    lines.extend(
        [
            "",
            _heading_with_anchor(2, "Executive Summary", "executive-summary"),
        ]
    )
    if objective.strip():
        lines.extend(
            [
                "",
                f"> Objective: {objective.strip()}",
            ]
        )
    lines.extend(
        [
            "",
            f"_Generated on {generated_at}_",
            "",
        ]
    )
    lines.extend([executive_summary.strip(), ""])
    for item in section_outputs:
        section_anchor = _section_anchor(item)
        lines.append(_heading_with_anchor(2, f"{item['index']}. {item['title']}", section_anchor))
        lines.append("")
        lines.append(_strip_leading_section_heading(item["section_text"], section_title=item["title"], section_index=int(item["index"])))
        lines.append("")
    bibliography_md = _bibliography_markdown(consolidated_references, style=citation_style).strip()
    lines.extend([bibliography_md, ""])
    lines.extend([_heading_with_anchor(2, "Appendix A - Source List", "appendix-a-source-list"), ""])
    if source_entries:
        for item in source_entries:
            lines.append(f"- [{item.get('id', '')}] {item.get('label', '')} - {item.get('url', '')}")
    else:
        lines.append("- No source list was available.")
    lines.append("")
    lines.append(_plagiarism_report_markdown(plagiarism_report).strip())
    lines.append("")
    lines.extend([_heading_with_anchor(2, "Appendix C - Research Log", "appendix-c-research-log"), ""])
    if research_log_lines:
        for line in research_log_lines:
            lines.append(f"- {line}")
    else:
        lines.append("- No research log lines were captured.")
    lines.append("")
    if methodology_text.strip():
        lines.extend(
            [
                _heading_with_anchor(3, "Run Configuration", "appendix-c-run-configuration"),
                "",
                methodology_text.strip(),
                "",
            ]
        )
    lines.append(_image_appendix_markdown(image_entries or []))
    lines.extend(
        [
            "",
            "---",
            "",
            f"_Generated by Kendr on {generated_at}. Visit at [Kendr.org](https://kendr.org)._",
            "",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def long_document_agent(state):
    active_task, task_content, _ = begin_agent_session(state, "long_document_agent")
    state["long_document_calls"] = state.get("long_document_calls", 0) + 1
    call_number = state["long_document_calls"]
    _raise_if_cancelled(state, phase="startup")

    objective = str(state.get("current_objective") or task_content or state.get("user_query", "")).strip()
    if not objective:
        raise ValueError("long_document_agent requires a non-empty objective.")

    # Ensure long-running document generation has persistent session memory files
    # (including soul.md) available for coherence anchoring.
    soul_file = str(state.get("memory_soul_file", "")).strip()
    if not soul_file or not Path(soul_file).exists():
        state = bootstrap_file_memory(state)

    target_pages = _safe_int(state.get("long_document_pages") or state.get("report_target_pages"), 50, 5, 500)
    research_depth_config = _research_depth_config(state.get("research_depth_mode"), target_pages)
    research_depth_mode = str(research_depth_config["mode"])
    research_depth_label = str(research_depth_config["label"])
    section_pages = _safe_int(state.get("long_document_section_pages"), 5, 2, 20)
    section_count_default = max(3, math.ceil(target_pages / max(1, section_pages)))
    section_count = _safe_int(state.get("long_document_sections"), section_count_default, 1, 40)
    initial_outline = state.get("long_document_outline", {})
    if not isinstance(initial_outline, dict):
        initial_outline = {}
    title = (
        str(state.get("long_document_title", "") or initial_outline.get("title", "")).strip()
        or _generate_report_title(objective)
    )
    state["research_depth_mode"] = research_depth_mode

    draft_selection = model_selection_for_agent("long_document_agent")
    explicit_research_provider = str(state.get("research_provider") or "").strip().lower()
    selected_provider = explicit_research_provider or str(state.get("provider") or draft_selection.get("provider") or "").strip().lower()
    selected_model = str(state.get("model") or draft_selection.get("model") or "").strip()

    _raw_research_model = str(state.get("research_model") or selected_model).strip()
    _PROVIDER_NAMES = {"openai", "anthropic", "google", "xai", "ollama", "openrouter", "custom", "minimax", "qwen", "glm"}
    if not _raw_research_model or _raw_research_model in _PROVIDER_NAMES:
        if _raw_research_model in _PROVIDER_NAMES:
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"state.research_model='{_raw_research_model}' looks like a provider name, not a model ID — ignoring and using the selected deep-research model/default.",
            )
        research_model = selected_model or DEFAULT_DEEP_RESEARCH_MODEL
    else:
        research_model = _raw_research_model
    max_tool_calls = _safe_int(state.get("research_max_tool_calls"), 8, 1, 64)
    max_output_tokens = state.get("research_max_output_tokens")
    max_output_tokens_int = _safe_int(max_output_tokens, 0, 0, 200000) if max_output_tokens is not None else None
    poll_interval_seconds = _safe_int(state.get("research_poll_interval_seconds"), 5, 1, 60)
    max_wait_seconds = _safe_int(state.get("research_max_wait_seconds"), 3600, 60, 86400)
    heartbeat_seconds = _safe_int(state.get("research_heartbeat_seconds"), 120, 30, 3600)
    words_per_page = _safe_int(state.get("long_document_words_per_page"), 250, 150, 700)
    disable_visuals = bool(state.get("long_document_disable_visuals", False))
    include_section_references = bool(state.get("long_document_section_references", False))
    section_search_flag = state.get("long_document_section_search")
    web_search_enabled = bool(state.get("research_web_search_enabled", True))
    native_web_search_enabled = bool(web_search_enabled and supports_native_web_search(research_model, selected_provider))
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if native_web_search_enabled and not api_key:
        raise ValueError("OPENAI_API_KEY is required when the selected Deep Research model uses native web search.")
    if web_search_enabled and not native_web_search_enabled:
        log_task_update(
            DEEP_RESEARCH_LABEL,
            (
                f"{selected_provider or 'selected'} / {research_model} does not expose native web search here. "
                "Using the Kendr web search client for source gathering instead."
            ),
        )
    web_search_mode = (
        "native_model"
        if native_web_search_enabled
        else "kendr_search"
        if web_search_enabled
        else "disabled"
    )
    use_section_search = web_search_enabled if section_search_flag is None else bool(section_search_flag) and web_search_enabled
    section_search_results_count = _safe_int(state.get("long_document_section_search_results"), 6, 1, 20)
    deep_research_mode = bool(state.get("deep_research_mode", False))
    citation_style = _normalize_citation_style(state.get("research_citation_style", "apa"))
    output_formats = _normalize_research_formats(state.get("research_output_formats", DEFAULT_RESEARCH_FORMATS))
    plagiarism_enabled = bool(state.get("research_enable_plagiarism_check", True))
    date_range = str(state.get("research_date_range", "all_time") or "all_time").strip() or "all_time"
    requested_sources = [str(item).strip().lower() for item in state.get("research_sources", []) if str(item).strip()] if isinstance(state.get("research_sources", []), list) else []
    source_family_display = requested_sources or (["web"] if web_search_enabled else ["local"])
    search_backend = normalize_research_search_backend(state.get("research_search_backend", "auto"))
    max_sources = _safe_int(state.get("research_max_sources"), 0, 0, 400)
    checkpoint_enabled = bool(state.get("research_checkpoint_enabled", False))
    research_kb_enabled = bool(state.get("research_kb_enabled", False))
    research_kb_id = str(state.get("research_kb_id", "") or "").strip()
    research_kb_top_k = _safe_int(state.get("research_kb_top_k"), 8, 1, 50)
    intent, source_strategy = _deep_research_strategy(
        state,
        objective=objective,
        max_files=_safe_int(state.get("local_drive_max_files"), 200, 1, 1000),
    )

    research_instructions = str(
        state.get(
            "research_instructions",
            (
                "Perform exhaustive web-grounded research for this section. Favor primary and authoritative sources, "
                "surface disagreements, and avoid unsupported claims."
            ),
        )
    ).strip()

    log_task_update(
        DEEP_RESEARCH_LABEL,
        (
            f"Deep research pass #{call_number} started. depth={research_depth_mode}, "
            f"sections={section_count}, section_pages~{section_pages}, model={research_model}, "
            f"formats={','.join(output_formats)}, citation_style={citation_style}, search_backend={search_backend}"
        ),
        objective,
    )
    _trace_research_event(
        state,
        title="Deep research run started",
        detail=(
            f"Preparing a {research_depth_label} run across about {section_count} sections with "
            f"{citation_style.upper()} citations and {', '.join(output_formats)} exports."
        ),
        status="running",
        metadata={
            "phase": "startup",
            "call_number": call_number,
            "target_pages": target_pages,
            "section_count": section_count,
            "research_model": research_model,
            "web_search_mode": web_search_mode,
            "search_backend": search_backend,
            "research_kb_enabled": research_kb_enabled,
            "research_kb_id": research_kb_id,
        },
        subtask="Initialize deep research pipeline",
    )

    state["deep_research_mode"] = deep_research_mode
    state["long_document_mode"] = True
    state["workflow_type"] = "deep_research" if deep_research_mode else "long_document"
    state["research_provider"] = selected_provider
    state["research_model"] = research_model
    state["research_output_formats"] = output_formats
    state["research_citation_style"] = citation_style
    state["research_enable_plagiarism_check"] = plagiarism_enabled
    state["research_web_search_enabled"] = web_search_enabled
    state["research_web_search_mode"] = web_search_mode
    state["research_search_backend"] = search_backend
    state["research_date_range"] = date_range
    state["research_checkpoint_enabled"] = checkpoint_enabled
    if max_sources > 0:
        state["research_max_sources"] = max_sources
    state["deep_research_source_urls"] = _normalize_research_urls(state.get("deep_research_source_urls", []))

    collect_sources_first = bool(state.get("long_document_collect_sources_first", True if deep_research_mode else False))
    artifact_dir = f"deep_research_runs/deep_research_run_{call_number}"
    resume_reuse_enabled = bool(
        state.get("resume_source_run_id")
        or state.get("resume_requested", False)
        or state.get("resume_output_dir")
        or state.get("resume_checkpoint_payload")
    )
    if resume_reuse_enabled:
        resumed_artifact_dir = _resolve_resume_artifact_dir(
            state,
            call_number=call_number,
            default_artifact_dir=artifact_dir,
            objective=objective,
        )
        if resumed_artifact_dir != artifact_dir:
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Resume detected: reusing artifact directory {resumed_artifact_dir} instead of {artifact_dir}.",
            )
            artifact_dir = resumed_artifact_dir

    intent_json_path = _artifact_file(artifact_dir, "deep_research_intent.json")
    strategy_json_path = _artifact_file(artifact_dir, "deep_research_source_strategy.json")
    write_text_file(intent_json_path, json.dumps(intent, indent=2, ensure_ascii=False))
    write_text_file(strategy_json_path, json.dumps(source_strategy, indent=2, ensure_ascii=False))
    _trace_research_event(
        state,
        title="Research intent discovered",
        detail=str(intent.get("summary", "")).strip(),
        status="completed",
        kind="intent",
        metadata={"phase": "intent", "intent": intent},
        subtask="Discover research intent",
    )
    _trace_research_event(
        state,
        title="Source strategy planned",
        detail=str(source_strategy.get("summary", "")).strip(),
        status="completed",
        kind="source_strategy",
        metadata={"phase": "source_strategy", "strategy": source_strategy},
        subtask="Plan source strategy",
    )
    _ensure_local_source_manifest(
        state,
        objective=objective,
        artifact_dir=artifact_dir,
        source_strategy=source_strategy,
    )

    if bool(state.get("long_document_addendum_requested", False)):
        return _run_long_document_addendum(
            state,
            objective=objective,
            call_number=call_number,
            artifact_dir=artifact_dir,
        )

    analysis = state.get("deep_research_analysis", {})
    analysis_signature = {
        "objective": objective,
        "target_pages": target_pages,
        "depth_mode": research_depth_mode,
        "requested_sources": list(requested_sources),
        "date_range": date_range or "all_time",
        "max_sources": max_sources,
    }
    existing_signature = analysis.get("request_signature", {}) if isinstance(analysis, dict) else {}
    if not isinstance(analysis, dict) or not analysis or existing_signature != analysis_signature:
        _raise_if_cancelled(state, phase="analysis")
        analysis = _research_depth_analysis(
            objective,
            target_pages=target_pages,
            depth_mode=research_depth_mode,
            requested_sources=requested_sources,
            date_range=date_range,
            max_sources=max_sources,
        )
        state["deep_research_analysis"] = analysis
        state["deep_research_tier"] = int(analysis.get("tier", 0) or 0)

    if (
        deep_research_mode
        and bool(analysis.get("requires_deep_research", False))
        and not bool(state.get("deep_research_confirmed", False))
        and not bool(state.get("auto_approve", False))
    ):
        analysis_md = _deep_research_analysis_markdown(
            analysis,
            formats=output_formats,
            citation_style=citation_style,
            plagiarism_enabled=plagiarism_enabled,
            web_search_enabled=web_search_enabled,
            local_source_count=len(_collect_local_drive_evidence(state)),
            provided_url_count=len(_normalize_research_urls(state.get("deep_research_source_urls", []))) if web_search_enabled else 0,
        )
        version = int(state.get("deep_research_confirmation_version", 0) or 0) + 1
        analysis_json_path = _artifact_file(artifact_dir, "deep_research_analysis.json")
        analysis_md_path = _artifact_file(artifact_dir, "deep_research_analysis.md")
        write_text_file(analysis_json_path, json.dumps(analysis, indent=2, ensure_ascii=False))
        write_text_file(analysis_md_path, analysis_md)
        prompt = _deep_research_analysis_summary_prompt(
            title=title,
            analysis=analysis,
            formats=output_formats,
            citation_style=citation_style,
            plagiarism_enabled=plagiarism_enabled,
            web_search_enabled=web_search_enabled,
            local_source_count=len(_collect_local_drive_evidence(state)),
            provided_url_count=len(_normalize_research_urls(state.get("deep_research_source_urls", []))) if web_search_enabled else 0,
            analysis_storage_path=resolve_output_path(analysis_md_path),
            version=version,
        )
        approval_request = _build_deep_research_analysis_request(
            title=title,
            analysis=analysis,
            formats=output_formats,
            citation_style=citation_style,
            plagiarism_enabled=plagiarism_enabled,
            web_search_enabled=web_search_enabled,
            local_source_count=len(_collect_local_drive_evidence(state)),
            provided_url_count=len(_normalize_research_urls(state.get("deep_research_source_urls", []))) if web_search_enabled else 0,
            analysis_storage_path=resolve_output_path(analysis_md_path),
            version=version,
        )
        state["deep_research_confirmation_version"] = version
        state["pending_user_input_kind"] = "deep_research_confirmation"
        state["approval_pending_scope"] = "deep_research_confirmation"
        state["approval_request"] = approval_request
        state["pending_user_question"] = approval_request_to_text(approval_request) or prompt
        state["draft_response"] = state["pending_user_question"]
        state["deep_research_result_card"] = {
            "kind": "analysis",
            "title": title,
            "tier": analysis.get("tier", 0),
            "depth_mode": analysis.get("depth_mode", research_depth_mode),
            "depth_label": analysis.get("depth_label", research_depth_label),
            "estimated_sources": analysis.get("estimated_sources", 0),
            "estimated_duration_minutes": analysis.get("estimated_duration_minutes", 0),
            "subtopics": analysis.get("subtopics", []),
            "formats": output_formats,
            "citation_style": citation_style,
            "plagiarism_enabled": plagiarism_enabled,
            "web_search_enabled": web_search_enabled,
            "local_sources": len(_collect_local_drive_evidence(state)),
            "provided_urls": len(_normalize_research_urls(state.get("deep_research_source_urls", []))) if web_search_enabled else 0,
            "date_range": date_range,
            "intent_summary": intent.get("summary", ""),
            "strategy_summary": source_strategy.get("summary", ""),
            "research_kind": intent.get("research_kind", ""),
            "target_deliverable": intent.get("target_deliverable", ""),
            "source_needs": intent.get("source_needs", []),
            "risk_level": intent.get("risk_level", ""),
            "strategy_mode": source_strategy.get("mode", ""),
            "family_budgets": source_strategy.get("family_budgets", {}),
            "selection_rationale": source_strategy.get("selection_rationale", []),
        }
        update_planning_file(
            state,
            status="awaiting_deep_research_confirmation",
            objective=objective,
            plan_text=state.get("plan", ""),
            clarifications=state.get("plan_clarification_questions", []),
            execution_note="Deep research analysis completed; awaiting confirmation before expensive research starts.",
        )
        log_task_update(DEEP_RESEARCH_LABEL, "Prepared the deep research analysis card for approval.", analysis_md)
        _trace_research_event(
            state,
            title="Deep research analysis prepared",
            detail="Analysis card is ready for approval before the expensive research pipeline starts.",
            status="completed",
            metadata={"phase": "analysis", "requires_confirmation": True, "tier": analysis.get("tier", 0)},
            subtask="Review deep research analysis",
        )
        return publish_agent_output(
            state,
            "long_document_agent",
            state["draft_response"],
            f"deep_research_analysis_{version}",
            recipients=["orchestrator_agent", "reviewer_agent", "report_agent"],
        )

    if deep_research_mode:
        state["deep_research_confirmed"] = True
        if str(state.get("approval_pending_scope", "") or "").strip() == "deep_research_confirmation":
            state["approval_request"] = {}

    auto_approve_subplan = bool(state.get("auto_approve", False)) or bool(state.get("auto_approve_plan", False))

    approved_outline = state.get("long_document_outline", {})
    if not isinstance(approved_outline, dict):
        approved_outline = {}
    title = _normalize_title(approved_outline.get("title", ""), title)
    needs_outline_approval = (
        state.get("long_document_plan_status") != "approved"
        or not approved_outline
        or bool(state.get("long_document_replan_requested", False))
    )

    if needs_outline_approval:
        _raise_if_cancelled(state, phase="outline_generation")
        outline_objective = objective
        feedback = str(state.get("long_document_plan_feedback", "") or "").strip()
        if feedback:
            outline_objective = (
                f"{objective}\n\nUser requested these section-plan changes:\n{feedback}"
            )
        outline = _build_outline(
            outline_objective,
            title=title,
            section_count=section_count,
            section_pages=section_pages,
            subtopics=analysis.get("subtopics", []),
        )
        title = _normalize_title(outline.get("title", ""), title)
        outline_md = _outline_markdown(outline, fallback_title=title)
        subplan_data = _long_document_subplan(outline, objective=objective, target_pages=target_pages, research_model=research_model)
        subplan_md = outline_md.rstrip() + "\n\n" + plan_as_markdown(subplan_data)
        subplan_version = int(state.get("long_document_plan_version", 0) or 0) + 1

        write_text_file(_artifact_file(artifact_dir, "deep_research_outline.json"), json.dumps(outline, indent=2, ensure_ascii=False))
        write_text_file(_artifact_file(artifact_dir, "deep_research_outline.md"), outline_md)
        write_text_file(_artifact_file(artifact_dir, "deep_research_subplan.json"), json.dumps(subplan_data, indent=2, ensure_ascii=False))
        write_text_file(_artifact_file(artifact_dir, "deep_research_subplan.md"), subplan_md + "\n")

        outline_storage_path = resolve_output_path(_artifact_file(artifact_dir, "deep_research_outline.md"))
        subplan_storage_path = resolve_output_path(_artifact_file(artifact_dir, "deep_research_subplan.md"))
        approval_prompt = _deep_research_subplan_summary_prompt(
            title=title,
            analysis=analysis,
            outline=outline,
            target_pages=target_pages,
            section_pages=section_pages,
            formats=output_formats,
            web_search_enabled=web_search_enabled,
            outline_storage_path=outline_storage_path,
            subplan_storage_path=subplan_storage_path,
            version=subplan_version,
        )
        approval_request = _build_deep_research_subplan_request(
            title=title,
            analysis=analysis,
            outline=outline,
            target_pages=target_pages,
            section_pages=section_pages,
            formats=output_formats,
            web_search_enabled=web_search_enabled,
            outline_storage_path=outline_storage_path,
            subplan_storage_path=subplan_storage_path,
            version=subplan_version,
        )

        state["long_document_outline"] = outline
        state["long_document_plan_data"] = subplan_data
        state["long_document_plan_markdown"] = subplan_md
        state["long_document_plan_version"] = subplan_version
        state["long_document_outline_md_path"] = outline_storage_path
        state["long_document_subplan_md_path"] = subplan_storage_path
        state["long_document_replan_requested"] = False
        state["deep_research_result_card"] = {
            "kind": "plan",
            "status": "approved" if auto_approve_subplan else "pending_approval",
            "title": title,
            "tier": analysis.get("tier", 0),
            "depth_mode": analysis.get("depth_mode", research_depth_mode),
            "depth_label": analysis.get("depth_label", research_depth_label),
            "subtopics": analysis.get("subtopics", []),
            "section_count": len(outline.get("sections", [])),
            "formats": output_formats,
            "web_search_enabled": web_search_enabled,
            "intent_summary": intent.get("summary", ""),
            "strategy_summary": source_strategy.get("summary", ""),
            "research_kind": intent.get("research_kind", ""),
            "target_deliverable": intent.get("target_deliverable", ""),
            "source_needs": intent.get("source_needs", []),
            "risk_level": intent.get("risk_level", ""),
            "strategy_mode": source_strategy.get("mode", ""),
            "family_budgets": source_strategy.get("family_budgets", {}),
            "selection_rationale": source_strategy.get("selection_rationale", []),
        }

        if auto_approve_subplan:
            approved_outline = outline
            state["long_document_plan_waiting_for_approval"] = False
            state["long_document_plan_status"] = "approved"
            state["long_document_execute_from_saved_outline"] = True
            state["pending_user_input_kind"] = ""
            state["approval_pending_scope"] = ""
            state["approval_request"] = {}
            state["pending_user_question"] = ""
            state["draft_response"] = _deep_research_subplan_auto_approved_message(
                title=title,
                outline=outline,
                analysis=analysis,
                target_pages=target_pages,
                section_pages=section_pages,
                web_search_enabled=web_search_enabled,
                outline_storage_path=outline_storage_path,
                subplan_storage_path=subplan_storage_path,
                version=subplan_version,
            )
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Auto-approved section plan v{subplan_version}; continuing to execution.",
            )
            if _deep_research_handoff_only_enabled(state):
                update_planning_file(
                    state,
                    status="subplan_auto_approved",
                    objective=objective,
                    plan_text=state.get("plan", ""),
                    clarifications=state.get("plan_clarification_questions", []),
                    execution_note=(
                        f"Deep research subplan v{subplan_version} auto-approved and handoff prepared for evidence collection."
                    ),
                )
                _trace_research_event(
                    state,
                    title="Research handoff prepared",
                    detail=(
                        f"Auto-approved subplan v{subplan_version} and stopped after the brief-to-report handoff."
                    ),
                    status="completed",
                    metadata={"phase": "planning", "section_count": len(outline.get("sections", [])), "version": subplan_version},
                    subtask="Prepare execution handoff",
                )
                return publish_agent_output(
                    state,
                    "long_document_agent",
                    state["draft_response"],
                    f"deep_research_subplan_{subplan_version}",
                    recipients=["orchestrator_agent", "reviewer_agent", "report_agent"],
                )
        else:
            state["long_document_plan_waiting_for_approval"] = True
            state["long_document_plan_status"] = "pending"
            state["pending_user_input_kind"] = "subplan_approval"
            state["approval_pending_scope"] = "long_document_plan"
            state["approval_request"] = approval_request
            state["pending_user_question"] = approval_request_to_text(approval_request) or approval_prompt
            state["draft_response"] = state["pending_user_question"]
            state["_skip_review_once"] = True
            state["_hold_planned_step_completion_once"] = True

        if not auto_approve_subplan:
            update_planning_file(
                state,
                status="awaiting_subplan_approval",
                objective=objective,
                plan_text=state.get("plan", ""),
                clarifications=state.get("plan_clarification_questions", []),
                execution_note=f"Deep research subplan v{subplan_version} generated and queued for approval.",
            )
            log_task_update(DEEP_RESEARCH_LABEL, "Prepared a section-by-section deep research plan for approval.", subplan_md)
            _trace_research_event(
                state,
                title="Research section plan prepared",
                detail=f"Generated a section-by-section plan with {len(outline.get('sections', []))} sections and queued it for approval.",
                status="completed",
                metadata={"phase": "planning", "section_count": len(outline.get("sections", [])), "version": subplan_version},
                subtask="Review section plan",
            )
            return publish_agent_output(
                state,
                "long_document_agent",
                state["draft_response"],
                f"deep_research_subplan_{subplan_version}",
                recipients=["orchestrator_agent", "reviewer_agent", "report_agent"],
            )
    if str(state.get("approval_pending_scope", "") or "").strip() == "long_document_plan":
        state["approval_request"] = {}

    evidence_bank_md = ""
    evidence_sources: list[dict] = []
    evidence_excerpt = ""
    local_entries: list[dict] = _collect_local_drive_evidence(state)
    url_entries: list[dict] = _collect_user_url_evidence(objective, state)
    kb_grounding: dict[str, Any] = {"requested": research_kb_enabled}
    kb_warning = ""
    if research_kb_enabled:
        try:
            kb_grounding = build_research_grounding(
                objective,
                kb_ref=research_kb_id,
                top_k=research_kb_top_k,
                use_active_if_empty=True,
                require_indexed=True,
            )
            kb_grounding["requested"] = True
            if int(kb_grounding.get("hit_count", 0) or 0) <= 0:
                kb_warning = (
                    f"Knowledge base '{kb_grounding.get('kb_name', 'unknown')}' returned no relevant results for this report objective."
                )
        except Exception as exc:
            kb_warning = str(exc)
            has_other_report_evidence = bool(web_search_enabled or local_entries or url_entries)
            if has_other_report_evidence:
                log_task_update(
                    DEEP_RESEARCH_LABEL,
                    f"Knowledge base grounding unavailable; continuing with other report sources. {kb_warning}",
                )
            else:
                raise ValueError(
                    f"Knowledge base grounding failed and no other report evidence sources are available. {kb_warning}"
                ) from exc
    explicit_source_entries = _merge_sources(
        _sources_from_local_entries(local_entries),
        _sources_from_url_entries(url_entries),
        max_sources=max_sources or 120,
    )
    explicit_source_entries = _merge_sources(
        explicit_source_entries,
        _sources_from_kb_grounding(kb_grounding),
        max_sources=max_sources or 120,
    )
    if research_kb_enabled and not web_search_enabled and not local_entries and not url_entries and not explicit_source_entries:
        raise ValueError(
            (
                f"Knowledge base '{str(kb_grounding.get('kb_name', '') or research_kb_id or 'active KB')}' "
                "returned no usable evidence and no other sources were provided for this deep research run."
            ).strip()
        )
    if collect_sources_first:
        cached_evidence = _load_cached_evidence_bank(
            artifact_dir=artifact_dir,
            objective=objective,
            max_sources=max_sources,
        )
        if cached_evidence:
            evidence_bank_md = str(cached_evidence.get("evidence_markdown", "") or "")
            evidence_sources = list(cached_evidence.get("evidence_sources", []) or [])
            evidence_excerpt = _trim_text(evidence_bank_md, 18000) if evidence_bank_md else ""
            state["long_document_sources_collected"] = True
            state["long_document_evidence_bank_path"] = str(cached_evidence.get("evidence_path", "") or "")
            state["long_document_evidence_bank_json_path"] = str(cached_evidence.get("evidence_json_path", "") or "")
            state["long_document_evidence_bank_excerpt"] = evidence_excerpt
            state["long_document_evidence_sources"] = evidence_sources
            _trace_research_event(
                state,
                title="Collecting evidence bank",
                detail=f"Reused existing evidence bank from {artifact_dir}.",
                command=objective,
                status="completed",
                metadata={
                    "phase": "evidence_bank",
                    "sources": len(evidence_sources),
                    "resumed_from_cache": True,
                },
                subtask="Build cross-report evidence bank",
            )
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Reused cached evidence bank with {len(evidence_sources)} source(s).",
            )
    if collect_sources_first:
        if not state.get("long_document_sources_collected", False):
            evidence_started_at = _trace_now()
            _trace_research_event(
                state,
                title="Collecting evidence bank",
                detail="Gathering shared sources, contradictions, and benchmarks before section drafting begins.",
                command=objective,
                status="running",
                started_at=evidence_started_at,
                metadata={"phase": "evidence_bank", "web_search_enabled": web_search_enabled},
                subtask="Build cross-report evidence bank",
            )
            log_task_update(DEEP_RESEARCH_LABEL, "Collecting the cross-report evidence bank.")
            if web_search_enabled:
                log_task_update(
                    DEEP_RESEARCH_LABEL,
                    f"Searching the web for shared report evidence via {search_backend}.",
                )
                search_results = _collect_google_search_evidence(
                    objective,
                    num=int(state.get("long_document_search_results", 10) or 10),
                    search_backend=search_backend,
                    progress_callback=lambda url, status, payload, position, total: _log_web_review(
                        url,
                        status=status,
                        position=position,
                        total=total,
                        payload=payload,
                        context="search result website",
                    ),
                )
                _trace_research_event(
                    state,
                    title="Google search results gathered",
                    detail=f"Cross-report search returned {len((search_results or {}).get('results', []))} result(s).",
                    command=objective,
                    status="completed" if not str((search_results or {}).get("error", "")).strip() else "failed",
                    metadata={
                        "phase": "evidence_bank_search",
                        "search_query": objective,
                        "search_provider": str((search_results or {}).get("provider", "")).strip(),
                        "search_providers_tried": list((search_results or {}).get("providers_tried", []) or []),
                        "candidate_urls": _trace_url_list([str(item.get("url", "")).strip() for item in (search_results or {}).get("results", [])]),
                        "viewed_urls": _trace_url_list([str(item.get("url", "")).strip() for item in (search_results or {}).get("viewed_pages", [])]),
                        "urls": _trace_url_list([str(item.get("url", "")).strip() for item in (search_results or {}).get("results", [])]),
                    },
                        subtask="Search for shared report evidence",
                    )
                log_task_update(
                    DEEP_RESEARCH_LABEL,
                    (
                        f"Web search gathered {len((search_results or {}).get('results', []))} result(s); "
                        f"provider={str((search_results or {}).get('provider', '') or 'none')}, "
                        f"providers_tried={', '.join(list((search_results or {}).get('providers_tried', []) or [])) or 'none'}."
                    ),
                )
                evidence_instructions = str(
                    state.get(
                        "long_document_evidence_instructions",
                        (
                            "Collect a comprehensive evidence bank with source-backed facts, contradictions, and benchmarks. "
                            "Prefer primary sources, provide concrete numbers, and include URLs for every major claim."
                        ),
                    )
                ).strip()
                evidence_query_lines = [
                    f"Objective: {objective}",
                    f"Date range preference: {date_range}",
                    f"Requested source families: {source_family_display}",
                    "Local drive rollup summary:",
                    _trim_text(state.get("local_drive_rollup_summary", ""), 2000),
                    _kb_grounding_summary(kb_grounding),
                    "Task: Build a source-backed evidence bank with URLs, disagreements, numbers, and explicit citations.",
                ]
                if native_web_search_enabled:
                    evidence_pass = _run_research_pass(
                        OpenAI(api_key=api_key),
                        query="\n".join(line for line in evidence_query_lines if str(line).strip()),
                        model=research_model,
                        instructions=evidence_instructions,
                        max_tool_calls=max_tool_calls,
                        max_output_tokens=max_output_tokens_int,
                        poll_interval_seconds=poll_interval_seconds,
                        max_wait_seconds=max_wait_seconds,
                        heartbeat_interval_seconds=heartbeat_seconds,
                        heartbeat_label="Evidence bank research in progress",
                        heartbeat_callback=lambda elapsed: _trace_research_event(
                            state,
                            title="Collecting evidence bank",
                            detail=f"Shared evidence bank research still running ({elapsed}s elapsed).",
                            command=objective,
                            status="running",
                            metadata={"phase": "evidence_bank", "elapsed_seconds": elapsed},
                            subtask="Build cross-report evidence bank",
                        ),
                        cancel_check=lambda: _raise_if_cancelled(state, phase="evidence_bank"),
                    )
                    evidence_sources = _merge_sources(
                        explicit_source_entries,
                        _extract_source_entries(evidence_pass, max_sources=max_sources or 60),
                        max_sources=max_sources or 120,
                    )
                    evidence_notes_text = str(evidence_pass.get("output_text", "")).strip()
                else:
                    evidence_pass = {
                        "response_id": "kendr_search_evidence_bank",
                        "status": "fallback_web_search",
                        "elapsed_seconds": 0,
                        "output_text": _build_local_only_research_notes(
                            objective=objective,
                            focus="overall report evidence bank",
                            local_entries=local_entries,
                            url_entries=url_entries,
                        ),
                        "raw": {
                            "reason": "native_web_search_unavailable",
                            "search_provider": str((search_results or {}).get("provider", "")).strip(),
                            "search_providers_tried": list((search_results or {}).get("providers_tried", []) or []),
                        },
                    }
                    evidence_sources = _merge_sources(
                        explicit_source_entries,
                        _sources_from_kb_grounding(kb_grounding),
                        max_sources=max_sources or 120,
                    )
                    evidence_sources = _merge_sources(
                        evidence_sources,
                        _sources_from_search_results(search_results.get("results", [])),
                        max_sources=max_sources or 120,
                    )
                    evidence_notes_text = str(evidence_pass.get("output_text", "")).strip()
                evidence_pass["research_kb"] = kb_grounding
                evidence_pass["research_kb_warning"] = kb_warning
            else:
                search_results = {"results": [], "error": "Web search disabled by user."}
                evidence_pass = {
                    "response_id": "local_only",
                    "status": "local_only",
                    "elapsed_seconds": 0,
                    "output_text": _kb_grounding_summary(kb_grounding),
                    "raw": {"reason": "web_search_disabled"},
                }
                evidence_pass["research_kb"] = kb_grounding
                evidence_pass["research_kb_warning"] = kb_warning
                evidence_sources = list(explicit_source_entries)
                evidence_notes_text = _build_local_only_research_notes(
                    objective=objective,
                    focus="overall report evidence bank",
                    local_entries=local_entries,
                    url_entries=url_entries,
                )
            evidence_bank_md = _format_evidence_bank_markdown(
                objective=objective,
                local_entries=local_entries,
                url_entries=url_entries,
                kb_grounding=kb_grounding,
                kb_warning=kb_warning,
                search_results=search_results,
                research_notes_text=evidence_notes_text,
                source_ledger=evidence_sources,
                insufficiency_note="" if web_search_enabled else "Web search was disabled. This report used only local files, explicit links, and any selected knowledge-base evidence.",
            )
            evidence_payload = {
                "objective": objective,
                "local_drive_entries": local_entries,
                "provided_url_entries": url_entries,
                "research_kb": kb_grounding,
                "research_kb_warning": kb_warning,
                "search_results": search_results.get("results", []) if isinstance(search_results, dict) else [],
                "web_research": evidence_pass,
                "source_ledger": evidence_sources,
                "analysis": analysis,
            }
            evidence_md_filename = _artifact_file(artifact_dir, "evidence_bank.md")
            evidence_json_filename = _artifact_file(artifact_dir, "evidence_bank.json")
            write_text_file(evidence_md_filename, evidence_bank_md)
            write_text_file(evidence_json_filename, json.dumps(evidence_payload, indent=2, ensure_ascii=False))
            if url_entries:
                write_text_file(_artifact_file(artifact_dir, "provided_url_sources.json"), json.dumps(url_entries, indent=2, ensure_ascii=False))
            state["long_document_sources_collected"] = True
            state["long_document_evidence_bank_path"] = _output_file_path(evidence_md_filename)
            state["long_document_evidence_bank_json_path"] = _output_file_path(evidence_json_filename)
            evidence_excerpt = _trim_text(evidence_bank_md, 18000)
            state["long_document_evidence_bank_excerpt"] = evidence_excerpt
            state["long_document_evidence_sources"] = evidence_sources
            evidence_status = str(evidence_pass.get("status", "")).strip() or "completed"
            evidence_completed = "completed" if evidence_status in {"completed", "local_only", "evidence_bank", "fallback_web_search"} else "failed"
            evidence_detail = (
                f"Evidence bank ready with {len(evidence_sources)} consolidated sources."
                if evidence_completed == "completed"
                else f"Evidence bank finished with status '{evidence_status}'. Partial results will be used."
            )
            _trace_research_event(
                state,
                title="Collecting evidence bank",
                detail=evidence_detail,
                command=objective,
                status=evidence_completed,
                started_at=evidence_started_at,
                completed_at=_trace_now(),
                metadata={
                    "phase": "evidence_bank",
                    "response_status": evidence_status,
                    "sources": len(evidence_sources),
                    "elapsed_seconds": int(evidence_pass.get("elapsed_seconds", 0) or 0),
                    "urls": _trace_url_list([str(item.get("url", "")).strip() for item in evidence_sources]),
                },
                subtask="Build cross-report evidence bank",
            )
        else:
            evidence_path = _resolve_existing_output_path(state.get("long_document_evidence_bank_path", ""))
            evidence_bank_md = _read_text_file(evidence_path, "")
            evidence_excerpt = state.get("long_document_evidence_bank_excerpt", "") or _trim_text(evidence_bank_md, 18000)
            evidence_sources = state.get("long_document_evidence_sources", []) or []

    outline = approved_outline
    outline_md = _outline_markdown(outline, fallback_title=title)
    write_text_file(_artifact_file(artifact_dir, "deep_research_outline.json"), json.dumps(outline, indent=2, ensure_ascii=False))
    write_text_file(_artifact_file(artifact_dir, "deep_research_outline.md"), outline_md)
    update_planning_file(
        state,
        status="executing",
        objective=objective,
        plan_text=state.get("plan", ""),
        clarifications=state.get("plan_clarification_questions", []),
        execution_note="Deep research plan approved. Executing research, correlation, writing, and verification phases.",
    )

    research_log_lines = [
        f"Tier {analysis.get('tier', 0)} analysis confirmed.",
        f"Research depth: {analysis.get('depth_label', research_depth_label)}",
        f"Output formats: {', '.join(output_formats)}",
        f"Citation style: {citation_style.upper()}",
        f"Plagiarism check: {'enabled' if plagiarism_enabled else 'disabled'}",
        f"Date range: {date_range}",
    ]

    continuity_notes: list[str] = []
    section_outputs: list[dict] = []
    section_packages: list[dict] = []
    all_section_images: list[dict] = []
    image_collection_enabled = not disable_visuals and bool(state.get("long_document_enable_source_images", True))
    image_search_fallback_enabled = bool(state.get("long_document_image_search_fallback", False))
    image_vision_verification_enabled = bool(state.get("long_document_verify_source_images", True))
    if image_collection_enabled:
        log_task_update(
            DEEP_RESEARCH_LABEL,
            "Image collection enabled: preferring source-grounded images from reviewed web pages.",
        )
        if image_vision_verification_enabled:
            log_task_update(
                DEEP_RESEARCH_LABEL,
                "Vision verification is enabled for downloaded images when a vision backend is available.",
            )
        if image_search_fallback_enabled:
            log_task_update(
                DEEP_RESEARCH_LABEL,
                "External image-search fallback is enabled if no grounded page image qualifies.",
            )
    else:
        log_task_update(DEEP_RESEARCH_LABEL, "Image collection disabled.")
    coherence_context_md = _coherence_base_context(state, objective)
    write_text_file(_artifact_file(artifact_dir, "deep_research_coherence_base.md"), coherence_context_md)

    parallel_sections = outline.get("sections", []) if isinstance(outline.get("sections", []), list) else []
    total_parallel_sections = len(parallel_sections)
    section_parallelism = min(
        max(total_parallel_sections, 1),
        _parallelism(
            state.get("research_section_concurrency"),
            env_key="KENDR_RESEARCH_SECTION_CONCURRENCY",
            default=3,
            minimum=1,
            maximum=8,
        ),
    )
    if total_parallel_sections:
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"Phase 1/4 - researching {total_parallel_sections} sections with {section_parallelism} worker(s).",
        )
        section_started_at_by_index: dict[int, str] = {}
        completed_section_packages: dict[int, dict[str, Any]] = {}

        def _finalize_section_package(package: dict[str, Any]) -> None:
            recorded = _record_section_research_package(
                state,
                package=package,
                total_sections=total_parallel_sections,
                artifact_dir=artifact_dir,
                research_log_lines=research_log_lines,
                section_research_started_at=section_started_at_by_index.get(int(package.get("index", 0) or 0), _trace_now()),
            )
            completed_section_packages[int(recorded.get("index", 0) or 0)] = recorded

        reused_research_sections = 0
        if resume_reuse_enabled:
            for index, section in enumerate(parallel_sections, start=1):
                cached_package = _load_cached_section_research_package(
                    artifact_dir=artifact_dir,
                    objective=objective,
                    section=section,
                    section_index=index,
                    section_pages=section_pages,
                    collect_sources_first=collect_sources_first,
                    evidence_excerpt=evidence_excerpt,
                    evidence_sources=evidence_sources,
                    explicit_source_entries=explicit_source_entries,
                    max_sources=max_sources,
                )
                if not cached_package:
                    continue
                section_title = str(section.get("title", f"Section {index}")).strip() or f"Section {index}"
                section_started_at_by_index[index] = _trace_now()
                log_task_update(
                    DEEP_RESEARCH_LABEL,
                    f"Phase 1/4 - reusing cached research for section {index}/{total_parallel_sections}: {section_title}",
                )
                _finalize_section_package(cached_package)
                reused_research_sections += 1
        if reused_research_sections:
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Resume reuse: skipped fresh research for {reused_research_sections}/{total_parallel_sections} section(s).",
            )

        if section_parallelism <= 1:
            for index, section in enumerate(parallel_sections, start=1):
                if index in completed_section_packages:
                    continue
                section_title = str(section.get("title", f"Section {index}")).strip() or f"Section {index}"
                section_objective = str(section.get("objective", objective)).strip() or objective
                log_task_update(
                    DEEP_RESEARCH_LABEL,
                    f"Phase 1/4 - researching section {index}/{total_parallel_sections}: {section_title}",
                )
                section_started_at_by_index[index] = _trace_now()
                _trace_research_event(
                    state,
                    title=f"Researching section {index}/{total_parallel_sections}",
                    detail=section_title,
                    command=section_objective,
                    status="running",
                    started_at=section_started_at_by_index[index],
                    metadata={"phase": "section_research", "section_index": index, "section_title": section_title},
                    subtask=f"Gather evidence for {section_title}",
                )
                _finalize_section_package(
                    _collect_section_research_package(
                        api_key=api_key,
                        objective=objective,
                        section=section,
                        section_index=index,
                        total_sections=total_parallel_sections,
                        section_pages=section_pages,
                        use_section_search=use_section_search,
                        section_search_results_count=section_search_results_count,
                        collect_sources_first=collect_sources_first,
                        evidence_excerpt=evidence_excerpt,
                        evidence_sources=evidence_sources,
                        explicit_source_entries=explicit_source_entries,
                        local_entries=local_entries,
                        url_entries=url_entries,
                        continuity_notes=continuity_notes,
                        coherence_context_md=coherence_context_md,
                        web_search_enabled=web_search_enabled,
                        native_web_search_enabled=native_web_search_enabled,
                        research_model=research_model,
                        research_instructions=research_instructions,
                        max_tool_calls=max_tool_calls,
                        max_output_tokens_int=max_output_tokens_int,
                        poll_interval_seconds=poll_interval_seconds,
                        max_wait_seconds=max_wait_seconds,
                        heartbeat_seconds=heartbeat_seconds,
                        max_sources=max_sources,
                        research_kb_enabled=research_kb_enabled,
                        research_kb_id=research_kb_id,
                        research_kb_top_k=research_kb_top_k,
                        cancel_check=lambda idx=index: _raise_if_cancelled(state, phase=f"section_{idx}_research"),
                    )
                )
        else:
            with ThreadPoolExecutor(max_workers=section_parallelism, thread_name_prefix="kendr-section") as executor:
                future_map = {}
                for index, section in enumerate(parallel_sections, start=1):
                    if index in completed_section_packages:
                        continue
                    section_title = str(section.get("title", f"Section {index}")).strip() or f"Section {index}"
                    section_objective = str(section.get("objective", objective)).strip() or objective
                    log_task_update(
                        DEEP_RESEARCH_LABEL,
                        f"Phase 1/4 - queued section {index}/{total_parallel_sections}: {section_title}",
                    )
                    section_started_at_by_index[index] = _trace_now()
                    _trace_research_event(
                        state,
                        title=f"Researching section {index}/{total_parallel_sections}",
                        detail=f"{section_title} queued on the research worker pool.",
                        command=section_objective,
                        status="running",
                        started_at=section_started_at_by_index[index],
                        metadata={
                            "phase": "section_research",
                            "section_index": index,
                            "section_title": section_title,
                            "parallelism": section_parallelism,
                        },
                        subtask=f"Gather evidence for {section_title}",
                    )
                    future = executor.submit(
                        _collect_section_research_package,
                        api_key=api_key,
                        objective=objective,
                        section=section,
                        section_index=index,
                        total_sections=total_parallel_sections,
                        section_pages=section_pages,
                        use_section_search=use_section_search,
                        section_search_results_count=section_search_results_count,
                        collect_sources_first=collect_sources_first,
                        evidence_excerpt=evidence_excerpt,
                        evidence_sources=evidence_sources,
                        explicit_source_entries=explicit_source_entries,
                        local_entries=local_entries,
                        url_entries=url_entries,
                        continuity_notes=continuity_notes,
                        coherence_context_md=coherence_context_md,
                        web_search_enabled=web_search_enabled,
                        native_web_search_enabled=native_web_search_enabled,
                        research_model=research_model,
                        research_instructions=research_instructions,
                        max_tool_calls=max_tool_calls,
                        max_output_tokens_int=max_output_tokens_int,
                        poll_interval_seconds=poll_interval_seconds,
                        max_wait_seconds=max_wait_seconds,
                        heartbeat_seconds=heartbeat_seconds,
                        max_sources=max_sources,
                        research_kb_enabled=research_kb_enabled,
                        research_kb_id=research_kb_id,
                        research_kb_top_k=research_kb_top_k,
                        cancel_check=lambda idx=index: _raise_if_cancelled(state, phase=f"section_{idx}_research"),
                    )
                    future_map[future] = index

                for future in as_completed(future_map):
                    _raise_if_cancelled(state, phase="section_research")
                    _finalize_section_package(future.result())

        _raise_if_cancelled(state, phase="section_research_complete")
        section_packages = [completed_section_packages[index] for index in sorted(completed_section_packages)]

    sections: list[dict[str, Any]] = []
    total_sections = len(sections)
    for index, section in enumerate(sections, start=1):
        section_title = str(section.get("title", f"Section {index}")).strip() or f"Section {index}"
        section_objective = str(section.get("objective", objective)).strip() or objective
        section_questions = section.get("key_questions", [])
        if not isinstance(section_questions, list):
            section_questions = []
        target_section_pages = _safe_int(section.get("target_pages"), section_pages, 1, 30)

        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"Phase 1/4 - researching section {index}/{total_sections}: {section_title}",
        )
        section_research_started_at = _trace_now()
        _trace_research_event(
            state,
            title=f"Researching section {index}/{total_sections}",
            detail=section_title,
            command=section_objective,
            status="running",
            started_at=section_research_started_at,
            metadata={"phase": "section_research", "section_index": index, "section_title": section_title},
            subtask=f"Gather evidence for {section_title}",
        )

        section_search_results: dict[str, Any] = {}
        section_search_sources: list[dict] = []
        section_search_text = ""
        if use_section_search:
            search_query = f"{section_title} {section_objective}"
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Searching the web for section {index} via {search_backend}: {section_title}",
            )
            section_search_results = _collect_google_search_evidence(
                search_query,
                num=section_search_results_count,
                search_backend=search_backend,
                progress_callback=lambda url, status, payload, position, total: _log_web_review(
                    url,
                    status=status,
                    position=position,
                    total=total,
                    payload=payload,
                    context="search result website",
                ),
            )
            _trace_research_event(
                state,
                title=f"Google search results for section {index}",
                detail=f"{section_title} returned {len((section_search_results or {}).get('results', []))} result(s).",
                command=search_query,
                status="completed" if not str((section_search_results or {}).get("error", "")).strip() else "failed",
                metadata={
                    "phase": "section_search",
                    "section_index": index,
                    "section_title": section_title,
                    "search_query": search_query,
                    "search_provider": str((section_search_results or {}).get("provider", "")).strip(),
                    "search_providers_tried": list((section_search_results or {}).get("providers_tried", []) or []),
                    "candidate_urls": _trace_url_list([str(item.get("url", "")).strip() for item in (section_search_results or {}).get("results", [])]),
                    "viewed_urls": _trace_url_list([str(item.get("url", "")).strip() for item in (section_search_results or {}).get("viewed_pages", [])]),
                    "urls": _trace_url_list([str(item.get("url", "")).strip() for item in (section_search_results or {}).get("results", [])]),
                },
                subtask=f"Search for {section_title}",
            )
            log_task_update(
                DEEP_RESEARCH_LABEL,
                (
                    f"Section {index} web search gathered {len((section_search_results or {}).get('results', []))} result(s); "
                    f"provider={str((section_search_results or {}).get('provider', '') or 'none')}, "
                    f"providers_tried={', '.join(list((section_search_results or {}).get('providers_tried', []) or [])) or 'none'}."
                ),
            )
            if isinstance(section_search_results, dict) and str(section_search_results.get("error", "")).strip():
                log_task_update(
                    DEEP_RESEARCH_LABEL,
                    f"Section {index} web search error: {section_search_results.get('error')}",
                )
            section_search_sources = _sources_from_search_results(section_search_results.get("results", []))
            section_search_text = _search_results_markdown(section_search_results.get("results", []))

        if collect_sources_first and evidence_excerpt:
            local_section_notes = _build_local_only_research_notes(
                objective=objective,
                focus=f"section {index}: {section_title} — {section_objective}",
                local_entries=local_entries,
                url_entries=url_entries,
                continuity_notes=continuity_notes,
            ) if explicit_source_entries else ""
            research_pass = {
                "response_id": "evidence_bank" if web_search_enabled else f"local_only_section_{index}",
                "status": "evidence_bank" if web_search_enabled else "local_only",
                "elapsed_seconds": 0,
                "output_text": local_section_notes or evidence_excerpt,
                "raw": {"evidence_bank_path": state.get("long_document_evidence_bank_path", "")},
            }
            research_output = local_section_notes or evidence_excerpt
            if section_search_text:
                research_output = f"{section_search_text}\n\n{research_output}"
            section_sources = list(evidence_sources)
            section_sources = _merge_sources(section_sources, section_search_sources)
            source_ledger_md = _source_ledger_markdown(section_sources)
        else:
            local_section_notes = _build_local_only_research_notes(
                objective=objective,
                focus=f"section {index}: {section_title} — {section_objective}",
                local_entries=local_entries,
                url_entries=url_entries,
                continuity_notes=continuity_notes,
            ) if explicit_source_entries else ""
            if web_search_enabled:
                query_lines = [
                    f"Global objective: {objective}",
                    f"Section {index} title: {section_title}",
                    f"Section objective: {section_objective}",
                    "Key questions:",
                    *[f"- {item}" for item in section_questions[:12]],
                    "Continuity notes from previous sections:",
                    *[f"- {item}" for item in continuity_notes[-10:]],
                    "Markdown coherence anchors:",
                    _trim_text(coherence_context_md, 4000),
                ]
                research_pass = _run_research_pass(
                    client,
                    query="\n".join(query_lines),
                    model=research_model,
                    instructions=research_instructions,
                    max_tool_calls=max_tool_calls,
                    max_output_tokens=max_output_tokens_int,
                    poll_interval_seconds=poll_interval_seconds,
                    max_wait_seconds=max_wait_seconds,
                    heartbeat_interval_seconds=heartbeat_seconds,
                    heartbeat_label=f"Section {index} research in progress",
                    heartbeat_callback=lambda elapsed, idx=index, name=section_title: _trace_research_event(
                        state,
                        title=f"Researching section {idx}/{total_sections}",
                        detail=f"{name} still running ({elapsed}s elapsed).",
                        command=section_objective,
                        status="running",
                        metadata={"phase": "section_research", "section_index": idx, "section_title": name, "elapsed_seconds": elapsed},
                        subtask=f"Gather evidence for {name}",
                    ),
                    cancel_check=lambda idx=index: _raise_if_cancelled(state, phase=f"section_{idx}_research"),
                )

                if str(research_pass.get("status", "")).strip() not in {"completed", ""}:
                    log_task_update(
                        DEEP_RESEARCH_LABEL,
                        f"Section {index} research status: {research_pass.get('status')}",
                    )

                research_output = str(research_pass.get("output_text", "")).strip()
            else:
                research_pass = {
                    "response_id": f"local_only_section_{index}",
                    "status": "local_only",
                    "elapsed_seconds": 0,
                    "output_text": local_section_notes,
                    "raw": {"reason": "web_search_disabled"},
                }
                research_output = local_section_notes
            if section_search_text:
                research_output = f"{section_search_text}\n\n{research_output}".strip()
            if not research_output:
                if section_search_text:
                    research_output = section_search_text
                else:
                    research_output = "Research output was empty. Use only explicitly supported claims and call out uncertainty."

            section_sources = _merge_sources(explicit_source_entries, _extract_source_entries(research_pass))
            section_sources = _merge_sources(section_sources, section_search_sources)
            source_ledger_md = _source_ledger_markdown(section_sources)
        if max_sources > 0:
            section_sources = section_sources[:max_sources]
        section_research_status = str(research_pass.get("status", "")).strip() or "completed"
        _trace_research_event(
            state,
            title=f"Researching section {index}/{total_sections}",
            detail=(
                f"{section_title} completed with {len(section_sources)} sources."
                if section_research_status in {"completed", "local_only", "evidence_bank"}
                else f"{section_title} finished with status '{section_research_status}'. Using partial research output."
            ),
            command=section_objective,
            status="completed" if section_research_status in {"completed", "local_only", "evidence_bank"} else "failed",
            started_at=section_research_started_at,
            completed_at=_trace_now(),
            metadata={
                "phase": "section_research",
                "section_index": index,
                "section_title": section_title,
                "response_status": section_research_status,
                "sources": len(section_sources),
                "elapsed_seconds": int(research_pass.get("elapsed_seconds", 0) or 0),
                "search_query": search_query if use_section_search else "",
                "urls": _trace_url_list([str(item.get("url", "")).strip() for item in section_sources]),
            },
            subtask=f"Gather evidence for {section_title}",
        )

        write_text_file(
            _artifact_file(artifact_dir, f"section_{index:02d}/research.json"),
            json.dumps(research_pass, indent=2, ensure_ascii=False),
        )
        write_text_file(_artifact_file(artifact_dir, f"section_{index:02d}/sources.json"), json.dumps(section_sources, indent=2, ensure_ascii=False))
        write_text_file(_artifact_file(artifact_dir, f"section_{index:02d}/sources.md"), _references_markdown(section_sources))
        log_task_update(
            DEEP_RESEARCH_LABEL,
            (
                f"Saved section {index}/{total_sections} research bundle: {section_title} "
                f"({len(section_sources)} sources)."
            ),
        )
        section_packages.append(
            {
                "index": index,
                "title": section_title,
                "objective": section_objective,
                "key_questions": section_questions,
                "target_pages": target_section_pages,
                "research_pass": research_pass,
                "research_text": research_output,
                "sources": section_sources,
                "source_ledger_md": source_ledger_md,
            }
        )
        research_log_lines.append(f"Researched section {index}: {section_title} ({len(section_sources)} sources)")

    if checkpoint_enabled:
        write_text_file(
            _artifact_file(artifact_dir, "checkpoint_after_research.json"),
            json.dumps({"phase": "research", "sections": section_packages, "analysis": analysis}, indent=2, ensure_ascii=False),
        )

    log_task_update(DEEP_RESEARCH_LABEL, "Phase 2/4 - correlating subtopics and section dependencies.")
    _raise_if_cancelled(state, phase="correlation")
    correlation_started_at = _trace_now()
    _trace_research_event(
        state,
        title="Correlating sections",
        detail="Mapping cross-cutting themes, contradictions, and the final section order.",
        status="running",
        started_at=correlation_started_at,
        metadata={"phase": "correlation"},
        subtask="Correlate sections and dependencies",
    )
    correlation = _build_correlation_package(objective, section_packages)
    write_text_file(_artifact_file(artifact_dir, "correlation_briefing.md"), str(correlation.get("briefing", "")).strip() + "\n")
    write_text_file(_artifact_file(artifact_dir, "knowledge_graph.json"), json.dumps(correlation.get("knowledge_graph", {}), indent=2, ensure_ascii=False))
    research_log_lines.append(
        f"Correlation complete: {len(correlation.get('cross_cutting_themes', []))} themes, "
        f"{len(correlation.get('contradictions', []))} contradictions."
    )
    _trace_research_event(
        state,
        title="Correlating sections",
        detail=(
            f"Correlation complete with {len(correlation.get('cross_cutting_themes', []))} themes and "
            f"{len(correlation.get('contradictions', []))} contradictions."
        ),
        status="completed",
        started_at=correlation_started_at,
        completed_at=_trace_now(),
        metadata={
            "phase": "correlation",
            "theme_count": len(correlation.get("cross_cutting_themes", [])),
            "contradiction_count": len(correlation.get("contradictions", [])),
        },
        subtask="Correlate sections and dependencies",
    )

    packages_by_title = {str(item.get("title", "")).strip(): item for item in section_packages}
    ordered_packages = [packages_by_title[title_key] for title_key in correlation.get("section_order", []) if title_key in packages_by_title]
    if len(ordered_packages) != len(section_packages):
        ordered_packages = section_packages

    for write_index, package in enumerate(ordered_packages, start=1):
        _raise_if_cancelled(state, phase=f"section_{write_index}_draft")
        section_title = str(package.get("title", f"Section {write_index}")).strip() or f"Section {write_index}"
        words_target = max(500, int(package.get("target_pages", section_pages) or section_pages) * words_per_page)
        cached_draft = {}
        if resume_reuse_enabled:
            cached_draft = _load_cached_section_draft(
                artifact_dir=artifact_dir,
                section_index=write_index,
                section_title=section_title,
            )
        if cached_draft:
            log_task_update(
                DEEP_RESEARCH_LABEL,
                f"Phase 3/4 - reusing cached draft for section {write_index}/{len(ordered_packages)}: {section_title}",
            )
            section_text = str(cached_draft.get("section_text", "")).strip()
            note = str(cached_draft.get("continuity_note", "")).strip() or _fallback_continuity_note(section_text)
            visual_assets = cached_draft.get("visual_assets", {})
            if not isinstance(visual_assets, dict):
                visual_assets = {"tables": [], "flowcharts": [], "notes": ""}
            flowchart_files = cached_draft.get("flowchart_files", [])
            if not isinstance(flowchart_files, list):
                flowchart_files = []
            section_images = cached_draft.get("section_images", [])
            if not isinstance(section_images, list):
                section_images = []
            all_section_images.extend([item for item in section_images if isinstance(item, dict)])

            section_md_filename = _artifact_file(artifact_dir, f"section_{write_index:02d}/section.md")
            visual_assets_json_filename = _artifact_file(artifact_dir, f"section_{write_index:02d}/visual_assets.json")
            visual_assets_md_filename = _artifact_file(artifact_dir, f"section_{write_index:02d}/visual_assets.md")
            write_text_file(section_md_filename, section_text)
            write_text_file(visual_assets_json_filename, json.dumps(visual_assets, indent=2, ensure_ascii=False))
            write_text_file(visual_assets_md_filename, _render_visual_assets_md(visual_assets) + "\n")

            continuity_notes.append(f"{section_title}:\n{note}")
            write_text_file(_artifact_file(artifact_dir, f"section_{write_index:02d}/continuity.txt"), note)
            bridge_md = (
                f"# Section {write_index} Coherence Bridge\n\n"
                f"## Section\n{section_title}\n\n"
                "## Carry-Forward Notes\n"
                f"{note.strip()}\n"
            )
            write_text_file(_artifact_file(artifact_dir, f"section_{write_index:02d}/bridge.md"), bridge_md)
            coherence_context_md = coherence_context_md + "\n\n" + f"[Section {write_index} Bridge]\n" + _trim_text(bridge_md, 2000)
            write_text_file(_artifact_file(artifact_dir, "deep_research_coherence_live.md"), coherence_context_md)

            section_outputs.append(
                {
                    "index": write_index,
                    "title": section_title,
                    "objective": package.get("objective", objective),
                    "research_status": str((package.get("research_pass") or {}).get("status", "")),
                    "research_response_id": str((package.get("research_pass") or {}).get("response_id", "")),
                    "research_elapsed_seconds": int((package.get("research_pass") or {}).get("elapsed_seconds", 0) or 0),
                    "research_text_preview": str(package.get("research_text", ""))[:2000],
                    "section_text": section_text,
                    "continuity_note": note,
                    "references": package.get("sources", []),
                    "visual_assets": visual_assets,
                    "visual_assets_file": _output_file_path(visual_assets_json_filename),
                    "visual_assets_markdown_file": _output_file_path(visual_assets_md_filename),
                    "flowchart_files": flowchart_files,
                    "section_images": section_images,
                }
            )
            research_log_lines.append(f"Reused drafted section {write_index}: {section_title}")
            state["draft_response"] = (
                f"Deep research in progress: reused drafted section {write_index}/{len(ordered_packages)} "
                f"({section_title}). Sources this section: {len(package.get('sources', []))}."
            )
            progress_payload = {
                "title": title,
                "objective": objective,
                "target_pages": target_pages,
                "completed_sections": write_index,
                "total_sections": len(ordered_packages),
                "tier": analysis.get("tier", 0),
                "phase": "writing",
                "sections": [
                    {
                        "index": item["index"],
                        "title": item["title"],
                        "research_status": item["research_status"],
                        "references_count": len(item.get("references", [])),
                        "table_count": len((item.get("visual_assets") or {}).get("tables", [])),
                        "flowchart_count": len((item.get("visual_assets") or {}).get("flowcharts", [])),
                    }
                    for item in section_outputs
                ],
            }
            write_text_file(_artifact_file(artifact_dir, "deep_research_progress.json"), json.dumps(progress_payload, indent=2, ensure_ascii=False))
            _trace_research_event(
                state,
                title=f"Drafting section {write_index}/{len(ordered_packages)}",
                detail=f"{section_title} reused from resume artifacts.",
                status="completed",
                metadata={
                    "phase": "section_drafting",
                    "section_index": write_index,
                    "section_title": section_title,
                    "resumed_from_cache": True,
                    "source_count": len(package.get("sources", [])),
                    "table_count": len((visual_assets or {}).get("tables", [])),
                    "flowchart_count": len((visual_assets or {}).get("flowcharts", [])),
                },
                subtask=f"Draft {section_title}",
            )
            continue
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"Phase 3/4 - drafting section {write_index}/{len(ordered_packages)}: {section_title}",
        )
        section_draft_started_at = _trace_now()
        _trace_research_event(
            state,
            title=f"Drafting section {write_index}/{len(ordered_packages)}",
            detail=section_title,
            status="running",
            started_at=section_draft_started_at,
            metadata={"phase": "section_drafting", "section_index": write_index, "section_title": section_title},
            subtask=f"Draft {section_title}",
        )
        section_prompt = _format_section_prompt(
            global_objective=objective,
            section=package,
            section_index=write_index,
            section_count=len(ordered_packages),
            words_target=words_target,
            continuity_notes=continuity_notes,
            coherence_context_md=coherence_context_md,
            correlation_briefing=str(correlation.get("briefing", "")).strip(),
            source_ledger_md=package.get("source_ledger_md", ""),
            research_text=package.get("research_text", ""),
            evidence_bank_md=evidence_excerpt,
            citation_style=citation_style,
            date_range=date_range,
        )
        section_text = llm_text(section_prompt).strip()
        if not section_text:
            section_text = "No section text was generated."
        section_text = _strip_leading_section_heading(section_text, section_title=section_title, section_index=write_index)
        section_text = _ensure_section_citations(section_text, package.get("sources", []))
        existing_tables = _extract_markdown_tables(section_text)
        existing_flowcharts = _extract_mermaid_blocks(section_text)
        generated_visuals: dict[str, Any] = {"tables": [], "flowcharts": [], "notes": ""}
        if disable_visuals:
            log_task_update(DEEP_RESEARCH_LABEL, f"Skipping visual generation for section {write_index} (disabled).")
        elif not (existing_tables and existing_flowcharts):
            generated_visuals = _generate_visual_assets(section_title, section_text, package.get("research_text", ""))
        visual_assets = _normalize_visual_assets(existing_tables, existing_flowcharts, generated_visuals)
        # Replace inline Mermaid blocks in section text with rendered PNGs
        section_text, inline_png_files = _replace_mermaid_with_png(section_text, artifact_dir, write_index)

        # Render visual_assets flowcharts to PNG and store the relative path on the chart dict
        flowchart_files: list[str] = []
        for chart_index, chart in enumerate(visual_assets.get("flowcharts", []), start=1):
            flowchart_text = str(chart.get("mermaid", "")).strip()
            if not flowchart_text:
                continue
            mmd_filename = _artifact_file(artifact_dir, f"section_{write_index:02d}/flowchart_va_{chart_index:02d}.mmd")
            write_text_file(mmd_filename, flowchart_text + "\n")
            png_filename = _artifact_file(artifact_dir, f"section_{write_index:02d}/flowchart_va_{chart_index:02d}.png")
            png_abs = resolve_output_path(png_filename)
            os.makedirs(os.path.dirname(png_abs), exist_ok=True)
            if _render_mermaid_png(flowchart_text, png_abs):
                chart["png_path"] = f"section_{write_index:02d}/flowchart_va_{chart_index:02d}.png"
                flowchart_files.append(png_abs)
            else:
                flowchart_files.append(resolve_output_path(mmd_filename))
        flowchart_files.extend(inline_png_files)

        section_text = _append_generated_visuals(section_text, visual_assets)

        # --- Image collection (web-sourced diagrams / charts / visuals) ---
        section_images: list[dict] = []
        if image_collection_enabled:
            _trace_research_event(
                state,
                title=f"Collecting images for section {write_index}/{len(ordered_packages)}",
                detail=f"Searching for diagrams and visuals to illustrate '{section_title}'.",
                status="running",
                metadata={"phase": "image_collection", "section_index": write_index, "section_title": section_title},
                subtask=f"Collect images for {section_title}",
            )
            try:
                section_images = _collect_section_images(
                    section_title=section_title,
                    section_objective=str(package.get("objective", objective)),
                    section_text=section_text,
                    section_sources=package.get("sources", []) if isinstance(package.get("sources"), list) else [],
                    artifact_dir=artifact_dir,
                    section_index=write_index,
                    max_images=int(state.get("long_document_max_images_per_section", 3)),
                    search_num=int(state.get("long_document_image_search_num", 12)),
                    fallback_to_search=image_search_fallback_enabled,
                    verify_with_vision=image_vision_verification_enabled,
                )
                all_section_images.extend(section_images)
                _trace_research_event(
                    state,
                    title=f"Images collected for section {write_index}",
                    detail=f"{len(section_images)} image(s) downloaded for '{section_title}'.",
                    status="completed",
                    metadata={
                        "phase": "image_collection",
                        "section_index": write_index,
                        "section_title": section_title,
                        "image_count": len(section_images),
                        "images": [e.get("filename", "") for e in section_images],
                    },
                    subtask=f"Collect images for {section_title}",
                )
            except Exception as _img_exc:
                log_task_update(DEEP_RESEARCH_LABEL, f"Image collection failed for section {write_index}: {_img_exc}")
        if section_images:
            section_text = _embed_images_in_section(section_text, section_images)

        if include_section_references:
            section_text = _append_verified_references(section_text, package.get("sources", []))
        else:
            section_text = _strip_section_references(section_text)

        section_dir = _artifact_file(artifact_dir, f"section_{write_index:02d}")
        section_md_filename = _artifact_file(artifact_dir, f"section_{write_index:02d}/section.md")
        visual_assets_json_filename = _artifact_file(artifact_dir, f"section_{write_index:02d}/visual_assets.json")
        visual_assets_md_filename = _artifact_file(artifact_dir, f"section_{write_index:02d}/visual_assets.md")
        section_metadata_filename = _artifact_file(artifact_dir, f"section_{write_index:02d}/section_metadata.json")
        write_text_file(section_md_filename, section_text)
        write_text_file(visual_assets_json_filename, json.dumps(visual_assets, indent=2, ensure_ascii=False))
        write_text_file(visual_assets_md_filename, _render_visual_assets_md(visual_assets) + "\n")
        write_text_file(
            section_metadata_filename,
            json.dumps(
                {
                    "section_index": write_index,
                    "section_title": section_title,
                    "section_objective": package.get("objective", objective),
                    "flowchart_files": flowchart_files,
                    "section_images": section_images,
                },
                indent=2,
                ensure_ascii=False,
            ),
        )
        log_task_update(
            DEEP_RESEARCH_LABEL,
            (
                f"Saved section {write_index}/{len(ordered_packages)} artifacts: {section_title} "
                f"→ {_output_file_path(section_md_filename)}"
            ),
        )

        note = _section_continuity_note(section_title, section_text)
        continuity_notes.append(f"{section_title}:\n{note}")
        write_text_file(_artifact_file(artifact_dir, f"section_{write_index:02d}/continuity.txt"), note)
        bridge_md = (
            f"# Section {write_index} Coherence Bridge\n\n"
            f"## Section\n{section_title}\n\n"
            "## Carry-Forward Notes\n"
            f"{note.strip()}\n"
        )
        write_text_file(_artifact_file(artifact_dir, f"section_{write_index:02d}/bridge.md"), bridge_md)
        coherence_context_md = coherence_context_md + "\n\n" + f"[Section {write_index} Bridge]\n" + _trim_text(bridge_md, 2000)
        write_text_file(_artifact_file(artifact_dir, "deep_research_coherence_live.md"), coherence_context_md)

        section_outputs.append(
            {
                "index": write_index,
                "title": section_title,
                "objective": package.get("objective", objective),
                "research_status": str((package.get("research_pass") or {}).get("status", "")),
                "research_response_id": str((package.get("research_pass") or {}).get("response_id", "")),
                "research_elapsed_seconds": int((package.get("research_pass") or {}).get("elapsed_seconds", 0) or 0),
                "research_text_preview": str(package.get("research_text", ""))[:2000],
                "section_text": section_text,
                "continuity_note": note,
                "references": package.get("sources", []),
                "visual_assets": visual_assets,
                "visual_assets_file": _output_file_path(visual_assets_json_filename),
                "visual_assets_markdown_file": _output_file_path(visual_assets_md_filename),
                "flowchart_files": flowchart_files,
                "section_images": section_images,
            }
        )
        research_log_lines.append(
            f"Drafted section {write_index}: {section_title}"
            + (f" ({len(section_images)} images)" if section_images else "")
        )
        state["draft_response"] = (
            f"Deep research in progress: drafted section {write_index}/{len(ordered_packages)} "
            f"({section_title}). Sources this section: {len(package.get('sources', []))}."
        )
        progress_payload = {
            "title": title,
            "objective": objective,
            "target_pages": target_pages,
            "completed_sections": write_index,
            "total_sections": len(ordered_packages),
            "tier": analysis.get("tier", 0),
            "phase": "writing",
            "sections": [
                {
                    "index": item["index"],
                    "title": item["title"],
                    "research_status": item["research_status"],
                    "references_count": len(item.get("references", [])),
                    "table_count": len((item.get("visual_assets") or {}).get("tables", [])),
                    "flowchart_count": len((item.get("visual_assets") or {}).get("flowcharts", [])),
                }
                for item in section_outputs
            ],
        }
        write_text_file(_artifact_file(artifact_dir, "deep_research_progress.json"), json.dumps(progress_payload, indent=2, ensure_ascii=False))
        _trace_research_event(
            state,
            title=f"Drafting section {write_index}/{len(ordered_packages)}",
            detail=(
                f"{section_title} drafted with {len(package.get('sources', []))} sources, "
                f"{len((visual_assets or {}).get('tables', []))} tables, and "
                f"{len((visual_assets or {}).get('flowcharts', []))} flowcharts."
            ),
            status="completed",
            started_at=section_draft_started_at,
            completed_at=_trace_now(),
            metadata={
                "phase": "section_drafting",
                "section_index": write_index,
                "section_title": section_title,
                "source_count": len(package.get("sources", [])),
                "table_count": len((visual_assets or {}).get("tables", [])),
                "flowchart_count": len((visual_assets or {}).get("flowcharts", [])),
            },
            subtask=f"Draft {section_title}",
        )

    if checkpoint_enabled:
        write_text_file(
            _artifact_file(artifact_dir, "checkpoint_after_writing.json"),
            json.dumps({"phase": "writing", "sections": section_outputs}, indent=2, ensure_ascii=False),
        )

    log_task_update(DEEP_RESEARCH_LABEL, f"Phase 3/4 - generating executive summary from {len(section_outputs)} sections.")
    _trace_research_event(
        state,
        title="Generating executive summary",
        detail=f"Synthesizing executive summary from {len(section_outputs)} drafted sections.",
        status="running",
        metadata={"phase": "executive_summary", "section_count": len(section_outputs)},
        subtask="Generate executive summary",
    )
    summary_prompt = f"""
Create a concise executive summary for this deep research report.
Objective:
{objective}

Section continuity notes:
{json.dumps(continuity_notes, indent=2, ensure_ascii=False)}
"""
    executive_summary = llm_text(summary_prompt).strip()
    if not executive_summary:
        executive_summary = "Executive summary was not generated."
        log_task_update(DEEP_RESEARCH_LABEL, "WARNING: Executive summary was empty — using placeholder.")
    else:
        log_task_update(DEEP_RESEARCH_LABEL, f"Executive summary generated: {len(executive_summary.split())} words.")
    _trace_research_event(
        state,
        title="Generating executive summary",
        detail=f"Executive summary complete ({len(executive_summary.split())} words)." if executive_summary != "Executive summary was not generated." else "Executive summary generation returned empty — using placeholder.",
        status="completed" if executive_summary != "Executive summary was not generated." else "failed",
        metadata={"phase": "executive_summary", "word_count": len(executive_summary.split())},
        subtask="Generate executive summary",
    )

    log_task_update(DEEP_RESEARCH_LABEL, "Consolidating references and remapping citations across all sections.")
    consolidated_references = _consolidate_references(section_outputs)
    if not include_section_references:
        for item in section_outputs:
            item["section_text"] = _remap_section_citations(
                item.get("section_text", ""),
                item.get("references", []),
                consolidated_references,
            )
    log_task_update(DEEP_RESEARCH_LABEL, f"Consolidated {len(consolidated_references)} references across {len(section_outputs)} sections.")
    _trace_research_event(
        state,
        title="References consolidated",
        detail=f"Merged {len(consolidated_references)} unique references from {len(section_outputs)} sections using {citation_style.upper()} style.",
        status="completed",
        metadata={"phase": "references", "reference_count": len(consolidated_references), "citation_style": citation_style},
        subtask="Consolidate references",
    )
    references_md_filename = _artifact_file(artifact_dir, "deep_research_references.md")
    references_json_filename = _artifact_file(artifact_dir, "deep_research_references.json")
    write_text_file(references_md_filename, _bibliography_markdown(consolidated_references, style=citation_style))
    write_text_file(references_json_filename, json.dumps(consolidated_references, indent=2, ensure_ascii=False))
    visual_index = _build_visual_index(section_outputs)
    total_tables_vi = sum(len(item.get("tables", [])) for item in visual_index.get("sections", []))
    total_flowcharts_vi = sum(len(item.get("flowcharts", [])) for item in visual_index.get("sections", []))
    log_task_update(DEEP_RESEARCH_LABEL, f"Visual index built: {total_tables_vi} tables, {total_flowcharts_vi} flowcharts.")
    visual_index_json_filename = _artifact_file(artifact_dir, "deep_research_visual_index.json")
    visual_index_md_filename = _artifact_file(artifact_dir, "deep_research_visual_index.md")
    write_text_file(visual_index_json_filename, json.dumps(visual_index, indent=2, ensure_ascii=False))
    write_text_file(visual_index_md_filename, _visual_index_markdown(visual_index))

    methodology_lines = [
        f"- Research tier: {analysis.get('tier', 0)}",
        f"- Requested page target: {target_pages}",
        f"- Citation style: {citation_style.upper()}",
        f"- Web search: {'enabled' if web_search_enabled else 'disabled'}",
        f"- Date range: {date_range}",
        f"- Source families requested: {', '.join(source_family_display)}",
        f"- Local file sources: {len(local_entries)}",
        f"- Explicit URL sources: {len(url_entries)}",
        f"- Evidence bank pre-collection: {'enabled' if collect_sources_first else 'disabled'}",
        f"- Section search support: {'enabled' if use_section_search else 'disabled'}",
        f"- Cross-cutting themes identified: {len(correlation.get('cross_cutting_themes', []))}",
        f"- Contradictions identified: {len(correlation.get('contradictions', []))}",
    ]
    methodology_text = "\n".join(methodology_lines)

    if plagiarism_enabled:
        log_task_update(DEEP_RESEARCH_LABEL, f"Phase 3/4 - running plagiarism check across {len(section_outputs)} sections.")
        _trace_research_event(
            state,
            title="Running plagiarism check",
            detail=f"Checking {len(section_outputs)} sections for similarity and AI-generated content.",
            status="running",
            metadata={"phase": "plagiarism", "section_count": len(section_outputs)},
            subtask="Run plagiarism check",
        )
    else:
        log_task_update(DEEP_RESEARCH_LABEL, "Plagiarism check disabled — skipping.")
    plagiarism_sources = _build_plagiarism_source_texts(
        evidence_bank_md=evidence_bank_md,
        section_packages=section_packages,
        local_entries=local_entries,
        url_entries=url_entries,
    )
    plagiarism_report = (
        _build_plagiarism_report(section_outputs, plagiarism_sources)
        if plagiarism_enabled
        else {"overall_score": 0.0, "ai_content_score": 0.0, "status": "PASS", "sections": []}
    )
    if plagiarism_enabled:
        plag_score = plagiarism_report.get("overall_score", 0)
        plag_status = plagiarism_report.get("status", "PASS")
        plag_ai = plagiarism_report.get("ai_content_score", 0)
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"Plagiarism check complete: overall={plag_score}% ({plag_status}), AI content={plag_ai}%.",
        )
        _trace_research_event(
            state,
            title="Plagiarism check complete",
            detail=f"Overall similarity: {plag_score}% ({plag_status}). AI content score: {plag_ai}%.",
            status="completed",
            metadata={"phase": "plagiarism", "overall_score": plag_score, "status": plag_status, "ai_content_score": plag_ai},
            subtask="Run plagiarism check",
        )
    plagiarism_json_filename = _artifact_file(artifact_dir, "plagiarism_report.json")
    plagiarism_md_filename = _artifact_file(artifact_dir, "plagiarism_report.md")
    write_text_file(plagiarism_json_filename, json.dumps(plagiarism_report, indent=2, ensure_ascii=False))
    write_text_file(plagiarism_md_filename, _plagiarism_report_markdown(plagiarism_report))

    log_task_update(DEEP_RESEARCH_LABEL, f"Phase 4/4 - assembling compiled markdown from {len(section_outputs)} sections.")
    compile_started_at = _trace_now()
    _trace_research_event(
        state,
        title="Compiling final report",
        detail=f"Merging {len(section_outputs)} sections with executive summary, references, plagiarism appendix, and methodology.",
        status="running",
        started_at=compile_started_at,
        metadata={"phase": "compile", "section_count": len(section_outputs), "formats": output_formats},
        subtask="Compile and export final report",
    )
    compiled_markdown = _build_compiled_markdown(
        title,
        objective,
        section_outputs,
        executive_summary,
        consolidated_references,
        citation_style=citation_style,
        methodology_text=methodology_text,
        plagiarism_report=plagiarism_report,
        source_entries=consolidated_references,
        research_log_lines=research_log_lines,
        generated_at=dt.datetime.now(dt.UTC).strftime("%Y-%m-%d %H:%M UTC"),
        model_name=research_model,
        deep_research_tier=int(analysis.get("tier", 0) or 0),
        image_entries=all_section_images,
    )
    compiled_filename = _artifact_file(artifact_dir, "deep_research_report.md")
    write_text_file(compiled_filename, compiled_markdown)
    compiled_path = _output_file_path(compiled_filename)
    compiled_chars = len(compiled_markdown)
    compiled_words = len(compiled_markdown.split())
    log_task_update(
        DEEP_RESEARCH_LABEL,
        f"Compiled markdown written: {compiled_words} words, {compiled_chars} chars → {compiled_path}",
    )
    _trace_research_event(
        state,
        title="Compiled markdown ready",
        detail=f"Full report assembled: {compiled_words} words across {len(section_outputs)} sections. Writing export formats: {', '.join(output_formats)}.",
        status="running",
        started_at=compile_started_at,
        metadata={"phase": "compile_markdown", "word_count": compiled_words, "char_count": compiled_chars},
        subtask="Compile and export final report",
    )
    export_paths: dict[str, str] = {}
    export_errors: dict[str, str] = {}
    try:
        log_task_update(DEEP_RESEARCH_LABEL, f"Exporting document formats: {', '.join(output_formats)}.")
        export_paths = _export_long_document_formats(compiled_markdown, compiled_filename, requested_formats=output_formats)
        export_errors = dict(export_paths.pop("errors", {}) or {})
        succeeded_formats = [fmt for fmt in ("html", "docx", "pdf", "md") if export_paths.get(fmt)]
        failed_formats = list(export_errors.keys())
        if succeeded_formats:
            log_task_update(DEEP_RESEARCH_LABEL, f"Document export succeeded for: {', '.join(succeeded_formats).upper()}.")
            for fmt in succeeded_formats:
                log_task_update(DEEP_RESEARCH_LABEL, f"  {fmt.upper()} → {export_paths[fmt]}")
        if export_errors:
            err_summary = "; ".join(f"{fmt.upper()}: {msg}" for fmt, msg in export_errors.items())
            log_task_update(DEEP_RESEARCH_LABEL, f"Some format exports failed: {err_summary}")
            _trace_research_event(
                state,
                title="Document export — partial failure",
                detail=f"Succeeded: {', '.join(succeeded_formats) or 'none'}. Failed: {err_summary}",
                status="failed",
                metadata={"phase": "export", "succeeded": succeeded_formats, "failed": failed_formats, "errors": export_errors},
                subtask="Compile and export final report",
            )
        else:
            _trace_research_event(
                state,
                title="Document export complete",
                detail=f"All requested formats exported successfully: {', '.join(succeeded_formats).upper() if succeeded_formats else 'none'}.",
                status="completed",
                metadata={"phase": "export", "succeeded": succeeded_formats, "paths": {fmt: export_paths.get(fmt, "") for fmt in succeeded_formats}},
                subtask="Compile and export final report",
            )
    except Exception as exc:
        log_task_update(DEEP_RESEARCH_LABEL, f"Format export failed with unexpected exception: {exc}")
        _trace_research_event(
            state,
            title="Document export failed",
            detail=f"Unexpected error during format export: {exc}",
            status="failed",
            metadata={"phase": "export", "error": str(exc)},
            subtask="Compile and export final report",
        )
        export_errors = {"all": str(exc)}
    local_manifest = state.get("local_drive_manifest", {}) if isinstance(state.get("local_drive_manifest"), dict) else {}
    evidence_ledger = _build_evidence_ledger(
        consolidated_references=consolidated_references,
        section_outputs=section_outputs,
        local_entries=local_entries,
        url_entries=url_entries,
    )
    coverage_report = _build_coverage_report(
        objective=objective,
        intent=intent,
        source_strategy=source_strategy,
        local_manifest=local_manifest,
        local_entries=local_entries,
        url_entries=url_entries,
        kb_grounding=kb_grounding,
        kb_warning=kb_warning,
        evidence_sources=evidence_sources,
        consolidated_references=consolidated_references,
    )
    quality_report = _build_quality_report(
        section_outputs=section_outputs,
        evidence_ledger=evidence_ledger,
        coverage_report=coverage_report,
        consolidated_references=consolidated_references,
    )
    source_ledger_filename = _artifact_file(artifact_dir, "source_ledger.json")
    coverage_report_filename = _artifact_file(artifact_dir, "coverage_report.json")
    quality_report_filename = _artifact_file(artifact_dir, "quality_report.json")
    write_text_file(source_ledger_filename, json.dumps(evidence_ledger, indent=2, ensure_ascii=False))
    write_text_file(coverage_report_filename, json.dumps(coverage_report, indent=2, ensure_ascii=False))
    write_text_file(quality_report_filename, json.dumps(quality_report, indent=2, ensure_ascii=False))
    _trace_research_event(
        state,
        title="Coverage report updated",
        detail=f"Coverage status: {coverage_report.get('status', 'unknown')} with {len(coverage_report.get('gaps', []) or [])} gap(s).",
        status="completed",
        kind="coverage",
        metadata={"phase": "coverage", "coverage_report": coverage_report},
        subtask="Update coverage report",
    )
    _trace_research_event(
        state,
        title="Research quality gate evaluated",
        detail=str(quality_report.get("summary", "")).strip(),
        status="completed" if quality_report.get("status") == "pass" else "failed" if quality_report.get("status") == "fail" else "pending",
        kind="quality_gate",
        metadata={"phase": "quality_gate", "quality_report": quality_report},
        subtask="Run research quality gate",
    )
    for gap in coverage_report.get("gaps", []) or []:
        _trace_research_event(
            state,
            title="Coverage gap detected",
            detail=str(gap).strip(),
            status="running",
            kind="gap_detected",
            metadata={"phase": "coverage_gap", "gap": gap},
            subtask="Review research coverage gaps",
        )
    report_alias_md = _output_file_path(_artifact_file(artifact_dir, "report.md"))
    report_alias_html = _output_file_path(_artifact_file(artifact_dir, "report.html"))
    report_alias_pdf = _output_file_path(_artifact_file(artifact_dir, "report.pdf"))
    report_alias_docx = _output_file_path(_artifact_file(artifact_dir, "report.docx"))
    nested_alias_results = {
        "md": _copy_artifact_alias(compiled_path, report_alias_md),
        "html": _copy_artifact_alias(export_paths.get("html", ""), report_alias_html),
        "pdf": _copy_artifact_alias(export_paths.get("pdf", ""), report_alias_pdf),
        "docx": _copy_artifact_alias(export_paths.get("docx", ""), report_alias_docx),
    }
    root_report_results = _mirror_root_report_artifacts(compiled_path, export_paths, root_report_dir="reports")
    if root_report_results:
        mirrored_formats = ", ".join(fmt.upper() for fmt in ("pdf", "docx", "html", "md") if root_report_results.get(fmt))
        log_task_update(DEEP_RESEARCH_LABEL, f"Mirrored stable report downloads to run root reports/: {mirrored_formats}.")
    report_results = {
        "md": root_report_results.get("md", "") or nested_alias_results.get("md", "") or compiled_path,
        "html": root_report_results.get("html", "") or nested_alias_results.get("html", "") or export_paths.get("html", ""),
        "pdf": root_report_results.get("pdf", "") or nested_alias_results.get("pdf", "") or export_paths.get("pdf", ""),
        "docx": root_report_results.get("docx", "") or nested_alias_results.get("docx", "") or export_paths.get("docx", ""),
    }
    created_artifacts = [
        {"name": "report.md", "path": report_results.get("md", "") or compiled_path, "kind": "report"},
        {"name": "source_ledger.json", "path": _output_file_path(source_ledger_filename), "kind": "evidence"},
        {"name": "coverage_report.json", "path": _output_file_path(coverage_report_filename), "kind": "coverage"},
        {"name": "quality_report.json", "path": _output_file_path(quality_report_filename), "kind": "quality"},
        {"name": "evidence_bank.md", "path": state.get("long_document_evidence_bank_path", ""), "kind": "evidence"},
        {"name": "evidence_bank.json", "path": state.get("long_document_evidence_bank_json_path", ""), "kind": "evidence"},
    ]
    if str(state.get("long_document_source_manifest_path", "")).strip():
        created_artifacts.append(
            {"name": "source_manifest.md", "path": state.get("long_document_source_manifest_path", ""), "kind": "manifest"}
        )
    if str(state.get("long_document_source_manifest_json_path", "")).strip():
        created_artifacts.append(
            {"name": "source_manifest.json", "path": state.get("long_document_source_manifest_json_path", ""), "kind": "manifest"}
        )
    if report_results.get("pdf"):
        created_artifacts.append({"name": "report.pdf", "path": report_results["pdf"], "kind": "report"})
    if report_results.get("docx"):
        created_artifacts.append({"name": "report.docx", "path": report_results["docx"], "kind": "report"})
    if report_results.get("html"):
        created_artifacts.append({"name": "report.html", "path": report_results["html"], "kind": "report"})
    for artifact in created_artifacts:
        if not str(artifact.get("path", "")).strip():
            continue
        _trace_research_event(
            state,
            title=f"Artifact created: {artifact.get('name', '')}",
            detail=str(artifact.get("path", "")).strip(),
            status="completed",
            kind="artifact_created",
            metadata={"phase": "artifact", "artifact": artifact},
            subtask="Write research artifact",
        )
    total_images = len(all_section_images)
    manifest_filename = _artifact_file(artifact_dir, "deep_research_manifest.json")
    manifest_path = _output_file_path(manifest_filename)
    outline_json_path = _output_file_path(_artifact_file(artifact_dir, "deep_research_outline.json"))
    outline_md_path = _output_file_path(_artifact_file(artifact_dir, "deep_research_outline.md"))
    coherence_base_path = _output_file_path(_artifact_file(artifact_dir, "deep_research_coherence_base.md"))
    coherence_live_path = _output_file_path(_artifact_file(artifact_dir, "deep_research_coherence_live.md"))
    correlation_briefing_path = _output_file_path(_artifact_file(artifact_dir, "correlation_briefing.md"))
    knowledge_graph_path = _output_file_path(_artifact_file(artifact_dir, "knowledge_graph.json"))
    references_md_path = _output_file_path(references_md_filename)
    references_json_path = _output_file_path(references_json_filename)
    visual_index_md_path = _output_file_path(visual_index_md_filename)
    visual_index_json_path = _output_file_path(visual_index_json_filename)
    plagiarism_json_path = _output_file_path(plagiarism_json_filename)
    plagiarism_md_path = _output_file_path(plagiarism_md_filename)
    source_ledger_path = _output_file_path(source_ledger_filename)
    coverage_report_path = _output_file_path(coverage_report_filename)
    quality_report_path = _output_file_path(quality_report_filename)
    intent_path = _output_file_path(intent_json_path)
    strategy_path = _output_file_path(strategy_json_path)
    write_text_file(
        manifest_filename,
        json.dumps(
            {
                "title": title,
                "objective": objective,
                "tier": analysis.get("tier", 0),
                "target_pages": target_pages,
                "section_count": len(section_outputs),
                "compiled_markdown_file": compiled_path,
                "report_markdown_file": report_results.get("md", "") or compiled_path,
                "report_html_file": report_results.get("html", "") or export_paths.get("html", ""),
                "report_docx_file": report_results.get("docx", "") or export_paths.get("docx", ""),
                "report_pdf_file": report_results.get("pdf", "") or export_paths.get("pdf", ""),
                "outline_file": outline_json_path,
                "outline_markdown_file": outline_md_path,
                "coherence_base_file": coherence_base_path,
                "coherence_live_file": coherence_live_path,
                "correlation_briefing_file": correlation_briefing_path,
                "knowledge_graph_file": knowledge_graph_path,
                "references_markdown_file": references_md_path,
                "references_json_file": references_json_path,
                "visual_index_markdown_file": visual_index_md_path,
                "visual_index_json_file": visual_index_json_path,
                "plagiarism_report_json": plagiarism_json_path,
                "plagiarism_report_markdown": plagiarism_md_path,
                "source_ledger_file": source_ledger_path,
                "coverage_report_file": coverage_report_path,
                "quality_report_file": quality_report_path,
                "compiled_html_file": export_paths.get("html", ""),
                "compiled_docx_file": export_paths.get("docx", ""),
                "compiled_pdf_file": export_paths.get("pdf", ""),
                "evidence_bank_file": state.get("long_document_evidence_bank_path", ""),
                "evidence_bank_json_file": state.get("long_document_evidence_bank_json_path", ""),
                "source_manifest_file": state.get("long_document_source_manifest_path", ""),
                "source_manifest_json_file": state.get("long_document_source_manifest_json_path", ""),
                "intent_file": intent_path,
                "source_strategy_file": strategy_path,
                "created_artifacts": created_artifacts,
                "image_gallery_count": total_images,
                "image_gallery": [
                    {
                        "section_index": e.get("section_index"),
                        "section_title": e.get("section_title", ""),
                        "filename": e.get("filename", ""),
                        "abs_path": e.get("abs_path", ""),
                        "relative_path": e.get("relative_path", ""),
                        "source": e.get("source", ""),
                        "source_page": e.get("source_page", ""),
                        "title": e.get("title", ""),
                        "alt_text": e.get("alt_text", ""),
                    }
                    for e in all_section_images
                ],
            },
            indent=2,
            ensure_ascii=False,
        ),
    )

    total_tables = sum(len((item.get("visual_assets") or {}).get("tables", [])) for item in section_outputs)
    total_flowcharts = sum(len((item.get("visual_assets") or {}).get("flowcharts", [])) for item in section_outputs)
    total_words = sum(len(str(item.get("section_text", "")).split()) for item in section_outputs) + len(str(executive_summary).split())
    total_sources = len(consolidated_references)
    if total_images:
        log_task_update(
            DEEP_RESEARCH_LABEL,
            f"Image collection complete: {total_images} images across {len(section_outputs)} sections.",
        )
        _trace_research_event(
            state,
            title="Image gallery ready",
            detail=f"Collected {total_images} web-sourced image(s) across {len(section_outputs)} sections for Appendix D.",
            status="completed",
            metadata={
                "phase": "image_gallery",
                "total_images": total_images,
                "images_by_section": {
                    str(e.get("section_index", "?")): e.get("filename", "")
                    for e in all_section_images
                },
            },
            subtask="Build image gallery appendix",
        )

    downloadable_report_formats: list[str] = []
    if report_results.get("pdf"):
        downloadable_report_formats.append("PDF")
    if report_results.get("docx"):
        downloadable_report_formats.append("DOCX")
    if report_results.get("html"):
        downloadable_report_formats.append("HTML")
    downloadable_report_formats.append("MD")
    export_issue_line = f"- Export issues: {'; '.join(export_errors)}\n" if export_errors else ""
    artifact_summary_lines = render_artifact_lines(
        [
            ("Artifact bundle", _output_file_path(artifact_dir)),
            ("Markdown report", report_results.get("md", "") or compiled_path),
            ("PDF report", report_results.get("pdf", "")),
            ("DOCX report", report_results.get("docx", "")),
            ("HTML report", report_results.get("html", "")),
            ("Source manifest", state.get("long_document_source_manifest_path", "")),
            ("Research manifest", manifest_path),
        ],
        output_root=OUTPUT_DIR,
    )
    artifact_summary_block = "".join(f"{line}\n" for line in artifact_summary_lines)

    final_summary = (
        f"Deep research pipeline completed.\n"
        f"- Title: {title}\n"
        f"- Tier: {analysis.get('tier', 0)}\n"
        f"- Research depth: {analysis.get('depth_label', research_depth_label)}\n"
        f"- Sections produced: {len(section_outputs)}\n"
        f"- Words: {total_words}\n"
        f"- Sources: {total_sources}\n"
        f"- Web search: {'disabled' if not web_search_enabled else 'enabled via native model web search' if native_web_search_enabled else 'enabled via Kendr search client'}\n"
        f"- Local file sources: {len(local_entries)}\n"
        f"- Explicit URL sources: {len(url_entries)}\n"
        f"- Knowledge base: {str(kb_grounding.get('kb_name', '') or 'disabled') if research_kb_enabled else 'disabled'}\n"
        f"- Knowledge base hits: {int(kb_grounding.get('hit_count', 0) or 0)}\n"
        f"- Citations: {len(consolidated_references)}\n"
        f"- Plagiarism: {plagiarism_report.get('overall_score', 0)}% ({plagiarism_report.get('status', 'PASS')})\n"
        f"- Report downloads in chat: {', '.join(downloadable_report_formats)}\n"
        f"{artifact_summary_block}"
        f"- Supporting research files: source_manifest.md, deep_research_references.md, evidence_bank.md, source_ledger.json, coverage_report.json, quality_report.json\n"
        f"- Visual assets generated: {total_tables} tables, {total_flowcharts} flowcharts\n"
        f"- Web images collected: {total_images} (Appendix D)\n"
        f"{export_issue_line}"
        "\nExecutive summary:\n"
        f"{_trim_text(executive_summary, 1800)}\n"
    )

    state["long_document_mode"] = True
    state["deep_research_mode"] = deep_research_mode
    state["deep_research_confirmed"] = deep_research_mode or bool(state.get("deep_research_confirmed", False))
    state["deep_research_tier"] = int(analysis.get("tier", 0) or 0)
    state["long_document_title"] = title
    state["long_document_outline"] = outline
    state["long_document_sections_data"] = [
        {
            "index": item["index"],
            "title": item["title"],
            "objective": item["objective"],
            "research_status": item["research_status"],
            "research_response_id": item["research_response_id"],
            "research_elapsed_seconds": item["research_elapsed_seconds"],
            "continuity_note": item["continuity_note"],
            "references_count": len(item.get("references", [])),
            "table_count": len((item.get("visual_assets") or {}).get("tables", [])),
            "flowchart_count": len((item.get("visual_assets") or {}).get("flowcharts", [])),
            "visual_assets_file": item.get("visual_assets_file", ""),
            "visual_assets_markdown_file": item.get("visual_assets_markdown_file", ""),
            "flowchart_files": item.get("flowchart_files", []),
        }
        for item in section_outputs
    ]
    state["long_document_artifact_dir"] = _output_file_path(artifact_dir)
    state["long_document_compiled_path"] = compiled_path
    state["long_document_compiled_html_path"] = export_paths.get("html", "")
    state["long_document_compiled_docx_path"] = export_paths.get("docx", "")
    state["long_document_compiled_pdf_path"] = export_paths.get("pdf", "")
    state["deep_research_root_report_dir"] = _output_file_path("reports")
    state["deep_research_root_report_paths"] = dict(report_results)
    state["long_document_manifest_path"] = manifest_path
    state["long_document_source_manifest_path"] = str(state.get("long_document_source_manifest_path", "") or "")
    state["long_document_source_manifest_json_path"] = str(state.get("long_document_source_manifest_json_path", "") or "")
    state["long_document_outline_md_path"] = outline_md_path
    state["long_document_coherence_base_path"] = coherence_base_path
    state["long_document_coherence_live_path"] = coherence_live_path
    state["long_document_references_path"] = references_md_path
    state["long_document_references_json_path"] = references_json_path
    state["long_document_references"] = consolidated_references
    state["long_document_visual_index_path"] = visual_index_md_path
    state["long_document_visual_index_json_path"] = visual_index_json_path
    state["long_document_summary"] = executive_summary
    state["long_document_evidence_sources"] = evidence_sources or explicit_source_entries
    state["research_kb_used"] = bool(research_kb_enabled and int(kb_grounding.get("hit_count", 0) or 0) > 0)
    state["research_kb_name"] = str(kb_grounding.get("kb_name", "") or "")
    state["research_kb_hit_count"] = int(kb_grounding.get("hit_count", 0) or 0)
    state["research_kb_citations"] = list(kb_grounding.get("citations", []) or [])
    state["research_kb_warning"] = kb_warning
    state["deep_research_analysis"] = analysis
    state["deep_research_intent"] = intent
    state["deep_research_source_strategy"] = source_strategy
    state["deep_research_evidence_ledger"] = evidence_ledger
    state["deep_research_coverage_report"] = coverage_report
    state["deep_research_quality_report"] = quality_report
    state["deep_research_artifacts_manifest"] = {
        "artifact_dir": _output_file_path(artifact_dir),
        "manifest_path": manifest_path,
        "root_report_dir": _output_file_path("reports"),
        "root_report_paths": dict(report_results),
        "created_artifacts": created_artifacts,
    }
    state["deep_research_result_card"] = {
        "kind": "result",
        "title": title,
        "tier": analysis.get("tier", 0),
        "depth_mode": analysis.get("depth_mode", research_depth_mode),
        "depth_label": analysis.get("depth_label", research_depth_label),
        "pages": target_pages,
        "words": total_words,
        "sources": total_sources,
        "citations": len(consolidated_references),
        "plagiarism_score": plagiarism_report.get("overall_score", 0),
        "plagiarism_status": plagiarism_report.get("status", "PASS"),
        "ai_content_score": plagiarism_report.get("ai_content_score", 0),
        "duration_minutes": int(sum(int((item.get("research_elapsed_seconds") or 0)) for item in section_outputs) / 60),
        "web_search_enabled": web_search_enabled,
        "local_sources": len(local_entries),
        "provided_urls": len(url_entries),
        "research_kb_used": bool(research_kb_enabled and int(kb_grounding.get("hit_count", 0) or 0) > 0),
        "research_kb_name": str(kb_grounding.get("kb_name", "") or ""),
        "research_kb_hit_count": int(kb_grounding.get("hit_count", 0) or 0),
        "research_kb_citations": list(kb_grounding.get("citations", []) or []),
        "research_kb_warning": kb_warning,
        "intent_summary": intent.get("summary", ""),
        "strategy_summary": source_strategy.get("summary", ""),
        "research_kind": intent.get("research_kind", ""),
        "target_deliverable": intent.get("target_deliverable", ""),
        "source_needs": intent.get("source_needs", []),
        "risk_level": intent.get("risk_level", ""),
        "strategy_mode": source_strategy.get("mode", ""),
        "family_budgets": source_strategy.get("family_budgets", {}),
        "selection_rationale": source_strategy.get("selection_rationale", []),
        "why_selected": dict(source_strategy.get("selection_notes", {}) or {}),
        "why_skipped": dict(source_strategy.get("skip_notes", {}) or {}),
        "formats": output_formats,
        "report_path": compiled_path,
        "html_path": export_paths.get("html", ""),
        "docx_path": export_paths.get("docx", ""),
        "pdf_path": export_paths.get("pdf", ""),
        "report_md_path": report_results.get("md", "") or compiled_path,
        "report_docx_path": report_results.get("docx", "") or export_paths.get("docx", ""),
        "report_pdf_path": report_results.get("pdf", "") or export_paths.get("pdf", ""),
        "report_html_path": report_results.get("html", "") or export_paths.get("html", ""),
        "root_report_dir": _output_file_path("reports"),
        "root_report_paths": dict(report_results),
        "knowledge_graph_path": knowledge_graph_path,
        "plagiarism_report_path": plagiarism_json_path,
        "raw_json_path": manifest_path,
        "source_manifest_path": state.get("long_document_source_manifest_path", ""),
        "source_manifest_json_path": state.get("long_document_source_manifest_json_path", ""),
        "source_ledger_path": source_ledger_path,
        "coverage_report_path": coverage_report_path,
        "quality_report_path": quality_report_path,
        "coverage_status": coverage_report.get("status", ""),
        "coverage_gaps": coverage_report.get("gaps", []),
        "coverage_revisit_plan": coverage_report.get("revisit_plan", []),
        "failed_extractions": coverage_report.get("failed_extractions", []),
        "discovered_files": int(local_manifest.get("file_count", 0) or 0),
        "selected_local_files": int(local_manifest.get("selected_file_count", 0) or 0),
        "quality_status": quality_report.get("status", ""),
        "quality_flags": quality_report.get("flags", []),
        "created_artifacts": created_artifacts,
        "downloadable_reports": [artifact for artifact in created_artifacts if artifact.get("kind") == "report"],
        "export_errors": export_errors,
        "web_search_mode": web_search_mode,
        "image_count": total_images,
        "images": [
            {"section": e.get("section_title", ""), "file": e.get("filename", ""), "path": e.get("abs_path", "")}
            for e in all_section_images
        ],
    }
    state["draft_response"] = final_summary
    _trace_research_event(
        state,
        title="Compiling final report",
        detail=(
            f"Final report ready with {len(section_outputs)} sections, {total_sources} sources, and "
            f"{', '.join(output_formats)} exports."
        ),
        status="completed",
        started_at=compile_started_at,
        completed_at=_trace_now(),
        metadata={
            "phase": "compile",
            "section_count": len(section_outputs),
            "source_count": total_sources,
            "word_count": total_words,
            "research_kb_used": bool(research_kb_enabled and int(kb_grounding.get("hit_count", 0) or 0) > 0),
            "research_kb_name": str(kb_grounding.get("kb_name", "") or ""),
            "research_kb_hit_count": int(kb_grounding.get("hit_count", 0) or 0),
        },
        subtask="Compile and export final report",
    )

    log_task_update(DEEP_RESEARCH_LABEL, f"Completed deep research pass #{call_number}.", final_summary)
    state = publish_agent_output(
        state,
        "long_document_agent",
        final_summary,
        f"deep_research_result_{call_number}",
        recipients=["orchestrator_agent", "reviewer_agent", "report_agent"],
    )
    return state
